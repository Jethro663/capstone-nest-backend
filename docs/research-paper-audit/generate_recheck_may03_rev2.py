from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
AUDIT_DIR = ROOT / "docs" / "research-paper-audit"
EXTRACTED_DIR = AUDIT_DIR / "extracted"
DOCX_OUT = ROOT / "RESEARCH_PAPER_RECHECK_MAY03_REV2.docx"
MD_OUT = ROOT / "paper_recheck_may03_rev2.md"
JSON_OUT = AUDIT_DIR / "paper_recheck_may03_rev2.json"


def load_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


summary = load_json(EXTRACTED_DIR / "summary.json")
now = datetime.now()


improvements = [
    "The 74% mastery threshold correction remains intact across the core intervention sections and Figure 13.",
    "The old copied legacy block about `RND`, `nutritionist`, `food intake`, and `cloud database` is still gone.",
    "The scope is now narrower and more believable: it explicitly says Grades `7`, `8`, `9`, and `10` instead of `all high school grade levels`.",
    "The old offline/local-cache contradiction was cleaned up; the current draft consistently says offline functionality is excluded.",
    "The old duplicate use-case issue is fixed: `Track Student Performance` and `View Evaluations` are now separate table titles.",
    "The paper still honestly notes that teacher/admin mobile workflows are unsupported in the current repository build.",
]


findings = [
    {
        "id": "R2-001",
        "severity": "Major",
        "location": "Chapter 4, Figures 8 and 9",
        "page": "Approx. pp. 102-104",
        "exact": (
            "Figure 8 claims mobile OTP verification and mandatory first-login password update, "
            "while Figure 9 claims mobile forgot-password and account-recovery flows with reset-link handling."
        ),
        "why": (
            "The current `test-mobile` app does not expose OTP verification, password-reset, or account-recovery screens. "
            "Its auth stack contains only a `Login` route, and the login screen literally says `Forgot password? Contact your administrator`."
        ),
        "evidence": [
            "Research paper: docs/research-paper-audit/extracted/full_text.txt:545-564",
            "Repo: test-mobile/src/navigation/types.ts -> AuthStackParamList only contains `Login`",
            "Repo: test-mobile/src/screens/LoginScreen.tsx:431 -> `Forgot password? Contact your administrator`",
            "Repo search: no OTP or mobile reset/account-recovery screens found under test-mobile/src",
        ],
        "correction": (
            "Rewrite Figures 8 and 9 to match the actual mobile login flow, or explicitly move OTP and password-recovery logic to the web-only or future-scope section."
        ),
    },
    {
        "id": "R2-002",
        "severity": "Major",
        "location": "Chapter 3, software stack description",
        "page": "Approx. p. 36",
        "exact": "The frontend is built using Next.js 19.2.3 and React 19",
        "why": (
            "The repository uses Next `^16.2.4`, not Next `19.2.3`. The current paper is still inserting the React version into the Next.js slot."
        ),
        "evidence": [
            "Research paper: docs/research-paper-audit/extracted/full_text.txt:291",
            "Repo: next-frontend/package.json -> `next: ^16.2.4`, `react: 19.2.3`",
        ],
        "correction": "Change this to `Next.js 16.2.4 and React 19.2.3`, or state the major versions only.",
    },
    {
        "id": "R2-003",
        "severity": "Major",
        "location": "Chapter 4, Figure 11",
        "page": "Approx. p. 106",
        "exact": (
            "The system tracks the duration of assessment attempts in the database to monitor how long a student interacts with a module."
        ),
        "why": (
            "This is still mixing two different implementation concerns. The repo tracks lesson completions and separately stores `time_spent_seconds` on assessment attempts. "
            "That does not prove timed module-view engagement in the way the figure currently describes it."
        ),
        "evidence": [
            "Research paper: docs/research-paper-audit/extracted/full_text.txt:567-568",
            "Repo: backend/src/drizzle/schema/base.schema.ts -> `lesson_completions` table and separate `assessment_attempts.time_spent_seconds` field",
            "Repo: test-mobile/src/screens/AssessmentTakeScreen.tsx:190 submits `timeSpentSeconds` for an assessment attempt",
        ],
        "correction": (
            "Rewrite Figure 11 to describe lesson completion/progress honestly, or explicitly say the timing metric belongs to assessment attempts rather than module interaction."
        ),
    },
    {
        "id": "R2-004",
        "severity": "Moderate",
        "location": "Title page / full title",
        "page": "p. 1",
        "exact": "Nexora: A Learning Management With Learning Experience Platform Features...",
        "why": (
            "The title is still grammatically incomplete because it omits the word `System`. This is highly visible and easy for a panelist to catch immediately."
        ),
        "evidence": [
            "Research paper: docs/research-paper-audit/extracted/full_text.txt:2",
        ],
        "correction": (
            "Change the title to `Nexora: A Learning Management System with Learning Experience Platform Features for Targeted Student Intervention for Gat Andres Bonifacio High School`."
        ),
    },
    {
        "id": "R2-005",
        "severity": "Moderate",
        "location": "Chapter 3 architecture wording",
        "page": "Approx. p. 40",
        "exact": "NextJS Backend Core",
        "why": "The backend is NestJS, not NextJS. This is a technical naming error in the architecture narrative.",
        "evidence": [
            "Research paper: docs/research-paper-audit/extracted/full_text.txt:337",
            "Repo: backend/package.json is NestJS-based; next-frontend/package.json is the Next.js app",
        ],
        "correction": "Replace `NextJS Backend Core` with `NestJS backend core` or `NestJS API/backend`.",
    },
    {
        "id": "R2-006",
        "severity": "Minor",
        "location": "Chapter 4, figure narrative prose",
        "page": "Approx. pp. 50-51",
        "exact": (
            "`Figure 3 illustrates the web use case diagram depicts...` / `Figure 4 visualizes the mobile use case diagram illustrates...`"
        ),
        "why": "Both sentences have duplicated verbs and read like unfinished edits rather than polished academic prose.",
        "evidence": [
            "Research paper: docs/research-paper-audit/extracted/full_text.txt:392,395",
        ],
        "correction": (
            "Use a single clean verb: `Figure 3 depicts...` and `Figure 4 illustrates...`."
        ),
    },
]


readiness_score = 84
panel_risk = "Medium"
verdict = "Close, but still needs one short accuracy pass before panel."


def set_doc_defaults(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(22)
    styles["Title"].font.bold = True


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.color.rgb = RGBColor(24, 55, 89)
    return p


def add_bullets(doc: Document, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        p.add_run(item)


def add_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    p.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc: Document, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    for idx, header in enumerate(headers):
        add_cell_text(hdr.cells[idx], header, True)
    set_repeat_table_header(hdr)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            add_cell_text(cells[idx], value)
    doc.add_paragraph()


def build_markdown() -> str:
    lines = [
        "# May 03 Paper Recheck Rev2",
        "",
        f"Source DOCX: `{summary['docx']}`",
        "",
        f"Readiness score: **{readiness_score}/100**",
        f"Panel risk: **{panel_risk}**",
        f"Verdict: **{verdict}**",
        "",
        "## What Improved",
    ]
    lines.extend([f"- {item}" for item in improvements])
    lines += [
        "",
        "## Remaining Findings",
    ]
    for item in findings:
        lines += [
            f"### {item['id']} - {item['severity']}",
            f"- Location: {item['location']}",
            f"- Approx. page: {item['page']}",
            f"- Exact issue: {item['exact']}",
            f"- Why it matters: {item['why']}",
            "- Evidence:",
        ]
        lines.extend([f"  - {ev}" for ev in item["evidence"]])
        lines.append(f"- Correction: {item['correction']}")
        lines.append("")
    lines += [
        "## Bottom Line",
        "- This draft is much better than the last one and several old high-risk findings are now resolved.",
        "- The biggest remaining panel risks are mobile auth overclaiming, the wrong Next.js version, and Figure 11's inaccurate engagement-tracking description.",
        "- After one short correction pass, this should be substantially safer for panel review.",
    ]
    return "\n".join(lines) + "\n"


def build_docx():
    doc = Document()
    set_doc_defaults(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Research Paper Recheck - May 03 Draft Rev2")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(24, 55, 89)

    for line in [
        "System: Nexora",
        f"Research paper file: {Path(summary['docx']).name}",
        f"Generated: {now.strftime('%Y-%m-%d %H:%M:%S')}",
        "Auditor: Codex",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        p.add_run(line)

    doc.add_page_break()

    add_heading(doc, "Executive Verdict")
    for line in [
        f"Readiness score: {readiness_score}/100",
        f"Panel risk: {panel_risk}",
        verdict,
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.add_run(line)

    add_heading(doc, "What Improved From the Previous Draft")
    add_bullets(doc, improvements)

    add_heading(doc, "Remaining Findings")
    rows = []
    for item in findings:
        rows.append(
            [
                item["id"],
                item["severity"],
                item["location"],
                item["exact"],
                item["correction"],
            ]
        )
    add_table(
        doc,
        ["ID", "Severity", "Location", "Issue", "Correction"],
        rows,
    )

    add_heading(doc, "Repo Truth Anchors Used In This Recheck")
    add_bullets(
        doc,
        [
            "Web frontend version: `next-frontend/package.json` -> Next `^16.2.4`, React `19.2.3`.",
            "Mobile auth surface: `test-mobile/src/navigation/types.ts` -> auth stack contains only `Login`.",
            "Mobile forgot-password wording: `test-mobile/src/screens/LoginScreen.tsx` -> `Forgot password? Contact your administrator`.",
            "Teacher mobile limitation: `test-mobile/src/screens/TeacherUnsupportedScreen.tsx` -> teacher mobile is still coming soon.",
            "Intervention threshold: `backend/src/modules/lxp/lxp.service.ts` -> `INTERVENTION_THRESHOLD = 74`.",
            "Engagement/timing reality: `backend/src/drizzle/schema/base.schema.ts` separates `lesson_completions` from `assessment_attempts.time_spent_seconds`.",
            "Swagger path: `backend/src/main.ts` -> `SwaggerModule.setup('api', ...)`.",
        ],
    )

    add_heading(doc, "Bottom Line")
    for line in [
        "This draft is close.",
        "It no longer has the embarrassing copied-text and duplicated-table problems from the earlier versions.",
        "The remaining work is now a short accuracy-and-wording cleanup, not a full structural rescue.",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.add_run(line)

    doc.save(DOCX_OUT)


def main():
    md_text = build_markdown()
    MD_OUT.write_text(md_text, encoding="utf-8")
    build_docx()
    JSON_OUT.write_text(
        json.dumps(
            {
                "source_docx": summary["docx"],
                "readiness_score": readiness_score,
                "panel_risk": panel_risk,
                "verdict": verdict,
                "finding_counts": {
                    "Major": sum(1 for item in findings if item["severity"] == "Major"),
                    "Moderate": sum(1 for item in findings if item["severity"] == "Moderate"),
                    "Minor": sum(1 for item in findings if item["severity"] == "Minor"),
                },
            },
            indent=2,
        ),
        encoding="utf-8",
    )
    print(f"Wrote {DOCX_OUT}")
    print(f"Wrote {MD_OUT}")
    print(f"Wrote {JSON_OUT}")


if __name__ == "__main__":
    main()
