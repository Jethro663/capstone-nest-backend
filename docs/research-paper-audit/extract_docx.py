from __future__ import annotations

import json
import re
import shutil
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
AUDIT_DIR = ROOT / "docs" / "research-paper-audit"
OUT_DIR = AUDIT_DIR / "extracted"

NS = {
    "rel": "http://schemas.openxmlformats.org/package/2006/relationships",
}


def normalize_space(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def paragraph_style_name(paragraph) -> str | None:
    try:
        return paragraph.style.name
    except Exception:
        return None


def extract_docx(docx_path: Path) -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    media_dir = OUT_DIR / "media"
    if media_dir.exists():
        shutil.rmtree(media_dir)
    media_dir.mkdir(parents=True, exist_ok=True)

    doc = Document(str(docx_path))

    paragraphs = []
    full_lines = []
    headings = []
    figure_captions = []
    table_captions = []

    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.rstrip()
        style = paragraph_style_name(paragraph)
        entry = {"index": index, "style": style, "text": text}
        paragraphs.append(entry)
        full_lines.append(text)
        if style and "heading" in style.lower():
            headings.append(entry)
        if re.match(r"Figure\s+\d+[:.]", text):
            figure_captions.append(entry)
        if re.match(r"Table\s+\d+[:.]", text):
            table_captions.append(entry)

    tables = []
    for t_index, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([normalize_space(cell.text) for cell in row.cells])
        tables.append({"index": t_index, "rows": rows})

    rels = []
    media_inventory = []
    with zipfile.ZipFile(docx_path) as archive:
        for info in archive.infolist():
            if info.filename.startswith("word/media/") and not info.is_dir():
                target = media_dir / Path(info.filename).name
                with archive.open(info) as src, target.open("wb") as dst:
                    shutil.copyfileobj(src, dst)
                media_inventory.append(
                    {
                        "zip_path": info.filename,
                        "filename": target.name,
                        "size": target.stat().st_size,
                    }
                )

        rels_path = "word/_rels/document.xml.rels"
        if rels_path in archive.namelist():
            rel_root = ET.fromstring(archive.read(rels_path))
            for rel in rel_root.findall("rel:Relationship", NS):
                rels.append(rel.attrib)

    summary = {
        "docx": str(docx_path),
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "figure_caption_count": len(figure_captions),
        "table_caption_count": len(table_captions),
        "media_count": len(media_inventory),
    }

    (OUT_DIR / "summary.json").write_text(
        json.dumps(summary, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )
    (OUT_DIR / "paragraphs.json").write_text(
        json.dumps(paragraphs, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )
    (OUT_DIR / "headings.json").write_text(
        json.dumps(headings, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )
    (OUT_DIR / "figure_captions.json").write_text(
        json.dumps(figure_captions, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )
    (OUT_DIR / "table_captions.json").write_text(
        json.dumps(table_captions, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )
    (OUT_DIR / "tables.json").write_text(
        json.dumps(tables, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )
    (OUT_DIR / "rels.json").write_text(
        json.dumps(rels, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )
    (OUT_DIR / "media_inventory.json").write_text(
        json.dumps(media_inventory, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )
    (OUT_DIR / "full_text.txt").write_text("\n".join(full_lines), encoding="utf-8")


def main() -> int:
    if len(sys.argv) != 2:
        print("usage: extract_docx.py <path-to-docx>")
        return 2
    docx_path = Path(sys.argv[1]).resolve()
    if not docx_path.exists():
        print(f"missing file: {docx_path}")
        return 1
    extract_docx(docx_path)
    print(f"Extracted {docx_path} into {OUT_DIR}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
