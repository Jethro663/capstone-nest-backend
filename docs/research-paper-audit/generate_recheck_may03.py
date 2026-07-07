from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
AUDIT_DIR = ROOT / "docs" / "research-paper-audit"
EXTRACTED_DIR = AUDIT_DIR / "extracted"
DOCX_OUT = AUDIT_DIR / "RESEARCH_PAPER_RECHECK_MAY03.docx"
MD_OUT = AUDIT_DIR / "paper_recheck_may03.md"
JSON_OUT = AUDIT_DIR / "paper_recheck_may03.json"


def load_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


summary = load_json(EXTRACTED_DIR / "summary.json")
now = datetime.now()


improvements = [
    "The intervention threshold is now mostly corrected from `60%` / `c0%` to `74%` in core Chapter 1, Chapter 2, Chapter 3, and Figure 13 sections.",
    "The copied legacy paragraph about `RND`, `nutritionist`, `food intakes`, and `cloud database` is gone.",
    "The typo `Disccusion` was corrected to `Discussion Board`.",
    "The old duplicate `Student Profile` problem was partially fixed: Table 18 is now `Security and Verification` instead of another `Student Profile`.",
    "The mobile use-case narrative now honestly states that teacher/admin mobile workflows are unsupported in the current repository build.",
    "Figure 30 no longer claims immediate mobile push notifications; it now uses safer web/in-app notification wording.",
    "Figure 21 was corrected from generic `Account Locking` to `Web Profile Detail Locking`, which matches the actual student-profile lock behavior in the web app.",
    "Figure 11 was softened from a hard `30-Second` rule to `Time-Based Engagement Tracking`.",
]


findings = [
    {
        "id": "R-001",
        "severity": "Major",
        "location": "Chapter 3, Software stack description",
        "page": "Approx. p. 36",
        "exact": "The frontend is built using Next.js 19.2.3 and React 19",
        "why": "The repository uses Next `^16.2.4`, not Next `19.2.3`. The paper appears to have inserted the React version into the Next.js slot.",
        "evidence": [
            "Research paper: docs/research-paper-audit/extracted/full_text.txt:289",
            "Repo: next-frontend/package.json -> next ^16.2.4, react 19.2.3",
        ],
        "correction": "Change this to `Next.js 16.2.4 and React 19.2.3`, or state the major versions only.",
    },
    {
        "id": "R-002",
        "severity": "Major",
        "location": "Chapter 1, Scope and Delimitation",
        "page": "Approx. p. 16",
        "exact": "supports all high school grade levels and subjects",
        "why": "The repo grade model is constrained to grade levels `7`, `8`, `9`, and `10`. The implementation evidence does not prove full all-subject, all-high-school deployment breadth.",
        "evidence": [
            "Research paper: docs/research-paper-audit/extracted/full_text.txt:171",
            "Repo: backend/src/common/utils/grade-level.util.ts:5",
            "Repo: backend/src/drizzle/schema/base.schema.ts:68",
        ],
        "correction": "Narrow the statement to the implemented grade coverage or qualify it as intended deployment scope rather than confirmed implemented scope.",
    },
    {
        "id": "R-003",
        "severity": "Major",
        "location": "Chapter 4, Figure 11",
        "page": "Approx. p. 106",
        "exact": "The system utilizes the timeSpentSeconds field in the database to monitor how long a student interacts with a module.",
        "why": "The repo does contain `timeSpentSeconds`, but that field is tied to `assessment_attempts`, not lesson/module engagement. Lesson completion exists, but this specific field-to-module claim is unsupported.",
        "evidence": [
            "Research paper: docs/research-paper-audit/extracted/full_text.txt:563-564",
            "Repo: backend/src/drizzle/schema/base.schema.ts shows lesson_completions and separately assessment_attempts with time_spent_seconds",
            "Repo: mobile/src/screens/AssessmentTakeScreen.tsx: timeSpentSeconds is submitted for assessments",
        ],
        "correction": "Rewrite Figure 11 to describe the actual lesson completion/progress logic, or explicitly tie `timeSpentSeconds` to assessments rather than module viewing.",
    },
    {
        "id": "R-004",
        "severity": "Major",
        "location": "Chapter 4, Use-case tables",
        "page": "Approx. pp. 77 and 81",
        "exact": "Table 31: Use Case Narratives of View Evaluations / Table 36: Use Case Narratives of View Evaluations",
        "why": "The document still contains duplicate use-case titles later in the sequence. This is a cleaner version of the old numbering issue, but it is still a visible documentation integrity problem.",
        "evidence": [
            "Research paper: docs/research-paper-audit/extracted/full_text.txt:102,107,452,466",
        ],
        "correction": "Rename one of the duplicated tables to its intended use case and recheck the list of tables against in-body captions.",
    },
    {
        "id": "R-005",
        "severity": "Moderate",
        "location": "Chapter 1 delimitation vs Chapter 4 technical constraints",
        "page": "Approx. pp. 17 and 45",
        "exact": "offline functionality and synchronization are excluded / mobile application is designed to cache core lesson text and announcements locally",
        "why": "These two statements pull in different directions. The repository does not give clear evidence of a completed offline lesson/announcement cache feature matching this wording.",
        "evidence": [
            "Research paper: docs/research-paper-audit/extracted/full_text.txt:186,382",
            "Repo search found secure storage and file-system usage, but not a clear implemented offline lesson-and-announcement cache contract",
        ],
        "correction": "Either remove the local-cache claim or rewrite the delimitation so the intended offline/resilience behavior is consistent and accurately scoped.",
    },
    {
        "id": "R-006",
        "severity": "Moderate",
        "location": "Chapter 3 architecture wording",
        "page": "Approx. pp. 39-40",
        "exact": "NextJS Backend Core",
        "why": "The backend is NestJS, not NextJS. The web frontend is Next.js. This is a framework naming error in the architecture narrative.",
        "evidence": [
            "Research paper: docs/research-paper-audit/extracted/full_text.txt:325,335",
            "Repo: backend/package.json uses NestJS; next-frontend/package.json uses Next.js",
        ],
        "correction": "Replace `NextJS Backend Core` with `NestJS backend core` or `NestJS API/backend`.",
    },
    {
        "id": "R-007",
        "severity": "Moderate",
        "location": "Chapter 1 wording quality",
        "page": "Approx. pp. 14-16",
        "exact": "student to teacher ratios is often high / awkward repeated high-intensity phrasing",
        "why": "The paper is much cleaner than before, but there are still grammar and phrasing spots that sound rushed or inflated.",
        "evidence": [
            "Research paper: docs/research-paper-audit/extracted/full_text.txt:141 and surrounding prose",
        ],
        "correction": "Do a line edit pass for subject-verb agreement and over-intense phrasing before submission.",
    },
]


readiness_score = 71
panel_risk = "Medium"
verdict = "Better, but still needs revision before panel."


def set_doc_defaults(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(22)
    styles["Title"].font.bold = True


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.color.rgb = RGBColor(24, 55, 89)
    return p


def add_bullets(doc: Document, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        p.add_run(item)


def add_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    p.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc: Document, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    for idx, header in enumerate(headers):
        add_cell_text(hdr.cells[idx], header, True)
    set_repeat_table_header(hdr)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            add_cell_text(cells[idx], value)
    doc.add_paragraph()


def build_markdown() -> str:
    lines = [
        "# May 03 Paper Recheck",
        "",
        f"Source DOCX: `{summary['docx']}`",
        "",
        f"Readiness score: **{readiness_score}/100**",
        f"Panel risk: **{panel_risk}**",
        f"Verdict: **{verdict}**",
        "",
        "## What Improved",
    ]
    lines.extend([f"- {item}" for item in improvements])
    lines += [
        "",
        "## Remaining Findings",
    ]
    for item in findings:
        lines += [
            f"### {item['id']} - {item['severity']}",
            f"- Location: {item['location']}",
            f"- Approx. page: {item['page']}",
            f"- Exact issue: {item['exact']}",
            f"- Why it matters: {item['why']}",
            "- Evidence:",
        ]
        lines.extend([f"  - {ev}" for ev in item["evidence"]])
        lines.append(f"- Correction: {item['correction']}")
        lines.append("")
    lines += [
        "## Bottom Line",
        "- This draft is materially better than the previous one.",
        "- The old threshold typo cluster, copied nutritionist block, and push-notification overclaim were cleaned up.",
        "- The strongest remaining risks are the wrong Next.js version, the overbroad scope claim, the module-engagement tracking claim in Figure 11, and the lingering duplicate `View Evaluations` table title.",
    ]
    return "\n".join(lines) + "\n"


def build_docx():
    doc = Document()
    set_doc_defaults(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Research Paper Recheck - May 03 Draft")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(24, 55, 89)

    for line in [
        "System: Nexora",
        f"Research paper file: {Path(summary['docx']).name}",
        f"Generated: {now.strftime('%Y-%m-%d %H:%M:%S')}",
        "Auditor: Codex",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        p.add_run(line)

    doc.add_page_break()

    add_heading(doc, "Executive Verdict", 1)
    add_bullets(doc, [
        f"Readiness score: {readiness_score}/100",
        f"Panel risk: {panel_risk}",
        verdict,
    ])

    add_heading(doc, "What Improved From the Previous Draft", 1)
    add_bullets(doc, improvements)

    add_heading(doc, "Remaining Findings", 1)
    add_table(
        doc,
        ["ID", "Severity", "Approx. Page", "Issue", "Why It Still Matters", "Correction"],
        [
            [
                item["id"],
                item["severity"],
                item["page"],
                item["exact"],
                item["why"],
                item["correction"],
            ]
            for item in findings
        ],
    )

    add_heading(doc, "Repo Truth Anchors Used In This Recheck", 1)
    add_bullets(doc, [
        "Web frontend version: `next-frontend/package.json` -> Next `^16.2.4`, React `19.2.3`.",
        "Implemented grade levels: `backend/src/common/utils/grade-level.util.ts` and `backend/src/drizzle/schema/base.schema.ts` -> `7`, `8`, `9`, `10`.",
        "Teacher mobile limitation: `mobile/src/screens/TeacherUnsupportedScreen.tsx`.",
        "Mobile discussion limitation: `mobile/src/screens/ClassDetailScreen.tsx`.",
        "Intervention threshold: `backend/src/modules/lxp/lxp.service.ts` -> `INTERVENTION_THRESHOLD = 74`.",
        "Profile locking reality: `next-frontend/src/components/profile/StudentProfilePage.tsx`.",
        "Swagger path: `backend/src/main.ts` -> `SwaggerModule.setup('api', ...)`.",
    ])

    add_heading(doc, "Bottom Line", 1)
    add_bullets(doc, [
        "This draft is clearly improved and no longer has the most embarrassing copied-text problems.",
        "It is not yet panel-safe if the panel is strict about implementation accuracy.",
        "Fix the remaining six-to-seven issues and the paper will be far stronger than the previous submission.",
    ])

    doc.save(DOCX_OUT)


def main():
    MD_OUT.write_text(build_markdown(), encoding="utf-8")
    build_docx()
    JSON_OUT.write_text(
        json.dumps(
            {
                "generated_at": now.isoformat(),
                "docx": str(DOCX_OUT),
                "markdown": str(MD_OUT),
                "source_docx": summary["docx"],
                "readiness_score": readiness_score,
                "panel_risk": panel_risk,
                "finding_counts": {
                    "Major": sum(1 for item in findings if item["severity"] == "Major"),
                    "Moderate": sum(1 for item in findings if item["severity"] == "Moderate"),
                },
            },
            indent=2,
        ),
        encoding="utf-8",
    )


if __name__ == "__main__":
    main()
