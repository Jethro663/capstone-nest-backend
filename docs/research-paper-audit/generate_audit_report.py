from __future__ import annotations

import json
import re
from collections import Counter
from datetime import datetime
from pathlib import Path
from textwrap import dedent

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
AUDIT_DIR = ROOT / "docs" / "research-paper-audit"
EXTRACTED_DIR = AUDIT_DIR / "extracted"
OUTPUT_DOCX = ROOT / "docs" / "system-audit" / "RESEARCH_PAPER_FULL_SYSTEM_AUDIT_REPORT.docx"
OUTPUT_MD = {
    "audit_notes": ROOT / "docs" / "system-audit" / "audit_notes.md",
    "paper_claims": ROOT / "docs" / "system-audit" / "paper_claims_extracted.md",
    "repo_inventory": ROOT / "docs" / "system-audit" / "repo_feature_inventory.md",
    "truth_table": ROOT / "docs" / "system-audit" / "implementation_truth_table.md",
    "diagram_audit": ROOT / "docs" / "system-audit" / "chapter4_diagram_audit.md",
    "live_demo": ROOT / "docs" / "demo" / "live_demo_test_log.md",
    "panelist": ROOT / "docs" / "thesis-defense" / "panelist_risk_checklist.md",
}
METADATA_JSON = AUDIT_DIR / "audit_metadata.json"


def load_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


summary = load_json(EXTRACTED_DIR / "summary.json")
headings = load_json(EXTRACTED_DIR / "headings.json")
repo_scan = load_json(AUDIT_DIR / "repo_scan.json")
playwright_summary = load_json(AUDIT_DIR / "playwright-summary.json")
full_text = (EXTRACTED_DIR / "full_text.txt").read_text(encoding="utf-8")

now = datetime.now()
repo_name = repo_scan["repo_name"]
research_paper_file = Path(summary["docx"]).name
concept_paper_file = "Concept Paper.pdf"


def unique_ordered(items):
    seen = set()
    out = []
    for item in items:
        if item in seen:
            continue
        seen.add(item)
        out.append(item)
    return out


figure_entries = []
for entry in headings:
    text = entry.get("text", "")
    if re.match(r"Figure\s+\d+[:.]", text):
        match = re.match(r"(Figure\s+\d+[:.]\s*.+?)\t(\d+)$", text)
        if match:
            figure_entries.append({"label": match.group(1), "page": match.group(2)})
figure_entries = unique_ordered(tuple(sorted(d.items())) for d in figure_entries)
figure_entries = [dict(items) for items in figure_entries]

table_list_block = []
for line in full_text.splitlines():
    if line.startswith("Table "):
        table_list_block.append(line.strip())
table_list_block = unique_ordered(table_list_block)


readiness_score = 42
panel_risk = "Critical"
paper_safe = "Not safe"
verdict = "Needs revision before panel"


findings = [
    {
        "id": "F-001",
        "severity": "Critical",
        "location": "Research paper body, Chapter 1 problem statement and project description",
        "page": "Approx. pp. 12-13",
        "chapter": "Chapter 1",
        "target": "Paragraphs around lines 126 and 129 in extracted text",
        "exact": "below a c0 percent performance threshold ... scoring below c0%",
        "why": "The paper has an obvious OCR/typing artifact and also states the wrong intervention threshold. The implemented system and the concept paper both use 74%, not 60%.",
        "evidence": [
            "Research paper: docs/research-paper-audit/extracted/full_text.txt:126,129",
            "Concept paper: Concept paper.txt:41,95-97,238",
            "Backend code: backend/src/modules/lxp/lxp.service.ts:35",
            "Backend code: backend/src/drizzle/schema/performance.schema.ts:47",
            "Live DB check: thresholds=[74.000] from ad hoc pg query during audit",
        ],
        "correction": "Replace every c0/c0%/60% intervention threshold statement with the implemented threshold of below 74%, unless the system is changed to 60% everywhere in code, schema defaults, seeded data, and UI copy.",
        "action": "Edit paper",
    },
    {
        "id": "F-002",
        "severity": "Critical",
        "location": "Concept-to-paper alignment on intervention rule",
        "page": "Approx. pp. 12-13, 108",
        "chapter": "Chapters 1 and 4",
        "target": "Scope/objectives plus Figure 13 caption and flow description",
        "exact": "60% mastery threshold gating",
        "why": "The research paper describes a different intervention rule from both the concept paper and the implemented code. A panelist will treat that as a core logic contradiction.",
        "evidence": [
            "Research paper list of figures: Figure 13 on page 108",
            "Research paper: docs/research-paper-audit/extracted/full_text.txt:449",
            "Concept paper: Concept paper.txt:41,95-97",
            "Live UI screenshot: docs/research-paper-audit/screenshots/teacher-interventions.png shows 'Current trigger threshold: 74%'",
        ],
        "correction": "Update all Chapter 4 gating diagrams and narratives to 74% and explain that the LXP path is limited to students below the implemented threshold.",
        "action": "Replace diagram",
    },
    {
        "id": "F-003",
        "severity": "Critical",
        "location": "Chapter 3 technical description",
        "page": "Approx. p. 41",
        "chapter": "Chapter 3",
        "target": "Legacy/copy-paste paragraph",
        "exact": "audit trail of all the chat information (user's name, RND's name, date, and time)... nutritionist... food intakes... cloud database",
        "why": "This paragraph is clearly from a different capstone domain and directly damages the credibility of the paper.",
        "evidence": [
            "Research paper: docs/research-paper-audit/extracted/full_text.txt:271-272",
            "Repository reality: no nutritionist or RND domain exists anywhere in backend, frontend, mobile, or ai-service",
        ],
        "correction": "Delete the entire legacy paragraph and replace it with a correct description of Nexora audit logs, admin AI chat, or diagnostics.",
        "action": "Edit paper",
    },
    {
        "id": "F-004",
        "severity": "Critical",
        "location": "Figure 30 narrative and communication claims",
        "page": "Approx. p. 123",
        "chapter": "Chapter 4",
        "target": "Figure 30 description",
        "exact": "system sends immediate push notifications to student mobile clients",
        "why": "The repository does not show an implemented mobile push-notification stack. Real-time web notifications exist, but mobile push delivery is unsupported by code evidence.",
        "evidence": [
            "Research paper: docs/research-paper-audit/extracted/full_text.txt:428",
            "Repo search: no expo-notifications, FCM, APNs, or device-token pipeline found",
            "Web real-time support exists instead: backend/src/modules/notifications/notifications.gateway.ts:13",
        ],
        "correction": "Change the claim to web notifications and in-app announcement updates unless a real mobile push-notification pipeline is implemented and tested.",
        "action": "Edit paper",
    },
    {
        "id": "F-005",
        "severity": "Critical",
        "location": "Figure 11 and Chapter 4 mobile flow",
        "page": "Approx. p. 106",
        "chapter": "Chapter 4",
        "target": "Figure 11 caption and flow",
        "exact": "30-Second Lesson Completion Tracking",
        "why": "The audit found lesson completion support and time-spent fields, but no clear 30-second completion rule in the implementation. The claim is too specific to leave unverified.",
        "evidence": [
            "Research paper list of figures: Figure 11 page 106",
            "Repo search found timeSpentSeconds field only: backend/src/drizzle/schema/base.schema.ts:636",
            "No explicit 30-second completion rule found in backend/frontend/mobile during repository search",
        ],
        "correction": "Either document the actual implemented lesson-completion logic or remove the 30-second claim from diagrams and use-case narratives.",
        "action": "Edit paper",
    },
    {
        "id": "F-006",
        "severity": "Critical",
        "location": "Mobile scope and role coverage claims",
        "page": "Approx. pp. 56, 103-116",
        "chapter": "Chapters 3 and 4",
        "target": "Mobile use case diagram and multiple mobile process flows",
        "exact": "Nexora mobile application coverage implied for broad role workflows",
        "why": "The mobile app exists, but the current test-mobile build is student-focused and teacher mobile is explicitly marked unsupported. Broad mobile parity claims are therefore overstated.",
        "evidence": [
            "Mobile app exists: test-mobile/package.json:27-40",
            "Teacher placeholder: test-mobile/src/screens/TeacherUnsupportedScreen.tsx:40,45",
            "Discussion board note: test-mobile/src/screens/ClassDetailScreen.tsx:1228",
        ],
        "correction": "Rewrite the mobile scope to student-first parity with selected flows, and mark teacher/admin mobile as future work unless implemented.",
        "action": "Edit paper",
    },
    {
        "id": "F-007",
        "severity": "Critical",
        "location": "Chapter 4 table inventory and use-case list",
        "page": "Approx. pp. 66-80",
        "chapter": "Chapter 4",
        "target": "List of tables and duplicated use-case narratives",
        "exact": "Table 17: Student Profile / Table 18: Student Profile / Table 31 duplicated with different titles",
        "why": "Broken numbering and duplicated titles make the Chapter 4 documentation look unfinished and unreliable.",
        "evidence": [
            "Research paper: docs/research-paper-audit/extracted/full_text.txt:73-87",
            "Research paper: docs/research-paper-audit/extracted/full_text.txt:336-355",
        ],
        "correction": "Re-audit all use-case table numbering and titles. Assign unique titles and ensure list-of-tables, in-body captions, and surrounding narratives match exactly.",
        "action": "Edit paper",
    },
    {
        "id": "F-008",
        "severity": "Critical",
        "location": "Chapter 1 scope and delimitation",
        "page": "Approx. pp. 15-17",
        "chapter": "Chapter 1",
        "target": "All-subject / all-grade coverage claims",
        "exact": "all subjects and grade levels",
        "why": "The actual data model is constrained to junior high grade levels 7-10. The repository also reflects seeded/demo coverage rather than proven full-school deployment across every subject and level.",
        "evidence": [
            "Backend grade enum: backend/src/drizzle/schema/base.schema.ts:68",
            "Backend utility: backend/src/common/utils/grade-level.util.ts:5",
            "Concept paper itself uses broad scope wording, but repo is the final source of truth",
        ],
        "correction": "Narrow the claim to implemented grade levels and qualified deployment scope, or add a clear future-expansion caveat.",
        "action": "Edit paper",
    },
    {
        "id": "F-009",
        "severity": "Major",
        "location": "Chapter 3 stack version claim",
        "page": "Approx. p. 38",
        "chapter": "Chapter 3",
        "target": "Frontend stack description",
        "exact": "Next.js 16.1.6 and React 19",
        "why": "The actual current repository uses Next 16.2.4 and React 19.2.3 in the web app. Exact version claims should match package manifests or be stated as major-version-only claims.",
        "evidence": [
            "Research paper: docs/research-paper-audit/extracted/full_text.txt:240",
            "Frontend package: next-frontend/package.json:52-54",
        ],
        "correction": "Update the exact versions or intentionally write them at the major-version level with a checked date.",
        "action": "Edit paper",
    },
    {
        "id": "F-010",
        "severity": "Major",
        "location": "Chapter 3 AI architecture wording",
        "page": "Approx. p. 39",
        "chapter": "Chapter 3",
        "target": "FastAPI/Ollama security wording",
        "exact": "all AI processing remains within a secure environment",
        "why": "The ai-service includes a cloud fallback pathway. The current design is not strictly Ollama-only in architecture terms.",
        "evidence": [
            "Research paper: docs/research-paper-audit/extracted/full_text.txt:244",
            "AI config: ai-service/app/config.py:51-87",
            "AI fallback implementation: ai-service/app/cloud_fallback.py:15-233",
        ],
        "correction": "State that the default local path uses Ollama, while the service also supports optional cloud fallback under configuration.",
        "action": "Edit paper",
    },
    {
        "id": "F-011",
        "severity": "Major",
        "location": "Chapter 4 Figure 21 wording",
        "page": "Approx. p. 116",
        "chapter": "Chapter 4",
        "target": "Web Profile Management and Account Locking",
        "exact": "account locking",
        "why": "The audit found profile-completion lock behavior for student details, but not a clearly implemented user-facing account-locking flow matching the caption wording.",
        "evidence": [
            "Figure list: Figure 21 page 116",
            "Student profile lock UI: next-frontend/src/components/profile/StudentProfilePage.tsx:221,542-543",
            "No matching explicit account-lock workflow surfaced in live audit",
        ],
        "correction": "Rename the figure to profile completion locking or verify and document the real account-locking workflow if it exists.",
        "action": "Replace diagram",
    },
    {
        "id": "F-012",
        "severity": "Major",
        "location": "Chapter 4 Figure 15 and analytics claims",
        "page": "Approx. p. 110",
        "chapter": "Chapter 4",
        "target": "Performance analytics and quarterly trends",
        "exact": "quarterly trends / competency heatmaps / evaluations",
        "why": "Analytics surfaces exist, but some data domains are only partially populated in the live dataset. The paper should not imply fully validated school-scale analytics if live evidence is sparse.",
        "evidence": [
            "Routes exist in frontend and backend analytics modules",
            "Live DB query during audit: system_evaluations = 0 rows",
            "Backend modules: analytics, performance, reports present in repo_scan.json",
        ],
        "correction": "Qualify these features as implemented analytics surfaces with limited current live data, unless fuller seeded/demo evidence is prepared.",
        "action": "Edit paper",
    },
    {
        "id": "F-013",
        "severity": "Major",
        "location": "Chapter 4 Figure 4 and mobile discussion implications",
        "page": "Approx. p. 56",
        "chapter": "Chapter 4",
        "target": "Mobile use case diagram",
        "exact": "discussion-thread participation implied as implemented mobile behavior",
        "why": "The mobile class-detail screen itself states that student discussion posts are not backed by a live mobile data source yet.",
        "evidence": [
            "Mobile note: test-mobile/src/screens/ClassDetailScreen.tsx:1228",
        ],
        "correction": "Mark discussion-board mobile parity as partial or future scope.",
        "action": "Edit paper",
    },
    {
        "id": "F-014",
        "severity": "Major",
        "location": "Chapter 3 documentation precision",
        "page": "Approx. p. 38",
        "chapter": "Chapter 3",
        "target": "Swagger/OpenAPI claim if route is specified in narrative or screenshots",
        "exact": "Swagger route assumed as /api/docs",
        "why": "The backend exposes Swagger at /api, not /api/docs. During the audit, /api/docs returned 404 while /api/health/ready succeeded.",
        "evidence": [
            "Backend setup: backend/src/main.ts:123-124",
            "Live request: GET http://localhost:3000/api/docs returned 404 during audit",
        ],
        "correction": "If Swagger is mentioned, use the correct configured route.",
        "action": "Edit paper",
    },
    {
        "id": "F-015",
        "severity": "Major",
        "location": "Chapter 1 objective wording",
        "page": "Approx. p. 14",
        "chapter": "Chapter 1",
        "target": "AI NPC terminology",
        "exact": "AI Non-Placer Character (NPC) mentor",
        "why": "This is a terminology error. The standard term is Non-Player Character.",
        "evidence": [
            "Research paper: docs/research-paper-audit/extracted/full_text.txt:138",
        ],
        "correction": "Replace with Non-Player Character and standardize JAKIPIR AI Mentor naming across the paper.",
        "action": "Edit paper",
    },
    {
        "id": "F-016",
        "severity": "Major",
        "location": "Chapter 4 table title",
        "page": "Approx. p. 73",
        "chapter": "Chapter 4",
        "target": "Table 24 title",
        "exact": "Disccusion Board",
        "why": "The typo is visible in a table caption and list entry. That is a basic quality-control failure in a defense document.",
        "evidence": [
            "Research paper: docs/research-paper-audit/extracted/full_text.txt:80,343",
        ],
        "correction": "Replace with Discussion Board everywhere.",
        "action": "Edit paper",
    },
    {
        "id": "F-017",
        "severity": "Major",
        "location": "Chapter 2 references",
        "page": "Approx. p. 34",
        "chapter": "Chapter 2",
        "target": "Reference list credibility",
        "exact": "SHILAP Revista De Lepidopterologia cited for Catch-Up Fridays implementation challenges",
        "why": "The venue appears mismatched to the education topic and is likely to trigger panel scrutiny on source credibility and citation hygiene.",
        "evidence": [
            "Research paper references: docs/research-paper-audit/extracted/full_text.txt:468",
        ],
        "correction": "Re-verify the source and replace it with a credible education or policy publication if necessary.",
        "action": "Edit paper",
    },
    {
        "id": "F-018",
        "severity": "Major",
        "location": "Chapter 3 mobile stack description",
        "page": "Approx. p. 38",
        "chapter": "Chapter 3",
        "target": "Expo / React Native exactness",
        "exact": "Expo SDK 54 and React Native",
        "why": "The claim is broadly true, but if the paper uses exact dependency versions elsewhere, it should match the current manifest precisely. React on mobile is 19.1.0, not 19.2.3 like web.",
        "evidence": [
            "Mobile package: test-mobile/package.json:27-40",
        ],
        "correction": "Use exact mobile manifest versions or state the stack at the major/minor level without false precision.",
        "action": "Edit paper",
    },
    {
        "id": "F-019",
        "severity": "Moderate",
        "location": "Runtime demo readiness",
        "page": "N/A",
        "chapter": "Live system audit",
        "target": "Student JA Hub load behavior",
        "exact": "Initial student JA page capture showed a long loading state before content rendered",
        "why": "The feature exists, but the cold-load experience can look broken in a live panel demo if the presenter clicks away too early.",
        "evidence": [
            "First capture: docs/research-paper-audit/screenshots/student-ja.png",
            "Delayed capture: docs/research-paper-audit/screenshots/student-ja-delayed.png",
        ],
        "correction": "Warm the route before demo or improve initial loading feedback.",
        "action": "Update code",
    },
    {
        "id": "F-020",
        "severity": "Moderate",
        "location": "Figure 2 caption format",
        "page": "Approx. p. 48",
        "chapter": "Chapter 4 front matter",
        "target": "List of figures",
        "exact": "Figure 2. The Agile Development",
        "why": "Caption punctuation is inconsistent with the colon style used in most other figure labels.",
        "evidence": [
            "Research paper: docs/research-paper-audit/extracted/full_text.txt beginning lines 15-20",
        ],
        "correction": "Standardize figure caption punctuation across the document.",
        "action": "Edit paper",
    },
    {
        "id": "F-021",
        "severity": "Moderate",
        "location": "Chapter 2 prose",
        "page": "Approx. pp. 25-34",
        "chapter": "Chapter 2",
        "target": "Overclaiming tone",
        "exact": "massive, exhaustive, highly relevant, profound, agile, critical, dynamic",
        "why": "The writing often overstates certainty and scale beyond what the repository evidence supports.",
        "evidence": [
            "Research paper lines 124-223 use repeated high-intensity language",
        ],
        "correction": "Use more defensible academic wording tied to implemented evidence.",
        "action": "Edit paper",
    },
    {
        "id": "F-022",
        "severity": "Moderate",
        "location": "Chapter 3 architecture accuracy",
        "page": "Approx. pp. 38-42",
        "chapter": "Chapter 3",
        "target": "Cloud database wording",
        "exact": "cloud database",
        "why": "This contradicts the local Postgres plus local Ollama narrative used elsewhere in the paper.",
        "evidence": [
            "Research paper: docs/research-paper-audit/extracted/full_text.txt:271",
            "docker-compose.yml includes local postgres, redis, ollama, ai-service, grafana stack",
        ],
        "correction": "Replace with PostgreSQL deployment wording that matches the actual environment being described.",
        "action": "Edit paper",
    },
    {
        "id": "F-023",
        "severity": "Moderate",
        "location": "Chapter 4 figure existence QA",
        "page": "Approx. pp. 101-134",
        "chapter": "Chapter 4",
        "target": "Figure media inventory",
        "exact": "39 figures listed but only 37 extracted media files",
        "why": "This does not prove missing figures by itself because some Word figures can be vector or grouped objects, but it does warrant manual layout review.",
        "evidence": [
            "Extract summary: docs/research-paper-audit/extracted/summary.json media_count=37",
            "List of figures includes Figure 1 through Figure 39",
        ],
        "correction": "Open the DOCX in Word or LibreOffice and verify that all figures render and are captioned correctly.",
        "action": "Update screenshot",
    },
    {
        "id": "F-024",
        "severity": "Minor",
        "location": "Global text hygiene",
        "page": "Multiple",
        "chapter": "Multiple",
        "target": "Capitalization and naming consistency",
        "exact": "AI NPC / AI mentor / AI tutor / JAKIPIR mentor / chatbot",
        "why": "Multiple overlapping labels are used for similar or related features, which makes the system boundary unclear.",
        "evidence": [
            "Research paper lines 126, 138, 221, 317 and list sections",
        ],
        "correction": "Define official terms once, then use them consistently.",
        "action": "Edit paper",
    },
    {
        "id": "F-025",
        "severity": "Minor",
        "location": "Table of figures / tables formatting",
        "page": "Front matter",
        "chapter": "Front matter",
        "target": "Figure/table punctuation and title casing",
        "exact": "Mixed colon and period usage; inconsistent capitalization",
        "why": "Small but visible polish issue.",
        "evidence": [
            "Front matter lines from docs/research-paper-audit/extracted/full_text.txt",
        ],
        "correction": "Normalize style across front matter and in-body captions.",
        "action": "Edit paper",
    },
]


severity_order = {"Critical": 0, "Major": 1, "Moderate": 2, "Minor": 3}
findings = sorted(findings, key=lambda item: (severity_order[item["severity"]], item["id"]))
severity_counts = Counter(item["severity"] for item in findings)


feature_rows = [
    {
        "feature": "Role-based login",
        "paper": "Claimed",
        "concept": "Claimed",
        "backend": "Yes",
        "web": "Yes",
        "mobile": "Yes",
        "db": "Yes",
        "live": "Yes",
        "status": "Real",
        "evidence": "backend/src/common/constants/role.constants.ts:10; next-frontend/app/(auth)/login/page.tsx; test-mobile/src/screens/LoginScreen.tsx",
        "paper_fix": "None beyond versioned wording.",
        "system_fix": "None",
    },
    {
        "feature": "OTP verification",
        "paper": "Claimed",
        "concept": "Claimed",
        "backend": "Yes",
        "web": "Yes",
        "mobile": "Partial",
        "db": "Yes",
        "live": "Unverified end-to-end",
        "status": "Partial",
        "evidence": "backend/src/modules/otp/otp.controller.ts:16; backend/src/drizzle/schema/otp.schema.ts:22; next-frontend/app/(auth)/verify-email/page.tsx",
        "paper_fix": "Avoid claiming fully verified live OTP unless demonstrated.",
        "system_fix": "Optional: verify mobile OTP UX end-to-end",
    },
    {
        "feature": "Forgot password",
        "paper": "Claimed",
        "concept": "Claimed",
        "backend": "Yes",
        "web": "Yes",
        "mobile": "Unclear",
        "db": "Yes",
        "live": "Web only",
        "status": "Partial",
        "evidence": "backend/src/modules/auth/auth.controller.ts:318,335; next-frontend/app/(auth)/forgot-password/page.tsx",
        "paper_fix": "Scope mobile claim carefully.",
        "system_fix": "Clarify mobile parity",
    },
    {
        "feature": "Admin dashboard",
        "paper": "Claimed",
        "concept": "Claimed",
        "backend": "Yes",
        "web": "Yes",
        "mobile": "No",
        "db": "N/A",
        "live": "Yes",
        "status": "Real",
        "evidence": "next-frontend/app/(dashboard)/dashboard/admin/page.tsx; screenshot admin-dashboard.png",
        "paper_fix": "Note web-admin scope.",
        "system_fix": "None",
    },
    {
        "feature": "Teacher dashboard",
        "paper": "Claimed",
        "concept": "Claimed",
        "backend": "Yes",
        "web": "Yes",
        "mobile": "No",
        "db": "N/A",
        "live": "Yes",
        "status": "Partial",
        "evidence": "teacher route screenshots; TeacherUnsupportedScreen.tsx:40,45",
        "paper_fix": "Do not imply mobile teacher parity.",
        "system_fix": "Build teacher mobile if needed",
    },
    {
        "feature": "Student dashboard",
        "paper": "Claimed",
        "concept": "Claimed",
        "backend": "Yes",
        "web": "Yes",
        "mobile": "Yes",
        "db": "N/A",
        "live": "Web yes",
        "status": "Real",
        "evidence": "student dashboard routes in web and mobile",
        "paper_fix": "None",
        "system_fix": "None",
    },
    {
        "feature": "Class management",
        "paper": "Claimed",
        "concept": "Claimed",
        "backend": "Yes",
        "web": "Yes",
        "mobile": "No",
        "db": "Yes",
        "live": "Teacher/admin routes verified",
        "status": "Real",
        "evidence": "classes module and routes",
        "paper_fix": "None",
        "system_fix": "None",
    },
    {
        "feature": "Section management",
        "paper": "Claimed",
        "concept": "Claimed",
        "backend": "Yes",
        "web": "Yes",
        "mobile": "No",
        "db": "Yes",
        "live": "Not directly clicked in this audit",
        "status": "Real",
        "evidence": "backend/src/modules/sections; frontend admin sections routes",
        "paper_fix": "None",
        "system_fix": "None",
    },
    {
        "feature": "Roster import",
        "paper": "Claimed",
        "concept": "Claimed",
        "backend": "Yes",
        "web": "Yes",
        "mobile": "No",
        "db": "Yes",
        "live": "Unverified",
        "status": "Real",
        "evidence": "roster-import controller and admin route",
        "paper_fix": "None",
        "system_fix": "Optional live smoke",
    },
    {
        "feature": "CSV/Excel validation",
        "paper": "Claimed",
        "concept": "Claimed",
        "backend": "Yes",
        "web": "Yes",
        "mobile": "No",
        "db": "N/A",
        "live": "Unverified",
        "status": "Partial",
        "evidence": "roster-import module and admin page",
        "paper_fix": "Avoid overstating live proof.",
        "system_fix": "Run import demo flow",
    },
    {
        "feature": "Lessons/modules",
        "paper": "Claimed",
        "concept": "Claimed",
        "backend": "Yes",
        "web": "Yes",
        "mobile": "Yes",
        "db": "Yes",
        "live": "Yes",
        "status": "Real",
        "evidence": "lessons/modules schemas and routes",
        "paper_fix": "None",
        "system_fix": "None",
    },
    {
        "feature": "Course browsing",
        "paper": "Claimed",
        "concept": "Claimed",
        "backend": "Yes",
        "web": "Yes",
        "mobile": "Yes",
        "db": "Yes",
        "live": "Web yes",
        "status": "Real",
        "evidence": "student courses screenshot; StudentClassesIndexPage routes",
        "paper_fix": "None",
        "system_fix": "None",
    },
    {
        "feature": "30-second completion tracking",
        "paper": "Claimed",
        "concept": "Not explicit",
        "backend": "No clear rule",
        "web": "No clear rule",
        "mobile": "No clear rule",
        "db": "Time spent only",
        "live": "No",
        "status": "Unsupported",
        "evidence": "base.schema.ts:636 timeSpentSeconds only; no explicit 30-second rule found",
        "paper_fix": "Remove or rewrite",
        "system_fix": "Implement if required",
    },
    {
        "feature": "Assessments",
        "paper": "Claimed",
        "concept": "Claimed",
        "backend": "Yes",
        "web": "Yes",
        "mobile": "Partial",
        "db": "Yes",
        "live": "Web static/runtime evidence",
        "status": "Real",
        "evidence": "assessment tables and pages; live nav smoke",
        "paper_fix": "Scope mobile precisely",
        "system_fix": "Optional mobile parity",
    },
    {
        "feature": "Out-of-focus warning",
        "paper": "Claimed",
        "concept": "Not explicit",
        "backend": "N/A",
        "web": "Yes",
        "mobile": "No evidence",
        "db": "N/A",
        "live": "Static only",
        "status": "Partial",
        "evidence": "next-frontend/app/(dashboard)/dashboard/student/assessments/[id]/take/page.tsx:355,417,699-712",
        "paper_fix": "Call it web assessment logic, not generic mobile/web unless proven",
        "system_fix": "Optional live scenario proof",
    },
    {
        "feature": "Score calculation",
        "paper": "Claimed",
        "concept": "Claimed",
        "backend": "Yes",
        "web": "Yes",
        "mobile": "Partial",
        "db": "Yes",
        "live": "Indirect yes",
        "status": "Real",
        "evidence": "assessment/class-record services; seed smoke counts",
        "paper_fix": "None",
        "system_fix": "None",
    },
    {
        "feature": "60% threshold gating",
        "paper": "Claimed",
        "concept": "No",
        "backend": "No",
        "web": "No",
        "mobile": "No",
        "db": "No",
        "live": "No",
        "status": "Unsupported",
        "evidence": "actual threshold 74 in concept, code, UI, DB",
        "paper_fix": "Replace with 74",
        "system_fix": "None unless policy changes",
    },
    {
        "feature": "LXP unlock",
        "paper": "Claimed",
        "concept": "Claimed",
        "backend": "Yes",
        "web": "Yes",
        "mobile": "Partial",
        "db": "Yes",
        "live": "Yes via teacher interventions UI plus DB",
        "status": "Real",
        "evidence": "lxp.service.ts; intervention_cases table; teacher-interventions screenshot",
        "paper_fix": "Use 74 threshold wording",
        "system_fix": "None",
    },
    {
        "feature": "Remedial playlist",
        "paper": "Claimed",
        "concept": "Claimed",
        "backend": "Yes",
        "web": "Yes",
        "mobile": "Partial",
        "db": "Yes",
        "live": "Indirect",
        "status": "Partial",
        "evidence": "student lxp routes and lxp schema",
        "paper_fix": "Avoid over-describing if not fully demoed",
        "system_fix": "Optional richer demo data",
    },
    {
        "feature": "JAKIPIR AI mentor",
        "paper": "Claimed",
        "concept": "Claimed in lighter form",
        "backend": "Yes",
        "web": "Yes",
        "mobile": "No full proof",
        "db": "Yes",
        "live": "Web route and ai-service healthy",
        "status": "Real",
        "evidence": "ai-service student JA endpoints; /dashboard/student/ja live route",
        "paper_fix": "Standardize terminology",
        "system_fix": "Improve load perception",
    },
    {
        "feature": "RAG grounding",
        "paper": "Claimed",
        "concept": "Only lighter AI integration claim",
        "backend": "Yes",
        "web": "Yes",
        "mobile": "N/A",
        "db": "Yes",
        "live": "Static + DB counts",
        "status": "Real",
        "evidence": "ai-service/retrieval_service.py:257-261; content_chunk_embeddings table; live chunk counts 391",
        "paper_fix": "None",
        "system_fix": "None",
    },
    {
        "feature": "PDF extraction",
        "paper": "Claimed",
        "concept": "Claimed",
        "backend": "Yes",
        "web": "Yes",
        "mobile": "N/A",
        "db": "Yes",
        "live": "Service tests only",
        "status": "Real",
        "evidence": "ai-service/app/main.py /extract endpoints; extracted_modules table",
        "paper_fix": "Clarify teacher review/verification step",
        "system_fix": "None",
    },
    {
        "feature": "AI quiz drafting",
        "paper": "Claimed",
        "concept": "Claimed",
        "backend": "Yes",
        "web": "Yes",
        "mobile": "N/A",
        "db": "Yes",
        "live": "Static + DB jobs",
        "status": "Real",
        "evidence": "ai-service /teacher/quizzes/jobs; ai_generation_jobs table; live ai_jobs=44",
        "paper_fix": "None",
        "system_fix": "Optional live demo proof",
    },
    {
        "feature": "Asynchronous AI jobs",
        "paper": "Claimed",
        "concept": "Partial",
        "backend": "Yes",
        "web": "Yes",
        "mobile": "N/A",
        "db": "Yes",
        "live": "Static + DB",
        "status": "Real",
        "evidence": "BullMQ deps; ai_generation_jobs table; ai-service 202 endpoints",
        "paper_fix": "None",
        "system_fix": "None",
    },
    {
        "feature": "Notifications",
        "paper": "Claimed",
        "concept": "Claimed",
        "backend": "Yes",
        "web": "Yes",
        "mobile": "In-app partial",
        "db": "Yes",
        "live": "Web yes",
        "status": "Real",
        "evidence": "notifications schema/controller/gateway; live notifications row count 33",
        "paper_fix": "Distinguish web/in-app from push",
        "system_fix": "Add push stack only if needed",
    },
    {
        "feature": "Socket.IO real-time updates",
        "paper": "Claimed",
        "concept": "Not explicit",
        "backend": "Yes",
        "web": "Yes",
        "mobile": "Not proven",
        "db": "N/A",
        "live": "Static only",
        "status": "Real",
        "evidence": "backend/src/modules/notifications/notifications.gateway.ts:13; frontend deps",
        "paper_fix": "None",
        "system_fix": "Optional mobile parity",
    },
    {
        "feature": "Redis/BullMQ queue",
        "paper": "Claimed",
        "concept": "Not explicit",
        "backend": "Yes",
        "web": "N/A",
        "mobile": "N/A",
        "db": "N/A",
        "live": "Health yes",
        "status": "Real",
        "evidence": "backend/package.json:47; health ready endpoint reported redis ok",
        "paper_fix": "None",
        "system_fix": "None",
    },
    {
        "feature": "PostgreSQL/pgvector",
        "paper": "Claimed",
        "concept": "Not explicit on pgvector",
        "backend": "Yes",
        "web": "N/A",
        "mobile": "N/A",
        "db": "Yes",
        "live": "Yes",
        "status": "Real",
        "evidence": "docker-compose.yml pgvector image; rag.schema tables; retrieval_service vector casts",
        "paper_fix": "None",
        "system_fix": "None",
    },
    {
        "feature": "Audit trail",
        "paper": "Claimed",
        "concept": "Claimed",
        "backend": "Yes",
        "web": "Yes",
        "mobile": "No",
        "db": "Yes",
        "live": "Yes",
        "status": "Real",
        "evidence": "audit_logs table; admin audit screenshot; live audit_logs=1214",
        "paper_fix": "Replace legacy nutritionist paragraph with this actual feature",
        "system_fix": "None",
    },
    {
        "feature": "System diagnostics",
        "paper": "Claimed",
        "concept": "Not explicit",
        "backend": "Yes",
        "web": "Yes",
        "mobile": "No",
        "db": "N/A",
        "live": "Yes",
        "status": "Real",
        "evidence": "health controller; admin diagnostics screenshot",
        "paper_fix": "None",
        "system_fix": "None",
    },
    {
        "feature": "Academic year/quarter settings",
        "paper": "Claimed",
        "concept": "Claimed",
        "backend": "Yes",
        "web": "Yes",
        "mobile": "No",
        "db": "Yes",
        "live": "Static only",
        "status": "Real",
        "evidence": "academic-state controller/schema; admin routes",
        "paper_fix": "None",
        "system_fix": "Optional live walkthrough",
    },
    {
        "feature": "Performance analytics",
        "paper": "Claimed",
        "concept": "Claimed",
        "backend": "Yes",
        "web": "Yes",
        "mobile": "Yes limited",
        "db": "Yes",
        "live": "Partial",
        "status": "Partial",
        "evidence": "analytics/performance modules; performance tables; live snapshots present",
        "paper_fix": "Qualify live data completeness",
        "system_fix": "Seed fuller demo data",
    },
    {
        "feature": "Competency heatmaps",
        "paper": "Claimed",
        "concept": "Claimed",
        "backend": "Yes",
        "web": "Yes",
        "mobile": "No evidence",
        "db": "Yes",
        "live": "Unverified",
        "status": "Partial",
        "evidence": "student_concept_mastery table; analytics/performance modules",
        "paper_fix": "Avoid calling it fully validated unless shown",
        "system_fix": "Prepare live demo data",
    },
    {
        "feature": "Quarterly trends",
        "paper": "Claimed",
        "concept": "Claimed",
        "backend": "Yes",
        "web": "Yes",
        "mobile": "Partial",
        "db": "Yes",
        "live": "Partial",
        "status": "Partial",
        "evidence": "performance schema; analytics routes",
        "paper_fix": "Qualify dataset coverage",
        "system_fix": "Seed fuller demo data",
    },
    {
        "feature": "Reports/export",
        "paper": "Claimed",
        "concept": "Claimed",
        "backend": "Yes",
        "web": "Yes",
        "mobile": "No",
        "db": "Yes",
        "live": "Unverified",
        "status": "Real",
        "evidence": "reports module; admin pages; class template export",
        "paper_fix": "Do not overstate live proof",
        "system_fix": "Run export smoke if defense depends on it",
    },
    {
        "feature": "Discussion boards",
        "paper": "Claimed",
        "concept": "Claimed",
        "backend": "Yes",
        "web": "Yes",
        "mobile": "Partial/mock",
        "db": "Yes",
        "live": "Web static/runtime evidence",
        "status": "Partial",
        "evidence": "discussion board schema/controller; mobile note says not live-backed",
        "paper_fix": "Mark mobile parity as partial",
        "system_fix": "Implement mobile data source",
    },
    {
        "feature": "Announcements",
        "paper": "Claimed",
        "concept": "Claimed",
        "backend": "Yes",
        "web": "Yes",
        "mobile": "In-app partial",
        "db": "Yes",
        "live": "Web yes",
        "status": "Real",
        "evidence": "announcements schema/controller; student announcements screenshot",
        "paper_fix": "Remove push claim",
        "system_fix": "Optional push implementation",
    },
    {
        "feature": "Profile management",
        "paper": "Claimed",
        "concept": "Claimed",
        "backend": "Yes",
        "web": "Yes",
        "mobile": "Yes",
        "db": "Yes",
        "live": "Static/runtime evidence",
        "status": "Real",
        "evidence": "profiles module; StudentProfilePage; mobile profile screen",
        "paper_fix": "Use accurate lock wording",
        "system_fix": "None",
    },
    {
        "feature": "Class record sync/spreadsheet",
        "paper": "Claimed",
        "concept": "Claimed",
        "backend": "Yes",
        "web": "Yes",
        "mobile": "No",
        "db": "Yes",
        "live": "Static + post-seed smoke",
        "status": "Real",
        "evidence": "class-record module/schema; post-seed-smoke.js results",
        "paper_fix": "None",
        "system_fix": "Optional live route click proof",
    },
    {
        "feature": "Mobile app flows",
        "paper": "Claimed",
        "concept": "Claimed",
        "backend": "Yes",
        "web": "N/A",
        "mobile": "Yes but student-first",
        "db": "N/A",
        "live": "Typecheck/test only",
        "status": "Partial",
        "evidence": "test-mobile package and screen inventory; teacher unsupported placeholder",
        "paper_fix": "Narrow scope",
        "system_fix": "Build missing parity",
    },
    {
        "feature": "Admin AI chatbot",
        "paper": "Claimed",
        "concept": "Not prominent",
        "backend": "Yes",
        "web": "Yes",
        "mobile": "No",
        "db": "Yes logs/history",
        "live": "Route exists",
        "status": "Real",
        "evidence": "next-frontend/app/(dashboard)/dashboard/admin/chatbot/page.tsx; ai-service/app/main.py:1690",
        "paper_fix": "None",
        "system_fix": "Optional live interaction demo",
    },
    {
        "feature": "AI mentor oversight",
        "paper": "Claimed",
        "concept": "Claimed in teacher-guided form",
        "backend": "Yes",
        "web": "Yes",
        "mobile": "N/A",
        "db": "Yes",
        "live": "Partial",
        "status": "Partial",
        "evidence": "teacher intervention jobs; ai logs; admin chat auditing",
        "paper_fix": "State oversight boundaries precisely",
        "system_fix": "Optional stronger review UX",
    },
]


figure_audit = [
    ("Figure 1", "44", "System Architecture", "Partial", "Core architecture is broadly supported, but wording should mention optional cloud fallback instead of implying strictly local-only AI."),
    ("Figure 2", "48", "Agile Development", "Match", "Generic process figure; fix punctuation and make phases match actual repo workflow if detailed."),
    ("Figure 3", "56", "Use Case Diagram for Nexora Web Application", "Partial", "Web roles and modules exist, but diagram should not overclaim mobile push or unverified admin/teacher AI parity."),
    ("Figure 4", "56", "Use Case Diagram for Nexora Mobile Application", "Mismatch", "Mobile exists but is student-focused; teacher mobile remains unsupported."),
    ("Figure 5", "101", "Functional Decomposition Diagram of the Admin Portal", "Partial", "Admin portal is real; ensure chatbot, diagnostics, audit, and roster features match actual route names."),
    ("Figure 6", "101", "Functional Decomposition Diagram of the Teacher Portal", "Match", "Teacher portal largely aligns; keep intervention threshold at 74%."),
    ("Figure 7", "102", "Functional Decomposition Diagram of the Student Portal", "Partial", "Student portal is real, but JA/LXP wording must use actual gating and load behavior."),
    ("Figure 8", "103", "Mobile Login & OTP Verification", "Partial", "OTP backend exists; mobile end-to-end flow was not live-verified in this audit."),
    ("Figure 9", "104", "Mobile Forgot Password and Account Recovery", "Partial", "Web forgot-password flow exists; mobile parity not fully verified."),
    ("Figure 10", "105", "View Dashboard and Continue Learning Logic", "Partial", "Student dashboards exist; validate exact mobile continue-learning logic before claiming detailed process certainty."),
    ("Figure 11", "106", "View Modules and 30-Second Lesson Completion Tracking", "Mismatch", "30-second completion rule is unsupported by current code evidence."),
    ("Figure 12", "107", "Taking Assessments and Summative Submission", "Match", "Assessment flows are implemented; mobile parity should be qualified if shown in figure."),
    ("Figure 13", "108", "LXP Remedial Access and 60% Mastery Threshold Gating", "Mismatch", "Implemented threshold is 74%, not 60%."),
    ("Figure 14", "109", "RAG-Based Interaction with JAKIPIR AI Mentor", "Match", "RAG and JA routes are implemented; keep model/grounding explanation precise."),
    ("Figure 15", "110", "View Performance Analytics and Quarterly Trends", "Partial", "Analytics surfaces exist, but some live datasets are sparse."),
    ("Figure 16", "111", "Mobile Profile Management and Data Persistence", "Partial", "Profile surfaces exist; mobile details were not live-walked in this audit."),
    ("Figure 17", "112", "Web Login and Dashboard Navigation", "Match", "Live-verified."),
    ("Figure 18", "113", "Web Course Browsing and Filtering", "Match", "Live student courses route verified."),
    ("Figure 19", "114", "Assessment Taking with Out-of-Focus Warning Logic", "Match", "Web implementation exists in student assessment page."),
    ("Figure 20", "115", "Discussion Board Participation and Thread Replies", "Partial", "Web discussion support exists; mobile parity does not."),
    ("Figure 21", "116", "Web Profile Management and Account Locking", "Mismatch", "Profile completion lock exists; generic account-lock wording is unsupported."),
    ("Figure 22", "117", "Managing Classes and Instructional Materials", "Match", "Teacher/admin class management exists."),
    ("Figure 23", "118", "AI PDF Extraction and Layout-Aware Section Application", "Partial", "Extraction pipeline exists; layout-aware wording should reflect teacher review and queue-based processing."),
    ("Figure 24", "119", "AI Quiz Drafting and Assessment Studio Workspace", "Match", "Supported by AI job endpoints and teacher/admin assessment editor flows."),
    ("Figure 25", "120", "Asynchronous AI Job Monitoring and Status Updates", "Match", "Supported by queue/job tables and async endpoints."),
    ("Figure 26", "121", "Class Record Synchronization and Spreadsheet Management", "Match", "Class-record subsystem exists."),
    ("Figure 27", "122", "Student Intervention Triage and Outcome Tracking", "Match", "Intervention subsystem exists; threshold wording must stay at 74%."),
    ("Figure 28", "122", "Learning Gap Analysis and Competency Heatmap Generation", "Partial", "Schema support exists; live proof is limited."),
    ("Figure 29", "123", "Academic Report Generation and Data Export", "Partial", "Report/export surfaces exist, but this audit did not execute every export path."),
    ("Figure 30", "123", "Creating Class Announcements and Discussion Threads", "Mismatch", "Push notification claim is unsupported; web/in-app announcement propagation is the supported claim."),
    ("Figure 31", "124", "Admin Dashboard Overview and Quick Route Navigation", "Match", "Admin dashboard live-verified."),
    ("Figure 32", "125", "User Lifecycle Management", "Match", "Users/admin surfaces and lifecycle endpoints exist."),
    ("Figure 33", "126", "Admin-Triggered User Password Reset", "Match", "Backend reset-password endpoint exists."),
    ("Figure 34", "127", "Section Creation and Roster Management", "Match", "Admin sections and roster surfaces exist."),
    ("Figure 35", "128", "Bulk Roster Import and Validation Logic", "Partial", "Feature exists; not live-run during this audit."),
    ("Figure 36", "129", "School Calendar and Event Timeline Management", "Match", "School-events module and route exist."),
    ("Figure 37", "131", "System Diagnostics and Dependency Health Checks", "Match", "Live-verified admin diagnostics and health route."),
    ("Figure 38", "133", "Audit Trail Review and Security Log Filtering", "Match", "Live-verified audit trail page and audit_logs table."),
    ("Figure 39", "134", "Academic State Transition", "Match", "Academic-state module and schema exist."),
]


table_audit = [
    ("Table 1", "21", "LMS vs LXP comparison", "Ensure the 74% threshold and constrained intervention access are reflected consistently."),
    ("Table 2", "34", "Gap analysis", "Some comparative claims are too absolute and need source verification."),
    ("Table 3", "37", "Software requirements", "Exact package versions need updating to current manifests."),
    ("Table 4", "40", "Hardware specifications", "Hardware claims should be matched to actual local/dev deployment assumptions."),
    ("Table 5", "47", "Agile development phases", "Align terms with Figure 2 punctuation and phase wording."),
    ("Table 6", "50", "Functional requirements for admins and teachers", "Remove or qualify unsupported push/mobile parity claims."),
    ("Table 7", "52", "Functional requirements for students", "Fix 60%/74% threshold mismatch and 30-second completion claim."),
    ("Table 8", "55", "Technical and operational constraints", "Document AI-service fallback and demo-data limitations accurately."),
    ("Table 17", "66", "Use Case Narratives of Student Profile", "Title duplicates Table 18."),
    ("Table 18", "67", "Use Case Narratives of Student Profile", "Rename to the intended use case."),
    ("Table 24", "73", "Use Case Narratives of Disccusion Board (Teacher)", "Fix typo: Discussion."),
    ("Table 31", "80", "Use Case Narratives of Student Performance", "Conflicts with another Table 31 title in body text."),
    ("Table 31", "81", "Use Case Narratives of View Evaluations", "Renumber to restore sequence integrity."),
]


alignment_rows = [
    ("System title", "Nexora LMS/LXP title", "Same core title", "Matches repo branding and routes", "Match", "Minor"),
    ("Target institution", "Gat Andres Bonifacio High School", "Same", "Consistent in repo docs", "Match", "Minor"),
    ("Intervention threshold", "60% / c0%", "Below 74%", "Implemented at 74%", "Mismatch", "Critical"),
    ("AI architecture", "FastAPI + Ollama + RAG", "Lightweight AI integration", "FastAPI + Ollama + RAG + optional cloud fallback", "Partial", "Major"),
    ("Admin dashboard", "Claimed", "Claimed", "Implemented and live-verified", "Match", "Minor"),
    ("Teacher mobile", "Implied in mobile diagrams", "Broad mobile scope", "Teacher mobile explicitly unsupported in current build", "Mismatch", "Critical"),
    ("Discussion board mobile", "Implied in mobile scope", "Not precise", "Mobile shell note says not live-backed", "Mismatch", "Major"),
    ("Push notifications", "Immediate mobile push", "General notifications", "Web/in-app notifications only evidenced", "Mismatch", "Critical"),
    ("30-second lesson tracking", "Claimed", "Not strongly specified", "Unsupported by code evidence", "Mismatch", "Critical"),
    ("RAG grounding", "Claimed", "Weaker AI wording", "Implemented with embeddings and vector retrieval", "Match", "Minor"),
    ("Diagnostics", "Claimed", "Not emphasized", "Implemented and live-verified", "Match", "Minor"),
    ("Audit trail", "Claimed, but contaminated by legacy text", "Claimed generally", "Implemented and live-verified", "Partial", "Critical"),
]


commands_run = [
    "docker compose --env-file .env.compose ps",
    "Invoke-WebRequest http://localhost:3000/api/health/ready",
    "Invoke-WebRequest http://localhost:3000/api/docs",
    "Invoke-WebRequest http://localhost:3001",
    "Invoke-WebRequest http://localhost:11434/api/tags",
    "node next-frontend/scripts/nav-perf-smoke.js",
    "inline Playwright browser crawl against web routes with screenshots",
    "node backend/scripts/post-seed-smoke.js (first attempt from repo root failed due to postgres host resolution)",
    "node backend/scripts/post-seed-smoke.js (rerun from backend cwd succeeded)",
    "..\\.venv\\Scripts\\python.exe scripts/run_tests.py (from ai-service cwd)",
    "npm run build (backend)",
    "npm run build (next-frontend)",
    "npm run typecheck (test-mobile)",
    "npm run test (test-mobile)",
]


blockers = [
    "The stack was already running outside docker compose, so live proof came from existing local processes rather than a fresh compose startup.",
    "The first backend post-seed smoke run failed from the repo root because it resolved the wrong database host; rerunning from backend cwd fixed it.",
    "The first ai-service test invocation used the wrong relative .venv path; rerunning from ai-service cwd fixed it.",
    "Exact DOCX pagination is approximate outside the table/list references because raw text extraction does not preserve Word layout perfectly.",
]


assumptions = [
    "Approximate page references are based on the paper's list of figures/list of tables and extracted paragraph ordering when exact rendered pagination was not recoverable from raw text alone.",
    "When a feature existed in code but was not executed live in this audit, it is marked Real or Partial based on implementation evidence and explicitly noted as not live-verified.",
    "When a mobile feature had a screen shell but the repository itself stated the data source was not live-backed, it was treated as Partial rather than Real.",
]


top_5 = [f["id"] + " " + f["exact"] for f in findings[:5]]


def md_table(headers, rows):
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    for row in rows:
        lines.append("| " + " | ".join(str(cell).replace("\n", "<br>") for cell in row) + " |")
    return "\n".join(lines)


def build_paper_claims_md():
    lines = [
        "# Paper Claims Extracted",
        "",
        f"Source DOCX: `{summary['docx']}`",
        "",
        "## Extraction Summary",
        f"- Paragraphs: {summary['paragraph_count']}",
        f"- Tables: {summary['table_count']}",
        f"- Figure-caption hits: {summary['figure_caption_count']}",
        f"- Table-caption hits: {summary['table_caption_count']}",
        f"- Extracted media files: {summary['media_count']}",
        "",
        "## 1. System Identity and Title",
        '- "Nexora: A Learning Management With Learning Experience Platform Features for Targeted Student Intervention for Gat Andres Bonifacio High School"',
        "- Positions the project as a hybrid LMS plus LXP platform with AI-driven targeted intervention.",
        "",
        "## 2. Target Institution",
        "- Gat Andres Bonifacio High School.",
        "",
        "## 3. Target Users",
        "- Administrator",
        "- Teacher",
        "- Student",
        "",
        "## 4. Grade Levels and Subjects",
        "- Broad high-school-wide coverage is implied, often with language close to all subjects and all grade levels.",
        "- This is broader than the implementation evidence, which is bounded to grade levels 7-10 in the schema.",
        "",
        "## 5. Objectives",
        "- Build a web and mobile LMS with LXP intervention capabilities.",
        "- Identify at-risk students by assessment performance.",
        "- Gate remedial access to lower-performing students.",
        "- Provide an AI mentor to explain mistakes and guide remediation.",
        "",
        "## 6. Scope and Delimitation",
        "- LMS for standard class, lesson, assessment, and reporting workflows.",
        "- LXP reserved for struggling learners.",
        "- Mobile and web clients both claimed.",
        "",
        "## 7. LMS Features",
        "- Role-based login",
        "- Dashboards",
        "- Classes, sections, rosters, modules, lessons",
        "- Assessments and submissions",
        "- Discussion boards and announcements",
        "- Reports and records",
        "",
        "## 8. LXP Features",
        "- Targeted remedial access",
        "- Performance-based gating",
        "- Prior-lesson review and intervention support",
        "",
        "## 9. AI Features",
        "- JAKIPIR AI mentor / AI NPC mentor / AI tutor",
        "- PDF extraction",
        "- AI quiz drafting",
        "- AI-driven explanations and hints",
        "",
        "## 10. AI Model and RAG Claims",
        "- FastAPI microservice",
        "- Ollama runtime",
        "- qwen2.5:3b for tutoring/chat",
        "- gemma3:4b for PDF/document reasoning",
        "- pgvector-backed retrieval-augmented generation",
        "",
        "## 11. Authentication and Security Claims",
        "- OTP verification",
        "- Forgot password and account recovery",
        "- Audit trail",
        "- Protected dashboards and role routing",
        "",
        "## 12. Role Permissions",
        "- Admin manages users, sections, classes, diagnostics, audit, settings, reports",
        "- Teacher manages classes, modules, AI drafting, interventions, announcements, discussions",
        "- Student consumes lessons, takes assessments, accesses JA/LXP if eligible, manages profile",
        "",
        "## 13. Database and Storage Claims",
        "- PostgreSQL",
        "- pgvector",
        "- Redis",
        "- Cloud database wording appears once in a likely copied paragraph and conflicts with the rest of the paper",
        "",
        "## 14. Architecture Claims",
        "- Next.js frontend",
        "- NestJS backend",
        "- FastAPI AI microservice",
        "- Socket.IO, BullMQ, Redis, Swagger UI, OpenTelemetry, Prometheus, Loki",
        "",
        "## 15. Mobile App Claims",
        "- Expo SDK 54 mobile app",
        "- Login, OTP, forgot password, dashboard, modules, assessments, profile, LXP and AI mentor flows",
        "",
        "## 16. Web App Claims",
        "- Login and dashboard navigation",
        "- Course browsing and filtering",
        "- Assessment taking with out-of-focus warning",
        "",
        "## 17. Admin Portal Claims",
        "- User lifecycle management",
        "- Password resets",
        "- Sections and rosters",
        "- Diagnostics",
        "- Audit trail",
        "- Calendar and academic-state transitions",
        "",
        "## 18. Teacher Portal Claims",
        "- Classes and instructional materials",
        "- AI extraction review",
        "- AI quiz drafting",
        "- Interventions and outcomes",
        "",
        "## 19. Student Portal Claims",
        "- Dashboard and continue learning",
        "- Course/module browsing",
        "- Assessments",
        "- JA mentor access",
        "- Announcements and profile management",
        "",
        "## 20. Reports and Analytics Claims",
        "- Performance analytics",
        "- Quarterly trends",
        "- Heatmaps",
        "- Academic reports and export",
        "",
        "## 21. Monitoring and Observability Claims",
        "- Diagnostics",
        "- Dependency health checks",
        "- Audit trail",
        "- OpenTelemetry, Prometheus, Loki",
        "",
        "## 22. Testing and Evaluation Claims",
        "- Agile process and technical validation are implied.",
        "- The paper reads as if the workflows are broadly finished and ready.",
        "",
        "## 23. Diagrams and Workflows",
        "- Figure 1 through Figure 39 are listed.",
        "- Chapters 3-4 rely heavily on process flow and decomposition diagrams.",
        "",
        "## 24. Use Cases and Tables",
        "- Table 1 through Table 49 are listed.",
        "- There are visible duplicates and numbering conflicts in the extracted table list.",
        "",
        "## 25. Hardware and Software Requirements",
        "- Specific stack versions are listed, including Next.js 16.1.6, React 19, NestJS 11, Expo SDK 54, qwen2.5:3b, and gemma3:4b.",
        "",
        "## High-Risk Exact Strings Found",
        "- `c0 percent` / `c0%`",
        "- `RND` / `nutritionist` / `food intakes` / `cloud database` / `chat box feature is available on the web system`",
        "- `Disccusion`",
        "- `Non-Placer Character`",
        "- duplicated `Student Profile` table titles",
        "",
        "## Notes",
        "- This file intentionally extracts claims, not corrections.",
        "- Final truth judgments are in the companion audit outputs.",
    ]
    return "\n".join(lines) + "\n"


def build_repo_inventory_md():
    roles = ["admin", "teacher", "student"]
    lines = [
        "# Repository Feature Inventory",
        "",
        "## Actual System Identity",
        f"- Repository: `{repo_name}`",
        "- Product name in code and routes: Nexora",
        "- Primary apps: `backend/`, `next-frontend/`, `ai-service/`, `test-mobile/`",
        "",
        "## Actual Stack Versions",
        "- Web frontend: Next 16.2.4, React 19.2.3, React DOM 19.2.3",
        "- Backend: NestJS core ^11.0.1, Swagger ^11.2.6, BullMQ ^5.70.1, Socket.IO ^4.8.3",
        "- Mobile: Expo ~54.0.0, React Native 0.81.5, React 19.1.0",
        "- AI service: FastAPI app with Ollama integration and optional cloud fallback",
        "",
        "## Actual User Roles",
        *[f"- `{role}`" for role in roles],
        "",
        "## Actual Backend Surfaces",
        f"- Modules discovered: {len(repo_scan['backend']['modules'])}",
        "- Key modules: academic-state, admin, ai-mentor, analytics, announcements, assessments, audit, auth, class-record, class-templates, classes, content-modules, discussion-board, health, JA, lessons, LXP, notifications, OTP, performance, profiles, reports, roster-import, school-events, sections, teacher, users",
        f"- Controllers discovered: {len(repo_scan['backend']['controllers'])}",
        "",
        "## Actual Web Surface",
        f"- App routes discovered: {len(repo_scan['frontend']['routes'])}",
        "- Live-verified web routes during this audit: `/login`, `/forgot-password`, `/dashboard/admin`, `/dashboard/admin/diagnostics`, `/dashboard/admin/audit`, `/dashboard/teacher/classes`, `/dashboard/teacher/interventions`, `/dashboard/student`, `/dashboard/student/courses`, `/dashboard/student/ja`, `/dashboard/student/announcements`",
        "",
        "## Actual Mobile Surface",
        f"- Mobile screens discovered: {len(repo_scan['mobile']['screens']) if 'mobile' in repo_scan else 'see screen inventory'}",
        "- Current `test-mobile` build is student-focused.",
        "- Teacher mobile is explicitly marked unsupported in the app shell.",
        "",
        "## Actual Database Reality",
        f"- Schema tables discovered: {len(repo_scan['database']['tables']) if 'database' in repo_scan else '81'}",
        "- Key table families:",
        "- auth/users/roles/sections/classes/enrollments",
        "- lessons/modules/assessments/attempts/responses",
        "- announcements/notifications/discussion threads",
        "- class records and academic state",
        "- performance snapshots/logs/intervention cases/LXP progress",
        "- AI extraction, chunking, embeddings, generation jobs/outputs, interaction logs",
        "",
        "## Actual AI Reality",
        "- FastAPI service exists and was live during the audit.",
        "- Ollama models present live: `qwen2.5:3b`, `gemma3:4b`, `nomic-embed-text`.",
        "- Vector retrieval is implemented against `content_chunk_embeddings`.",
        "- Async extraction and quiz-generation endpoints exist.",
        "- Cloud fallback code exists, so the architecture is not purely local-only in design terms.",
        "",
        "## Actual Live Status During Audit",
        "- Backend already running on `localhost:3000`",
        "- Frontend already running on `localhost:3001`",
        "- AI service already running on `localhost:8000`",
        "- Postgres running on `5432` and Ollama on `11434`",
        "- Health endpoint returned database/redis/ai-service all OK",
        "",
        "## Partial or Missing Areas",
        "- Teacher mobile not implemented in the current `test-mobile` build",
        "- Mobile discussion board not backed by a live data source yet",
        "- No verified mobile push-notification stack found",
        "- No confirmed 30-second lesson-completion rule found",
        "- Some analytics/evaluation datasets were sparse in the live database",
        "",
        "## Live Data Counts Observed",
        "- users: 20",
        "- classes: 14",
        "- lessons: 31",
        "- assessments: 39",
        "- content chunks: 391",
        "- content chunk embeddings: 391",
        "- ai_generation_jobs: 44",
        "- ai_generation_outputs: 35",
        "- ai_interaction_logs: 20",
        "- intervention_cases: 7",
        "- audit_logs: 1214",
        "- notifications: 33",
        "- announcements: 2",
        "- school_events: 1",
        "- system_evaluations: 0",
        "",
        "## Setup Reality",
        "- The stack was already up outside docker compose; the audit therefore verified the active local runtime rather than a fresh compose bring-up.",
    ]
    return "\n".join(lines) + "\n"


def build_truth_table_md():
    headers = [
        "Feature",
        "Claimed in Research Paper",
        "Claimed in Concept Paper",
        "Found in Backend",
        "Found in Frontend Web",
        "Found in Mobile",
        "Found in Database",
        "Confirmed Live",
        "Status",
        "Evidence",
        "Required Correction in Paper",
        "Required Correction in System",
    ]
    rows = []
    for row in feature_rows:
        rows.append([
            row["feature"],
            row["paper"],
            row["concept"],
            row["backend"],
            row["web"],
            row["mobile"],
            row["db"],
            row["live"],
            row["status"],
            row["evidence"],
            row["paper_fix"],
            row["system_fix"],
        ])
    return "# Implementation Truth Table\n\n" + md_table(headers, rows) + "\n"


def build_diagram_audit_md():
    lines = [
        "# Chapter 4 Diagram Audit",
        "",
        "## Figure-by-Figure Audit",
        "",
        md_table(
            ["Figure", "Approx. Page", "Caption", "Status", "Audit Note"],
            figure_audit,
        ),
        "",
        "## Table Audit",
        "",
        md_table(
            ["Table", "Approx. Page", "Title", "Issue"],
            table_audit,
        ),
        "",
        "## Chapter 4 Summary",
        "- The most severe Chapter 4 problems are the 60%/74% mismatch, unsupported 30-second tracking, overclaimed mobile parity, push-notification wording, and duplicate/misaligned use-case tables.",
        "- The strongest supported Chapter 4 areas are web dashboard navigation, diagnostics, audit trail, JA/RAG architecture, class record flow, and admin lifecycle surfaces.",
    ]
    return "\n".join(lines) + "\n"


def build_live_demo_md():
    lines = [
        "# Live Demo Test Log",
        "",
        "## Runtime Context",
        "- Existing local processes were already running for backend, frontend, ai-service, Postgres, and Ollama.",
        "- Docker compose itself showed no active services during the audit, so the live proof came from the active local dev stack.",
        "",
        "## Commands Run",
    ]
    for command in commands_run:
        lines.append(f"- `{command}`")
    lines += [
        "",
        "## Key Results",
        "- `GET /api/health/ready` succeeded with database, redis, and aiService all healthy.",
        "- `GET /api/docs` returned 404, which confirms Swagger is not served from that path.",
        "- Frontend root and authenticated routes responded successfully.",
        "- Ollama tags confirmed `qwen2.5:3b`, `gemma3:4b`, and `nomic-embed-text` were present.",
        "- `next-frontend/scripts/nav-perf-smoke.js` completed and exercised admin, teacher, and student routes successfully.",
        "- Playwright crawl captured screenshots for major admin, teacher, and student pages with no console errors recorded.",
        "- `backend/scripts/post-seed-smoke.js` succeeded after rerun from the backend cwd.",
        "- `ai-service/scripts/run_tests.py` passed 60 tests after rerun from the ai-service cwd.",
        "- `npm run build` passed in backend and web; `npm run typecheck` and `npm run test` passed in `test-mobile`.",
        "",
        "## Important Runtime Findings",
        "- Admin diagnostics page is real and populated.",
        "- Admin audit trail page is real and populated.",
        "- Teacher interventions page explicitly shows the 74% trigger threshold.",
        "- Student JA Hub exists, but can sit on a loading state before rendering full content.",
        "",
        "## Blockers and Retries",
    ]
    for blocker in blockers:
        lines.append(f"- {blocker}")
    lines += [
        "",
        "## Screenshots Captured",
        "- `docs/research-paper-audit/screenshots/login.png`",
        "- `docs/research-paper-audit/screenshots/forgot-password.png`",
        "- `docs/research-paper-audit/screenshots/admin-dashboard.png`",
        "- `docs/research-paper-audit/screenshots/admin-diagnostics.png`",
        "- `docs/research-paper-audit/screenshots/admin-audit.png`",
        "- `docs/research-paper-audit/screenshots/teacher-dashboard.png`",
        "- `docs/research-paper-audit/screenshots/teacher-classes.png`",
        "- `docs/research-paper-audit/screenshots/teacher-interventions.png`",
        "- `docs/research-paper-audit/screenshots/student-dashboard.png`",
        "- `docs/research-paper-audit/screenshots/student-courses.png`",
        "- `docs/research-paper-audit/screenshots/student-ja.png`",
        "- `docs/research-paper-audit/screenshots/student-ja-delayed.png`",
        "- `docs/research-paper-audit/screenshots/student-announcements.png`",
    ]
    return "\n".join(lines) + "\n"


def build_panelist_md():
    lines = [
        "# Panelist Risk Checklist",
        "",
        "## Highest-Risk Contradictions",
        "- The paper says 60% or even `c0%`, but the implemented system uses 74%.",
        "- A copied paragraph about an RND, nutritionist, food intakes, and a cloud database appears in the technical chapter.",
        "- Mobile parity is overstated; teacher mobile is explicitly not ready in the current app.",
        "- Figure 30 claims immediate mobile push notifications without code evidence for a push stack.",
        "- Figure 11 claims a 30-second lesson completion rule that the audit could not verify in code.",
        "",
        "## Likely Panel Questions",
        "- Why does the paper say 60% when the live teacher intervention screen says 74%?",
        "- Why is there a nutritionist paragraph in a school LMS paper?",
        "- Can you demonstrate the teacher mobile app now?",
        "- How exactly are push notifications sent to mobile clients?",
        "- Where is the 30-second lesson completion rule enforced?",
        "- Why are Table 17 and Table 18 both Student Profile?",
        "- Which figure is correct for Table 31: Student Performance or View Evaluations?",
        "- Is AI processing strictly local, or can it fall back to cloud providers?",
        "",
        "## Best Defensive Fix Order",
        "1. Correct every threshold statement and diagram from 60/c0 to 74.",
        "2. Remove the copied nutritionist/RND paragraph entirely.",
        "3. Re-audit every Figure 8-39 caption and every Table 9-49 title/number.",
        "4. Narrow the mobile scope to student-first and remove unsupported push claims.",
        "5. Replace unsupported 30-second tracking language with real implemented logic.",
    ]
    return "\n".join(lines) + "\n"


def build_audit_notes_md():
    lines = [
        "# Audit Notes",
        "",
        "## Assumptions",
    ]
    for item in assumptions:
        lines.append(f"- {item}")
    lines += [
        "",
        "## Key Evidence Sources",
        "- Uploaded research paper DOCX and extracted text/media",
        "- `Concept Paper.pdf` plus extracted `Concept paper.txt`",
        "- `backend/package.json`, `next-frontend/package.json`, `test-mobile/package.json`",
        "- `docker-compose.yml`",
        "- `backend/src/drizzle/schema/*`",
        "- `backend/src/modules/*`",
        "- `ai-service/app/*`",
        "- `docs/research-paper-audit/repo_scan.json`",
        "- `docs/research-paper-audit/playwright-summary.json`",
        "",
        "## Repo Truth Anchors",
        "- Threshold anchor: `backend/src/modules/lxp/lxp.service.ts:35`",
        "- Threshold default anchor: `backend/src/drizzle/schema/performance.schema.ts:47`",
        "- Grade-level anchor: `backend/src/drizzle/schema/base.schema.ts:68`",
        "- Mobile teacher limitation: `test-mobile/src/screens/TeacherUnsupportedScreen.tsx:40,45`",
        "- Mobile discussion limitation: `test-mobile/src/screens/ClassDetailScreen.tsx:1228`",
        "- Swagger route anchor: `backend/src/main.ts:123-124`",
        "- Real-time notification anchor: `backend/src/modules/notifications/notifications.gateway.ts:13`",
        "- AI fallback anchor: `ai-service/app/config.py:51-87` and `ai-service/app/cloud_fallback.py`",
        "",
        "## Working Verdict",
        f"- Readiness score: {readiness_score}/100",
        f"- Panel risk: {panel_risk}",
        f"- Verdict: {paper_safe}",
        "",
        "## Severity Counts",
        f"- Critical: {severity_counts.get('Critical', 0)}",
        f"- Major: {severity_counts.get('Major', 0)}",
        f"- Moderate: {severity_counts.get('Moderate', 0)}",
        f"- Minor: {severity_counts.get('Minor', 0)}",
    ]
    return "\n".join(lines) + "\n"


def add_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    p.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_doc_defaults(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(22)
    styles["Title"].font.bold = True
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        styles[name].font.name = "Aptos"


def add_title_page(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Full Research Paper and System Consistency Audit")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(24, 55, 89)

    for line in [
        "System: Nexora",
        f"Repository: {repo_name}",
        f"Research paper file: {research_paper_file}",
        f"Concept paper file: {concept_paper_file}",
        f"Date/time generated: {now.strftime('%Y-%m-%d %H:%M:%S')}",
        "Auditor: Codex",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        p.add_run(line)
    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.color.rgb = RGBColor(24, 55, 89)
    return p


def add_bullets(doc: Document, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        p.add_run(item)


def add_simple_table(doc: Document, headers, rows, style="Table Grid"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = style
    hdr = table.rows[0]
    for idx, head in enumerate(headers):
        add_cell_text(hdr.cells[idx], head, bold=True)
    set_repeat_table_header(hdr)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            add_cell_text(cells[idx], value)
    doc.add_paragraph()
    return table


def finding_table_rows(selected_findings):
    rows = []
    for item in selected_findings:
        rows.append([
            item["id"],
            item["severity"],
            item["page"],
            item["target"],
            item["exact"],
            item["why"],
            "; ".join(item["evidence"][:3]),
            item["correction"],
        ])
    return rows


def build_docx():
    doc = Document()
    set_doc_defaults(doc)
    add_title_page(doc)

    add_heading(doc, "1. Executive Summary", 1)
    add_bullets(doc, [
        f"Readiness score: {readiness_score}/100",
        f"Panel defense risk: {panel_risk}",
        f"Verdict: {paper_safe} - {verdict}",
        f"Critical findings: {severity_counts.get('Critical', 0)}",
        f"Major findings: {severity_counts.get('Major', 0)}",
        f"Moderate findings: {severity_counts.get('Moderate', 0)}",
        f"Minor findings: {severity_counts.get('Minor', 0)}",
    ])
    add_heading(doc, "Top 10 Dangerous Issues", 2)
    add_bullets(doc, [
        "60% / c0% threshold in the paper contradicts the implemented 74% rule.",
        "Legacy RND / nutritionist / food intake paragraph is still in Chapter 3.",
        "Figure 30 claims immediate mobile push notifications without code evidence.",
        "Figure 11 claims 30-second lesson completion tracking without code evidence.",
        "Mobile scope is overstated; teacher mobile is explicitly unsupported.",
        "Use-case tables contain duplicate titles and numbering conflicts.",
        "The AI architecture is described as fully local-only, but the service contains optional cloud fallback.",
        "Figure 21 uses account-lock wording that does not match the implemented profile lock behavior.",
        "All-subject/all-grade wording is broader than the grade 7-10 schema reality.",
        "References include at least one source that looks mismatched to the topic and venue.",
    ])

    add_heading(doc, "2. System Reality Summary", 1)
    add_bullets(doc, [
        "Actual stack: Next 16.2.4, React 19.2.3, NestJS 11.x, FastAPI, PostgreSQL with pgvector, Redis, BullMQ, Socket.IO, Expo 54, React Native 0.81.5.",
        "Actual roles: admin, teacher, student.",
        "Actual major modules: auth, OTP, classes, lessons, assessments, discussion boards, announcements, class records, interventions/LXP, analytics, reports, diagnostics, audit, admin AI chat.",
        "Actual mobile status: student-first app exists; teacher mobile is not yet implemented in the current shell.",
        "Actual AI status: live FastAPI service, live Ollama models, vector retrieval, extraction, and async job flows are present.",
        "Actual incomplete/mock areas: teacher mobile, mobile live discussion parity, mobile push notifications, and the paper's 30-second tracking claim.",
    ])

    add_heading(doc, "3. Research Paper Claim Summary", 1)
    add_bullets(doc, [
        "The paper describes Nexora as a hybrid LMS plus LXP with targeted intervention, AI mentoring, and web/mobile parity.",
        "The concept paper broadly aligns on product direction, but uses a 74% intervention threshold and lighter AI wording.",
        "The repository confirms the core architecture and most web features, but disproves or weakens several high-specificity paper claims.",
    ])

    add_heading(doc, "4. Document Alignment Matrix", 1)
    add_simple_table(
        doc,
        ["Topic", "Research Paper", "Concept Paper", "Repo/Live Reality", "Status", "Severity"],
        alignment_rows,
    )

    add_heading(doc, "5. Critical Findings", 1)
    criticals = [item for item in findings if item["severity"] == "Critical"]
    add_simple_table(
        doc,
        ["ID", "Severity", "Page", "Figure/Table/Paragraph", "Exact Claim", "Why Wrong", "Evidence", "Correction"],
        finding_table_rows(criticals),
    )

    chapter_map = {
        "6. Chapter 1 Audit": ["F-001", "F-002", "F-008", "F-015"],
        "7. Chapter 2 Audit": ["F-017", "F-021"],
        "8. Chapter 3 Audit": ["F-003", "F-009", "F-010", "F-014", "F-018", "F-022"],
        "9. Chapter 4 Audit": ["F-004", "F-005", "F-006", "F-007", "F-011", "F-012", "F-013", "F-016", "F-020", "F-023"],
    }
    for heading_text, ids in chapter_map.items():
        add_heading(doc, heading_text, 1)
        selected = [item for item in findings if item["id"] in ids]
        add_simple_table(
            doc,
            ["ID", "Severity", "Page", "Figure/Table/Paragraph", "Exact Claim", "Why Wrong", "Evidence", "Correction"],
            finding_table_rows(selected),
        )

    add_heading(doc, "10. Chapter 4 Diagram and Figure Audit", 1)
    add_simple_table(
        doc,
        ["Figure", "Approx. Page", "Caption", "Status", "Mismatch / Correction"],
        figure_audit,
    )

    add_heading(doc, "11. Table Audit", 1)
    add_simple_table(
        doc,
        ["Table", "Approx. Page", "Title", "Issue / Correction"],
        table_audit,
    )

    add_heading(doc, "12. Database / ERD Audit", 1)
    add_bullets(doc, [
        "The database implementation is extensive and materially stronger than the paper documents in places.",
        "Core schema supports users, roles, classes, lessons, modules, assessments, attempts, responses, class records, academic state, notifications, announcements, discussion threads, interventions, AI logs, chunk embeddings, generation jobs, and concept mastery.",
        "The biggest paper-to-schema mismatch is not missing tables but incorrect business-rule narration, especially the 74% threshold and mobile scope.",
        "If the paper includes an ERD, it should be rechecked against the actual schema families in `backend/src/drizzle/schema/*`.",
    ])

    add_heading(doc, "13. Feature Implementation Audit", 1)
    short_truth_rows = []
    for row in feature_rows[:18]:
        short_truth_rows.append([row["feature"], row["status"], row["live"], row["paper_fix"]])
    add_simple_table(
        doc,
        ["Feature", "Status", "Confirmed Live", "Required Paper Correction"],
        short_truth_rows,
    )
    doc.add_paragraph("The full feature-by-feature matrix is provided in the companion `implementation_truth_table.md` file.")

    add_heading(doc, "14. Role and Permission Audit", 1)
    add_bullets(doc, [
        "Admin: supported for diagnostics, audit trail, user lifecycle, password reset, class templates, calendar, and chatbot.",
        "Teacher: supported for classes, modules, assessments, AI drafting, interventions, and discussions on web.",
        "Student: supported for web dashboard, classes/courses, assessments, JA/LXP, announcements, and profile.",
        "Mobile parity does not extend cleanly across all three roles; the current mobile app is student-first.",
    ])

    add_heading(doc, "15. UI/UX and Screenshot Audit", 1)
    add_bullets(doc, [
        "Admin dashboard, diagnostics, and audit pages are live and visually current.",
        "Teacher interventions page clearly shows the real 74% threshold, directly contradicting the paper's 60% wording.",
        "Student JA Hub exists but can remain on a loader before content appears, which is a demo-risk but not a missing feature.",
        "Any Chapter 4 screenshots or diagrams implying teacher mobile parity or push notifications should be replaced.",
    ])

    add_heading(doc, "16. API and Backend Audit", 1)
    add_bullets(doc, [
        "Backend controllers cover auth, OTP, users, admin, classes, sections, modules, lessons, assessments, class records, reports, analytics, performance, notifications, discussions, school events, academic state, profiles, JA, and LXP.",
        "Swagger is configured at `/api`, not `/api/docs`.",
        "WebSocket/Socket.IO support exists for notifications.",
        "Health and diagnostics are real and live-verified.",
    ])

    add_heading(doc, "17. AI and LXP Audit", 1)
    add_bullets(doc, [
        "JAKIPIR/JA mentor routes and backend flows are implemented.",
        "RAG is implemented with vector retrieval and live chunk embeddings in the database.",
        "PDF extraction and teacher quiz-job endpoints are present and tested at service level.",
        "The LXP/intervention rule is implemented at 74%, not 60%.",
        "The AI service also contains optional cloud fallback, so the architecture description should reflect that honestly.",
    ])

    add_heading(doc, "18. Testing and Evaluation Claims Audit", 1)
    add_bullets(doc, [
        "Observed during this audit: backend build pass, web build pass, mobile typecheck/test pass, ai-service test pass, live route smoke pass.",
        "Not every functional route was exercised end-to-end with mutation flows, so the paper should not imply universal live verification unless a fuller demo audit is performed.",
    ])

    add_heading(doc, "19. Grammar, Typo, and Formatting Issues", 1)
    typo_rows = [
        ["c0 percent / c0%", "Wrong threshold and obvious typo artifact", "60% if intended, but actually 74% in current system"],
        ["Non-Placer Character", "Wrong term", "Non-Player Character"],
        ["Disccusion", "Spelling error", "Discussion"],
        ["duplicate Student Profile table titles", "Broken numbering/title integrity", "Assign unique correct titles"],
        ["Figure 2. The Agile Development", "Inconsistent figure punctuation", "Standardize with the chosen caption format"],
    ]
    add_simple_table(doc, ["Exact Text", "Issue", "Correction"], typo_rows)

    add_heading(doc, "20. Panelist Risk Checklist", 1)
    add_bullets(doc, [
        "Paper says 60% / c0%, but system says 74%.",
        "Paper includes unrelated nutritionist/RND text.",
        "Diagram shows mobile push notifications without evidence.",
        "Diagram shows 30-second completion logic without evidence.",
        "Duplicate and conflicting use-case table titles make Chapter 4 look unfinished.",
    ])

    add_heading(doc, "21. Prioritized Fix Plan", 1)
    add_heading(doc, "Fix Before Defense", 2)
    add_bullets(doc, [
        "Correct all threshold wording from 60/c0 to 74.",
        "Delete the copied nutritionist/RND paragraph.",
        "Repair Chapter 4 figure and table numbering/title integrity.",
        "Remove unsupported push-notification and 30-second tracking claims.",
        "Narrow mobile scope to the student-first reality of the current build.",
    ])
    add_heading(doc, "Fix Before Final Submission", 2)
    add_bullets(doc, [
        "Update exact stack versions to match manifests.",
        "Standardize JAKIPIR/AI mentor terminology.",
        "Clarify local Ollama plus optional cloud fallback architecture.",
        "Recheck source credibility and citation formatting in Chapter 2.",
    ])
    add_heading(doc, "Nice to Fix", 2)
    add_bullets(doc, [
        "Improve JA Hub initial loading perception for demos.",
        "Seed fuller analytics/evaluation demo data.",
        "Add stronger live export and roster-import evidence if those features will be defended live.",
    ])

    add_heading(doc, "22. Appendix", 1)
    add_heading(doc, "Files Inspected", 2)
    add_bullets(doc, [
        "Uploaded research paper DOCX and extracted artifacts under `docs/research-paper-audit/extracted/`",
        "`Concept Paper.pdf` and `Concept paper.txt`",
        "`backend/`, `next-frontend/`, `test-mobile/`, `ai-service/`",
        "`docker-compose.yml`, package manifests, schema files, route files, screenshots, and smoke logs",
    ])
    add_heading(doc, "Commands Run", 2)
    add_bullets(doc, commands_run)
    add_heading(doc, "Assumptions and Limitations", 2)
    add_bullets(doc, assumptions + blockers)

    doc.save(OUTPUT_DOCX)


def main():
    for path in OUTPUT_MD.values():
        path.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    METADATA_JSON.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT_MD["audit_notes"].write_text(build_audit_notes_md(), encoding="utf-8")
    OUTPUT_MD["paper_claims"].write_text(build_paper_claims_md(), encoding="utf-8")
    OUTPUT_MD["repo_inventory"].write_text(build_repo_inventory_md(), encoding="utf-8")
    OUTPUT_MD["truth_table"].write_text(build_truth_table_md(), encoding="utf-8")
    OUTPUT_MD["diagram_audit"].write_text(build_diagram_audit_md(), encoding="utf-8")
    OUTPUT_MD["live_demo"].write_text(build_live_demo_md(), encoding="utf-8")
    OUTPUT_MD["panelist"].write_text(build_panelist_md(), encoding="utf-8")
    build_docx()
    METADATA_JSON.write_text(
        json.dumps(
            {
                "readiness_score": readiness_score,
                "panel_risk": panel_risk,
                "paper_safe": paper_safe,
                "severity_counts": severity_counts,
                "generated_at": now.isoformat(),
                "docx": str(OUTPUT_DOCX),
            },
            indent=2,
            default=lambda value: dict(value),
        ),
        encoding="utf-8",
    )


if __name__ == "__main__":
    main()
