from __future__ import annotations

from collections import OrderedDict
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\jethr\Desktop\capstone-nest-react-lms")
ARTIFACTS = ROOT / "artifacts"
ARTIFACTS.mkdir(exist_ok=True)
NOW = datetime.now()
PAPER_FILE = Path(r"C:\Users\jethr\Downloads\May03-CHAPTER-1-4-FINAL (1).pdf")
REPO_NAME = ROOT.name

READINESS_SCORE = 79
RISK_LEVEL = "High"

COUNTS = {
    "general": 30,
    "technical": 30,
    "research": 25,
    "pedagogy": 25,
    "ai": 25,
    "privacy": 20,
    "database": 20,
    "mobile_web": 20,
    "demo": 20,
    "deployment": 15,
    "trap": 15,
}


INSIGHTS = {
    "overview_honest": {
        "best": "Nexora is a prototype LMS with targeted LXP-style intervention features. Its strongest demonstrated value is the integration between class workflows, assessment performance, and remedial follow-up inside one school-focused system.",
        "short": "Nexora is a school-focused LMS prototype with targeted intervention and AI-assisted support, not a claim of full platform replacement.",
        "long": "Within the current capstone scope, Nexora demonstrates an integrated learning workflow rather than a generic portal. The LMS side handles classes, lessons, assessments, records, and role-based dashboards. The LXP side is narrower: it uses performance signals to open guided remedial access for struggling learners. That is the core contribution we can defend confidently.",
        "evidence": [
            "Concept paper: LMS plus LXP for targeted intervention with below-74% access control.",
            "Repo: backend modules for classes, lessons, assessments, performance, LXP, JA, reports, audit, and AI proxy.",
            "Repo: next-frontend role dashboards for admin, teacher, and student."
        ],
        "avoid": "Do not say Nexora fully transforms education or fully replaces existing school systems.",
        "incomplete": "The safest honest phrasing is that the system is implemented in prototype scope and demonstrates the intended intervention workflow under controlled school-centered conditions.",
        "followups": ["Which exact features are already stable?", "Which claims are still only partially validated?"],
    },
    "scope_narrow": {
        "best": "The safe defense scope is Grades 7 to 10 at Gat Andres Bonifacio High School, with deployment breadth kept narrower than the broadest wording in the paper.",
        "short": "We should defend the system as scoped to Grades 7 to 10 in one school context.",
        "long": "The repository enforces grade-level values of 7, 8, 9, and 10, so the strongest defensible scope is that range inside Gat Andres Bonifacio High School. Some paper wording still sounds like all subjects and all high-school-wide deployment have already been proven. For defense, we should present that as intended institutional scope, not as fully validated breadth.",
        "evidence": [
            "Repo: backend/src/common/utils/grade-level.util.ts restricts grade levels to 7, 8, 9, and 10.",
            "Repo: backend/src/drizzle/schema/base.schema.ts uses grade_level enum ['7','8','9','10'].",
            "Paper extract: still contains 'all subjects and grade levels' wording."
        ],
        "avoid": "Do not say the system has already been proven across every subject and every possible high-school deployment scenario.",
        "incomplete": "We can say the design targets Grades 7 to 10 and is structurally extensible, but broader validation remains future work.",
        "followups": ["Why not include senior high?", "How many subjects were actually seeded or demonstrated?"],
    },
    "threshold_74": {
        "best": "The system currently uses 74% as a configurable mastery cutoff for remedial access, and the defense should frame that as a project policy choice aligned to the intervention design, not as a universal educational law.",
        "short": "Nexora currently uses a configurable 74% threshold to trigger intervention.",
        "long": "In code, the threshold is consistently 74 across performance snapshots, intervention cases, and teacher performance views. The defensible answer is that 74% is the project's current mastery cutoff used to operationalize targeted intervention. We should not pretend that 74 is permanently optimal for every school; instead, we should say it should be validated further against school policy and remediation practice.",
        "evidence": [
            "Repo: backend/src/modules/lxp/lxp.service.ts sets INTERVENTION_THRESHOLD = 74.",
            "Repo: backend/src/drizzle/schema/performance.schema.ts defaults threshold_applied to 74.",
            "Paper extract: Figure 13 and multiple sections now use 74% wording."
        ],
        "avoid": "Do not claim 74% is scientifically perfect or permanently correct for every institution.",
        "incomplete": "If asked for stronger justification, say the system already supports threshold-based intervention logic and the exact cutoff should be refined with school policy and future validation data.",
        "followups": ["Why not 75?", "What happens at exactly 74%?", "Can teachers override it?"],
    },
    "lxp_defense": {
        "best": "The honest answer is that Nexora has LXP features for targeted remedial guidance, but it is not trying to compete with a full enterprise LXP. Its LXP claim rests on guided review, learner support, personalized remediation signals, and AI-assisted follow-up inside school scope.",
        "short": "Nexora is best defended as an LMS with LXP-style intervention features, not as a full standalone enterprise LXP.",
        "long": "A hostile panelist may say this looks like an LMS with gating rather than a true LXP. The safest response is to agree partly and narrow the claim: Nexora is primarily an LMS, but it adds LXP-style remedial experience through targeted access, personalized checkpoints, review paths, JA support, and intervention progress tracking. That is defensible. Calling it a complete LXP replacement would be harder to sustain.",
        "evidence": [
            "Concept paper positions the LXP as an intervention component rather than general student use.",
            "Repo: LXP tables include intervention_cases, intervention_assignments, and lxp_progress.",
            "Repo: student JA/LXP surfaces route students into guided review and replay flows."
        ],
        "avoid": "Do not insist it has every hallmark of a commercial LXP if the panel challenges personalization depth.",
        "incomplete": "If pushed, say the current capstone scope focuses on intervention-oriented LXP features rather than broad enterprise personalization.",
        "followups": ["Where is learner autonomy?", "What makes this more than a remedial tab?"],
    },
    "ai_grounded": {
        "best": "JAKIPIR is assistive and grounded, not infallible. The safe defense is that it uses class evidence, retrieval, policy controls, and teacher oversight to reduce hallucination risk, while still acknowledging that AI output must be reviewed critically.",
        "short": "JAKIPIR is grounded and policy-controlled, but not claimed as perfectly accurate.",
        "long": "The code shows real retrieval, lesson bias, pgvector-backed similarity search, AI interaction logs, and class-level AI policies such as source scope and strict grounding. That supports a strong answer that the mentor is designed to stay within visible class material and avoid unsupported guesses. We should still say AI output is assistive, not an official grading authority, and that teacher oversight remains necessary.",
        "evidence": [
            "Repo: ai-service/app/retrieval_service.py implements vector retrieval and source filtering.",
            "Repo: backend/src/drizzle/schema/lxp.schema.ts defines class_ai_policies including strictGrounding and sourceScope.",
            "Repo: next-frontend teacher interventions page surfaces AI policy controls."
        ],
        "avoid": "Do not say the AI never hallucinates or always gives correct pedagogy.",
        "incomplete": "If reliability is challenged, say the system is designed to reduce hallucination risk through grounding and policy controls, but teacher review remains part of safe use.",
        "followups": ["What prevents cheating help?", "What happens when evidence is weak?", "Can the teacher disable AI?"],
    },
    "extraction_review": {
        "best": "The extraction feature is real, but it should be defended as teacher-reviewed AI assistance rather than one-click perfect content conversion.",
        "short": "PDF extraction exists, but the safe claim is teacher-reviewed AI-assisted drafting.",
        "long": "The strongest defense is that Nexora automates the first drafting pass from uploaded materials and then expects teachers to review, edit, and apply the output. That matches both technical reality and safer academic language. We should avoid describing extraction as a fully autonomous curriculum authoring engine.",
        "evidence": [
            "Repo: ai-service README and extraction pipeline implement extract, status, patch, and apply flows.",
            "Repo: backend/src/drizzle/schema/ai-mentor.schema.ts stores extracted_modules with statuses and isApplied flag.",
            "Repo: teacher extraction and AI draft routes exist in next-frontend."
        ],
        "avoid": "Do not say uploaded PDFs are always parsed perfectly, especially for scanned or messy files.",
        "incomplete": "If extraction quality is questioned, say the feature accelerates teacher preparation but still depends on teacher review before applied lesson content becomes part of class workflow.",
        "followups": ["What if the PDF is scanned?", "Can formulas and images fail?", "Can teachers reject AI output?"],
    },
    "mobile_partial": {
        "best": "The mobile story should be defended carefully: student workflows are the strongest, teacher mobile has meaningful surfaces in the current codebase, and admin mobile is still limited.",
        "short": "Mobile exists, but parity is not equal across roles.",
        "long": "The current test-mobile app contains real student flows, auth recovery screens, JA/LXP access, assessments, and profile screens. It also includes teacher navigation and detail screens, while admin routes remain placeholder-level. The safest defense is to present mobile as role-asymmetric rather than claiming full parity with the web system.",
        "evidence": [
            "Repo: test-mobile/src/navigation/types.ts includes Login, VerifyEmail, ForgotPassword, ResetPassword, SetInitialPassword.",
            "Repo: test-mobile/src/navigation/AppNavigator.tsx includes teacher tabs and teacher detail screens.",
            "Repo: test-mobile/src/screens/RoleWorkspaceScreen.tsx shows admin mobile placeholder sections."
        ],
        "avoid": "Do not say all web features are fully available on mobile.",
        "incomplete": "If asked about missing parity, say the capstone prioritized student mobile access first and role expansion remains staged.",
        "followups": ["Can admin work fully on mobile?", "Can teacher do every class action on mobile?"],
    },
    "security_controls": {
        "best": "The repository shows real security controls such as JWT guards, OTP flows, hashed OTP storage, validation pipes, throttling, CORS rules, and audit logging, but compliance claims should still stay practical and not absolute.",
        "short": "Security is implemented with concrete controls, but we should avoid absolute guarantees.",
        "long": "The backend uses a global JWT guard, a global throttler guard, validation pipes, CORS allowlists, and role-based route protection. OTP codes are hashed with HMAC-SHA256 and never stored in plaintext, and audit logs exist for administrative accountability. That is a strong technical defense. The safer framing is that the project applies core security controls and privacy-aware design, while formal compliance review and institutional policy rollout would still be required before production deployment.",
        "evidence": [
            "Repo: backend/src/app.module.ts wires global JwtAuthGuard and throttling.",
            "Repo: backend/src/main.ts configures validation, helmet, CORS, and cookie parsing.",
            "Repo: backend/src/modules/otp/otp.service.ts hashes OTP codes and enforces resend limits.",
            "Repo: backend/src/modules/audit/audit.service.ts writes audit log entries."
        ],
        "avoid": "Do not say the system is unhackable or fully compliant by thesis claim alone.",
        "incomplete": "If privacy law questions go deeper, say the system is built with role-based access, protected credentials, and auditability, but institutional deployment would still require policy, consent, and governance review.",
        "followups": ["How are refresh tokens handled?", "How do you protect minors' data?", "Are AI logs separated from grades?"],
    },
    "methodology_limited": {
        "best": "The main methodological weakness is not that the system is empty, but that implementation depth currently exceeds the strength of the measured evaluation evidence. The answer should admit that distinction directly.",
        "short": "Our implementation evidence is stronger than our formal outcome evidence, so we should avoid claiming proven educational effectiveness.",
        "long": "A strong defense answer is that this capstone primarily validates feasibility, workflow integration, and prototype functionality. The repository includes a system_evaluations feature and the paper promises evaluation dimensions such as usability, functionality, reliability, and portability, but that does not automatically prove actual learning gains yet. We should say the current study demonstrates system design and prototype behavior, while larger-scale educational impact evaluation remains future work.",
        "evidence": [
            "Paper extract: repeatedly promises functionality, reliability, usability, and portability evaluation.",
            "Repo: backend/src/drizzle/schema/lxp.schema.ts defines system_evaluations.",
            "Repo audits previously observed sparse or absent live evaluation records."
        ],
        "avoid": "Do not say the platform is already proven effective in improving school outcomes unless you have respondent data and statistical treatment ready.",
        "incomplete": "The honest fallback is that the capstone validates the system artifact and workflow logic first, while large-sample outcome validation should be treated as a subsequent study.",
        "followups": ["How many respondents did you have?", "What statistics did you compute?", "Where are your measured results?"],
    },
    "deployment_limited": {
        "best": "Nexora is better defended as deployable in pilot conditions than as already ready for unrestricted production rollout.",
        "short": "It is demo-capable and pilot-oriented, not something we should overstate as school-wide production-ready.",
        "long": "The repository has Docker, observability, pgvector, Redis, and AI service integration, which is more mature than a typical capstone. But live deployment claims should still stay modest because uptime, cost, training, operations, and policy approval are separate problems from code completeness. The safest answer is that the system is technically deployable for a controlled pilot and designed with production-minded components, but not yet defended as fully deployed institutional infrastructure.",
        "evidence": [
            "Repo: root docker-compose.yml provisions PostgreSQL, Redis, Ollama, backend, frontend, and monitoring stack.",
            "Repo: README documents observability, environment variables, and deployment notes.",
            "Current runtime check: frontend and ai-service reachable, backend not currently listening."
        ],
        "avoid": "Do not say it is already production-ready for any public school without operational validation.",
        "incomplete": "If asked about readiness, say the architecture is deployment-minded, but the defense claim is prototype readiness for controlled pilot use.",
        "followups": ["What hardware is required?", "Who will maintain it?", "What if Ollama is slow?"],
    },
    "backend_runtime_risk": {
        "best": "The codebase is broad, but the current local runtime is not fully healthy because backend port 3000 was unreachable during this audit. That makes demo discipline essential.",
        "short": "The biggest immediate demo risk is runtime stability, not absence of code.",
        "long": "During this run, the frontend on port 3001 and the AI service on port 8000 were reachable, but the backend on port 3000 was not. That means I could not rely on a full end-to-end live sweep and had to combine runtime checks with static evidence. For defense, the team should treat backend startup and seeded-auth verification as must-fix items before demo day.",
        "evidence": [
            "Current run: localhost:3001 returned HTTP 200.",
            "Current run: localhost:8000/ready returned ready with Ollama models available.",
            "Current run: localhost:3000/api/health/live and /ready were unreachable."
        ],
        "avoid": "Do not walk into defense saying the whole stack is already stable without rechecking ports, seeded logins, and backend health on the actual machine.",
        "incomplete": "If the backend is unstable, say the repository implementation is present but the local demo environment needs startup verification before presentation.",
        "followups": ["Can you show the health checks?", "What is your fallback if backend fails?"],
    },
    "not_official_records": {
        "best": "AI and LXP outputs are assistive and intentionally separated from official academic records.",
        "short": "The system keeps AI guidance separate from official grading records.",
        "long": "This is one of the safest technical answers in the repo. The code and schema separate official class-record behavior from AI and intervention surfaces, which is important ethically and defensively. The panel should hear clearly that AI recommendations do not directly overwrite official grades.",
        "evidence": [
            "Repo guidance: backend and ai-service AGENTS emphasize that AI features must not mutate official academic records.",
            "Repo: class-record, LXP, and AI logging are separate schema areas."
        ],
        "avoid": "Do not imply the AI directly changes grades or decides final marks.",
        "incomplete": "If pushed, say intervention feedback informs support decisions, while teachers and official records remain authoritative.",
        "followups": ["Can AI modify scores?", "Can a remedial result alter the official record automatically?"],
    },
    "fairness_access": {
        "best": "Restricting the remedial path to low-performing students is a design choice for targeted intervention, but it is not the same as saying stronger students could never benefit from similar supports.",
        "short": "The current rule targets scarce remedial support, but the design could be widened later.",
        "long": "A fair answer acknowledges the tradeoff. Nexora uses targeted access because the capstone problem is intervention for struggling learners, not enrichment for everyone. That is consistent with the paper and the LXP logic. If a panelist asks whether high performers might also benefit, the correct answer is yes, but that broader personalization model is outside the current intervention-focused scope.",
        "evidence": [
            "Concept paper and paper scope both define the LXP as an intervention component for selected students.",
            "Repo LXP logic is tied to at-risk and threshold-based case management."
        ],
        "avoid": "Do not argue that only struggling learners deserve personalized support.",
        "incomplete": "If challenged, say the current capstone scope focuses on targeted remediation first and can be expanded later for enrichment use cases.",
        "followups": ["Is that fair to high achievers?", "Could the threshold stigmatize students?"],
    },
    "reports_seeded": {
        "best": "Reports and analytics exist in code, but presentation quality depends on the available seeded or live data at demo time.",
        "short": "The reporting features are real, but their persuasiveness depends on data completeness.",
        "long": "The repo contains reports, exports, analytics, performance snapshots, and evaluation routes. However, these screens are only as strong as the data loaded into them. The safest defense is to claim the reporting workflow is implemented, while also preparing realistic seeded data so charts and exports do not look empty or trivial during demo.",
        "evidence": [
            "Repo: reports, analytics, performance, and evaluations modules exist across backend and frontend.",
            "Prior repo audits noted sparse datasets in some analytics areas."
        ],
        "avoid": "Do not imply the analytics have already been validated over long-term real school usage.",
        "incomplete": "If a chart looks sparse, say the workflow is implemented and currently demonstrated with controlled development data.",
        "followups": ["Is this real school data?", "How many records feed these charts?"],
    },
    "admin_mobile_placeholder": {
        "best": "Admin mobile support should be treated as limited. The current mobile codebase includes placeholder-style admin workspace sections rather than full parity.",
        "short": "Admin mobile is not a defense centerpiece.",
        "long": "The test-mobile navigator resolves admin roles, but the admin tabs point to a generic RoleWorkspaceScreen rather than full operational admin workflows. That is not a fatal capstone issue because the project is still web-first for administration, but it becomes a problem only if the team overclaims mobile parity.",
        "evidence": [
            "Repo: test-mobile/src/navigation/AppNavigator.tsx maps admin tabs to RoleWorkspaceScreen.",
            "Repo: test-mobile/src/screens/RoleWorkspaceScreen.tsx contains placeholder explanatory text."
        ],
        "avoid": "Do not offer to demo admin mobile unless explicitly required and clearly framed as limited.",
        "incomplete": "If asked, say administration remains strongest on web in the current scope.",
        "followups": ["Why include admin mobile at all?", "Is it functional or just a shell?"],
    },
    "teacher_mobile_partial": {
        "best": "Teacher mobile is materially better than a placeholder because teacher tabs and detail screens exist, but it still should not be described as full web parity unless verified live.",
        "short": "Teacher mobile exists in code, but it is still safer to treat web as the primary teacher surface.",
        "long": "The current test-mobile app includes teacher home, classes, assessments, announcements, profile, and detail screens in its navigator. That is a meaningful capability increase over older repo states. Still, without a live teacher walkthrough in this run, the defense-safe position is that teacher workflows exist on mobile in prototype form while the web interface remains the main teacher workspace.",
        "evidence": [
            "Repo: AppNavigator includes TeacherTabs and TeacherNavigator.",
            "Repo: teacher screens exist under test-mobile/src/screens."
        ],
        "avoid": "Do not promise full feature parity until you have live proof on the defense device.",
        "incomplete": "Use 'teacher mobile workflows are present in prototype scope' if you cannot verify every flow live.",
        "followups": ["Which teacher actions are proven?", "What remains web-only?"],
    },
}


@dataclass
class Question:
    question: str
    why: str
    risk: str
    insight: str
    evidence_extra: list[str] | None = None
    followups_extra: list[str] | None = None


@dataclass
class Weakpoint:
    title: str
    description: str
    why_dangerous: str
    probability: str
    impact: str
    severity: str
    verbal_defense: str
    fix_before_defense: str
    mention_strategy: str


def q(question: str, why: str, risk: str, insight: str, evidence_extra: Iterable[str] | None = None, followups_extra: Iterable[str] | None = None) -> Question:
    return Question(question, why, risk, insight, list(evidence_extra or []), list(followups_extra or []))


def render_question(item: Question) -> dict[str, object]:
    insight = INSIGHTS[item.insight]
    evidence = list(insight["evidence"])
    if item.evidence_extra:
        evidence.extend(item.evidence_extra)
    followups = list(insight["followups"])
    if item.followups_extra:
        followups.extend(item.followups_extra)
    return {
        "question": item.question,
        "why": item.why,
        "risk": item.risk,
        "best_safe_answer": insight["best"],
        "short_answer": insight["short"],
        "long_answer": insight["long"],
        "evidence": evidence,
        "what_not_to_say": insight["avoid"],
        "incomplete_answer": insight["incomplete"],
        "followups": followups,
    }


GENERAL_QUESTIONS = [
    q("What is Nexora in one sentence?", "The panel will test whether the team can explain the system simply and consistently.", "High", "overview_honest"),
    q("What exact school problem does Nexora solve?", "They want a concrete problem, not a vague EdTech mission statement.", "High", "overview_honest"),
    q("Why focus on Gat Andres Bonifacio High School specifically?", "Institutional fit is a common capstone scrutiny point.", "High", "scope_narrow"),
    q("What makes Nexora different from a normal LMS?", "This is the core novelty challenge.", "High", "lxp_defense"),
    q("Why combine LMS and LXP instead of building only one system?", "They want the architectural and product rationale.", "High", "lxp_defense"),
    q("Who are the primary users of the system?", "They are checking role clarity and feature alignment.", "Medium", "overview_honest"),
    q("What is the strongest contribution of the project?", "They want to hear a focused contribution, not a long feature dump.", "High", "overview_honest"),
    q("What is your real scope, as opposed to your ideal scope?", "This tests whether the team can defend limitations cleanly.", "High", "scope_narrow"),
    q("What is explicitly outside the current capstone scope?", "Good panelists expect a disciplined boundary.", "Medium", "deployment_limited"),
    q("Why did you include AI at all?", "They want to see whether AI is necessary or just fashionable.", "High", "ai_grounded"),
    q("Does Nexora replace teachers?", "This is both a pedagogy and ethics challenge.", "High", "not_official_records"),
    q("What is the intervention trigger in Nexora?", "A panelist will test whether the team knows the system rule cold.", "High", "threshold_74"),
    q("Why is the threshold 74 and not 75 or 70?", "This is one of the most predictable defense questions.", "Critical", "threshold_74"),
    q("What happens to a student who scores exactly 74%?", "They are probing rule precision.", "High", "threshold_74", followups_extra=["Is the comparison below 74 or at or below 74?"]),
    q("What happens to a student who scores below the threshold?", "They want a concrete workflow answer.", "High", "threshold_74"),
    q("Can high-performing students also benefit from the LXP?", "This tests fairness and pedagogy awareness.", "High", "fairness_access"),
    q("Why call the assistant JAKIPIR?", "They are probing branding versus actual function.", "Medium", "ai_grounded"),
    q("Is JAKIPIR an actual NPC mentor or mainly a guided chatbot?", "This exposes overclaim risk around the AI persona.", "High", "ai_grounded"),
    q("What part of the system is already strongest today?", "Panels often want the team to identify its most defensible surface.", "Medium", "overview_honest"),
    q("What part of the system still needs the most improvement?", "This tests honesty and technical maturity.", "High", "backend_runtime_risk"),
    q("If the AI is removed, what still remains valuable?", "This checks whether the project still has educational value without hype.", "Medium", "overview_honest"),
    q("Why is this still a capstone and not an overambitious startup pitch?", "They are looking for scope realism.", "High", "scope_narrow"),
    q("What objective does Nexora satisfy best right now?", "They want prioritization instead of a kitchen-sink answer.", "Medium", "overview_honest"),
    q("What claim about Nexora should be phrased most carefully during defense?", "This tests strategic communication.", "High", "deployment_limited"),
    q("What are the biggest limitations you would admit up front?", "Panels reward honest but controlled answers.", "High", "methodology_limited"),
    q("Why not just recommend Google Classroom plus a separate tutor bot?", "This attacks the need for integration.", "High", "overview_honest"),
    q("What school process becomes easier because of Nexora?", "They want a workflow-level benefit, not just a feature list.", "Medium", "overview_honest"),
    q("Why is targeted intervention better than generic review for everyone in this project?", "This probes the intervention philosophy.", "High", "fairness_access"),
    q("What evidence shows the project is more than mock UI?", "They want proof of implementation depth.", "High", "overview_honest", evidence_extra=["Repo: backend, next-frontend, ai-service, and test-mobile all have substantive modules and routes."]),
    q("If the panel remembers only one sentence, what should it be?", "This reveals the maturity of the team's narrative.", "Medium", "overview_honest"),
]

TECHNICAL_QUESTIONS = [
    q("What is the actual architecture of Nexora?", "A technical panelist will test whether the architecture matches the repo.", "Critical", "overview_honest", evidence_extra=["Repo: root README and docker-compose.yml show Next frontend, Nest backend, FastAPI AI service, PostgreSQL, Redis, Ollama."]),
    q("Why is the backend NestJS instead of Next.js API routes?", "They want evidence of intentional architecture separation.", "High", "overview_honest"),
    q("How does the web frontend communicate with the backend?", "API boundary clarity is a baseline technical expectation.", "Medium", "security_controls"),
    q("How does the backend communicate with the AI service?", "They want to see whether the AI service is truly separate.", "High", "ai_grounded", evidence_extra=["Repo: backend/src/modules/ai-mentor/ai-mentor.controller.ts forwards to AI proxy endpoints."]),
    q("Where is the intervention threshold enforced in code?", "This checks whether the team can tie business rules to implementation.", "High", "threshold_74"),
    q("How are background AI jobs processed?", "BullMQ and async workflow claims invite inspection.", "High", "extraction_review"),
    q("Why do you need Redis in this project?", "They are testing whether dependencies are justified.", "Medium", "overview_honest", evidence_extra=["Repo: backend uses BullMQ and Redis-backed coordination."]),
    q("Where is pgvector actually used?", "This is a classic 'you claimed RAG, prove it' question.", "Critical", "ai_grounded", evidence_extra=["Repo: ai-service/app/retrieval_service.py CASTs embeddings AS vector and searches content_chunk_embeddings."]),
    q("What does Drizzle ORM do in the backend?", "They want to know whether the team understands its own persistence layer.", "Medium", "overview_honest", evidence_extra=["Repo: backend/drizzle config and schema files define typed schema and migrations."]),
    q("How are JWT and role checks enforced?", "This is a core API security question.", "High", "security_controls"),
    q("How is OTP handled securely?", "Panels often target credential recovery flows.", "High", "security_controls"),
    q("How do you prevent AI outputs from modifying official grades?", "This checks domain boundaries.", "Critical", "not_official_records"),
    q("How are real-time notifications implemented?", "Socket.IO is an explicit architecture claim.", "Medium", "overview_honest", evidence_extra=["Repo: backend dependencies and notifications gateway indicate real-time notification support."]),
    q("What happens if the AI service is down?", "This is a demo and resilience question.", "High", "deployment_limited"),
    q("Is your AI stack strictly local-only?", "The panel may challenge architecture wording versus reality.", "High", "deployment_limited", evidence_extra=["Repo inventory and docs note optional cloud fallback code exists in addition to local Ollama."]),
    q("How are extracted lessons reviewed before becoming real content?", "This tests safety of automation.", "High", "extraction_review"),
    q("How do you avoid blocking the UI during long AI tasks?", "Async UX and system design matter here.", "Medium", "extraction_review", evidence_extra=["Repo: queued jobs, extraction status polling, and async AI job endpoints exist."]),
    q("What observability tools are actually included?", "They want to know if OpenTelemetry/Prometheus/Loki are real or decorative.", "High", "deployment_limited", evidence_extra=["Repo: docker-compose provisions Prometheus, Loki, Tempo, Grafana, promtail, exporters."]),
    q("What is the difference between liveness and readiness in your backend?", "This is a crisp architecture comprehension test.", "Medium", "deployment_limited", evidence_extra=["Repo: backend/src/modules/health/health.controller.ts exposes /health/live and /health/ready with dependency-aware readiness."]),
    q("How do you validate request payloads?", "This reveals baseline API discipline.", "Medium", "security_controls"),
    q("How do you separate admin, teacher, and student routes on the web?", "They want proof of role-aware app structure.", "Medium", "security_controls", evidence_extra=["Repo: next-frontend uses /dashboard/admin, /dashboard/teacher, /dashboard/student route conventions."]),
    q("How do you store AI chats and extraction logs?", "A technical panelist may target logging and audit design.", "High", "security_controls", evidence_extra=["Repo: backend/src/drizzle/schema/ai-mentor.schema.ts defines ai_interaction_logs and extracted_modules tables."]),
    q("What is the purpose of class AI policies?", "This is a good test of teacher oversight mechanisms.", "High", "ai_grounded"),
    q("How do you keep the mobile app aligned with backend contracts?", "This is a cross-platform maintainability question.", "Medium", "mobile_partial"),
    q("How does your retrieval stay inside allowed class material?", "Panels will test RAG containment claims.", "Critical", "ai_grounded"),
    q("How do you handle failed extractions or weak evidence?", "Failure-path design is a common technical challenge.", "High", "extraction_review"),
    q("How are reports and analytics generated from the database?", "This tests whether analytics are merely visual shells.", "Medium", "reports_seeded"),
    q("Why use a separate FastAPI service for AI instead of embedding everything in NestJS?", "This is a practical architecture rationale question.", "High", "overview_honest"),
    q("Which service currently looks most fragile from a demo-day perspective?", "A fair technical panelist may ask where the real operational risk lives.", "High", "backend_runtime_risk"),
    q("If you had one more week, what technical hardening would you do first?", "They want prioritization under pressure.", "High", "backend_runtime_risk"),
]

RESEARCH_QUESTIONS = [
    q("What is the exact research gap your study addresses?", "A research panelist wants a gap, not just a feature wish list.", "Critical", "overview_honest"),
    q("How do you justify the need for Nexora in the chosen school context?", "This tests whether the problem is locally grounded.", "High", "scope_narrow"),
    q("Why is your problem statement not too broad?", "Broad problem statements are a classic capstone weakness.", "High", "scope_narrow"),
    q("How do your objectives map to your implemented system?", "They want objective-to-artifact traceability.", "High", "overview_honest"),
    q("Which objective is fully demonstrated and which objective is still only partially validated?", "This tests methodological honesty.", "High", "methodology_limited"),
    q("Why did you choose Agile as your development methodology?", "Methodology choice must be defendable, not copied.", "Medium", "overview_honest"),
    q("What actual Agile evidence do you have beyond the diagram?", "They will probe whether Agile was truly followed.", "High", "methodology_limited"),
    q("How do you measure functionality, reliability, usability, and portability?", "The paper explicitly promises these dimensions.", "Critical", "methodology_limited"),
    q("How many respondents evaluated the system?", "A research panelist will ask this quickly.", "Critical", "methodology_limited"),
    q("What instrument did you use for evaluation?", "They want instrument rigor.", "High", "methodology_limited"),
    q("What statistical treatment did you apply to the evaluation data?", "This is a standard defense question.", "Critical", "methodology_limited"),
    q("How do you separate system implementation success from educational effectiveness?", "Strong panels distinguish artifact quality from learning outcomes.", "High", "methodology_limited"),
    q("How did you validate the 74% threshold academically?", "This crosses methodology and pedagogy.", "High", "threshold_74"),
    q("What are your delimitations and why are they reasonable?", "Scope discipline is a research strength.", "Medium", "scope_narrow"),
    q("Why are offline mode and third-party integrations excluded?", "Exclusions need rationale, not excuses.", "Medium", "deployment_limited"),
    q("How current and credible are your cited statistics?", "Literature quality is a likely attack line.", "High", "methodology_limited"),
    q("How do you defend the claim that existing LMS platforms are insufficient?", "This tests whether the gap is proven rather than assumed.", "High", "lxp_defense"),
    q("What is the significance of the study to teachers?", "Stakeholder significance must tie to actual workflows.", "Medium", "overview_honest"),
    q("What is the significance of the study to students?", "They want stakeholder-specific value.", "Medium", "overview_honest"),
    q("What is the significance of the study to the school?", "Institution-level value must be realistic.", "Medium", "overview_honest"),
    q("Why is your Chapter 4 not just design documentation?", "This challenges 'Results and Discussion' inflation.", "Critical", "methodology_limited"),
    q("What limitation would you highlight if asked about external validity?", "A good panelist will test transferability limits.", "High", "scope_narrow"),
    q("How do you defend your related-systems comparison table?", "Comparative tables are frequent overclaim zones.", "High", "methodology_limited"),
    q("What future study should follow this capstone?", "They want to see whether you know the next research step.", "Medium", "methodology_limited"),
    q("If the panel says your implementation is stronger than your study design, how do you answer?", "This is a likely accurate criticism.", "Critical", "methodology_limited"),
]

PEDAGOGY_QUESTIONS = [
    q("Why do you call this an LXP and not just an LMS with remediation?", "This is the most likely pedagogy challenge.", "Critical", "lxp_defense"),
    q("What pedagogical theory supports the intervention flow?", "They want theory behind system behavior.", "High", "threshold_74"),
    q("How is mastery learning reflected in your design?", "This probes depth of educational reasoning.", "High", "threshold_74"),
    q("Why is the remedial path limited to struggling students?", "This is a fairness question.", "High", "fairness_access"),
    q("Could restricting access create stigma?", "A pedagogy panelist may challenge the student experience.", "High", "fairness_access"),
    q("How does Nexora support learner autonomy?", "This is a common LXP criterion.", "High", "lxp_defense"),
    q("How does the system personalize remediation?", "Personalization claims must be defensible.", "High", "lxp_defense"),
    q("What role does the teacher still play after intervention is triggered?", "This tests whether the design remains teacher-centered.", "High", "not_official_records"),
    q("How does the system avoid over-reliance on AI during learning?", "This is a pedagogy plus ethics challenge.", "High", "ai_grounded"),
    q("Why should a student trust JAKIPIR?", "Trust and learning relationship matter in pedagogy defense.", "High", "ai_grounded"),
    q("How do students revisit previous lessons and assessments in a meaningful way?", "They want concrete remedial learning design.", "Medium", "lxp_defense"),
    q("Is the LXP designed for remediation only or also enrichment?", "Scope of pedagogy matters.", "Medium", "fairness_access"),
    q("How do you prevent students from using the AI only to get answers?", "Academic integrity challenge.", "High", "ai_grounded"),
    q("How does the system support 'Catch-Up Fridays' or similar school remediation periods?", "This ties project rhetoric to actual practice.", "Medium", "overview_honest"),
    q("What evidence do you have that the intervention path is pedagogically coherent?", "They want more than feature chaining.", "High", "methodology_limited"),
    q("Why is immediate targeted remediation better than waiting for manual teacher follow-up only?", "This tests the educational rationale for automation.", "High", "overview_honest"),
    q("How are weak concepts identified?", "They want to connect analytics with intervention design.", "Medium", "reports_seeded"),
    q("How do you balance motivation with accountability?", "This probes student experience design.", "Medium", "ai_grounded"),
    q("Could the system misclassify a student as needing remediation?", "Fairness and threshold design challenge.", "High", "threshold_74"),
    q("What happens when a student improves after remediation?", "They want the close-the-loop answer.", "Medium", "threshold_74"),
    q("How do class records and LXP records stay conceptually separate?", "Pedagogically and administratively important.", "High", "not_official_records"),
    q("How does the system help teachers who are not subject specialists in remediation?", "A paper claim directly invites this question.", "High", "overview_honest"),
    q("How do you justify using an anthropomorphic mentor for minors?", "This blends pedagogy and ethics.", "High", "ai_grounded"),
    q("What learning outcome should improve first if Nexora works as intended?", "They want realistic pedagogical expectations.", "Medium", "methodology_limited"),
    q("If the panel says this is targeted tutoring rather than a full LXP, how do you respond?", "This is a nuanced pedagogy defense moment.", "Critical", "lxp_defense"),
]

AI_QUESTIONS = [
    q("What exactly does JAKIPIR do?", "This is the baseline AI question.", "High", "ai_grounded"),
    q("Is JAKIPIR retrieval-augmented or just a plain chatbot?", "They want the real architecture.", "Critical", "ai_grounded"),
    q("Where does the AI get its evidence?", "Grounding claims must be explicit.", "Critical", "ai_grounded"),
    q("How do you prevent hallucinations?", "This is inevitable in an AI defense.", "Critical", "ai_grounded"),
    q("Can JAKIPIR give wrong answers?", "A hostile-but-fair panelist will ask this bluntly.", "Critical", "ai_grounded"),
    q("What happens when the AI is unsure?", "Good AI safety design should admit uncertainty.", "High", "ai_grounded"),
    q("How do you stop the AI from giving direct answer keys?", "Academic integrity question.", "High", "ai_grounded"),
    q("What models are actually used and for what tasks?", "This tests whether claimed models are real.", "High", "ai_grounded", evidence_extra=["Repo: ai-service README maps qwen2.5:3b to text tasks and gemma3:4b to document-oriented reasoning."]),
    q("Why use qwen2.5:3b for text tutoring?", "They want model selection rationale.", "Medium", "deployment_limited"),
    q("Why use gemma3:4b for document reasoning or extraction?", "They want task-model alignment.", "Medium", "extraction_review"),
    q("How do you store embeddings?", "This checks whether RAG is technically real.", "High", "ai_grounded"),
    q("What is the role of pgvector in Nexora?", "Another direct RAG proof question.", "High", "ai_grounded"),
    q("How does PDF extraction work end to end?", "This is a likely AI workflow defense topic.", "High", "extraction_review"),
    q("What if a PDF is scanned or poorly formatted?", "The panel will test a failure case.", "High", "extraction_review"),
    q("Can teachers edit AI-generated outputs before applying them?", "Teacher oversight is essential.", "High", "extraction_review"),
    q("Is AI-generated remedial content automatically trusted by the system?", "This probes automation boundaries.", "High", "extraction_review"),
    q("What AI logs are kept?", "Logging and accountability matter.", "High", "security_controls"),
    q("Can teachers control how strict the AI grounding is?", "This is an advanced but fair question.", "High", "ai_grounded"),
    q("What is the difference between JA, JAKIPIR, AI Tutor, and LXP in your system?", "Naming consistency can confuse a panel.", "High", "ai_grounded"),
    q("Is the AI feature safe for minors?", "This is both an ethics and AI question.", "Critical", "security_controls"),
    q("What happens if Ollama is slow or unavailable?", "This is a live demo risk.", "High", "deployment_limited"),
    q("Do you rely only on local AI or can the architecture fall back to cloud APIs?", "Architecture truthfulness matters.", "High", "deployment_limited"),
    q("How do you justify calling the AI adaptive?", "Adaptivity is an overclaim risk.", "High", "ai_grounded"),
    q("How do you measure AI quality?", "A research-quality AI question.", "High", "methodology_limited"),
    q("If the panel says your AI is useful but not yet trustworthy enough for strong claims, how do you answer?", "This is a likely nuanced criticism.", "Critical", "ai_grounded"),
]

PRIVACY_QUESTIONS = [
    q("How do you protect student data?", "A privacy panelist will ask this early.", "Critical", "security_controls"),
    q("How do JWT and refresh flows work in your system?", "They want concrete auth knowledge.", "High", "security_controls"),
    q("How are OTP codes stored?", "This is a precise security detail question.", "High", "security_controls"),
    q("How do you prevent account enumeration in recovery flows?", "A sharper panelist may test security maturity.", "High", "security_controls"),
    q("How is role-based access control enforced?", "RBAC is central to school systems.", "High", "security_controls"),
    q("How do you separate teacher, admin, and student data access?", "This tests authorization design.", "High", "security_controls"),
    q("How do you handle AI chats that may contain personal information?", "Sensitive data governance question.", "Critical", "security_controls"),
    q("Are AI interaction logs stored separately from official academic records?", "Important ethical and technical boundary.", "High", "not_official_records"),
    q("How would you answer a Data Privacy Act of 2012 question?", "In the Philippines context, this is very likely.", "Critical", "security_controls"),
    q("What consent assumptions exist when minors use AI features?", "The paper should not dodge this.", "Critical", "security_controls"),
    q("What if a student types private family or health information into JAKIPIR?", "Real-world safety scenario.", "Critical", "security_controls"),
    q("How do you audit sensitive system actions?", "Auditability is an explicit repo strength.", "High", "security_controls"),
    q("How do you secure cookies, tokens, and API access?", "This is a standard security panel question.", "High", "security_controls"),
    q("What happens if someone tampers with client-side requests?", "They want server-side trust boundaries.", "Medium", "security_controls"),
    q("How is rate limiting handled?", "This checks anti-abuse basics.", "Medium", "security_controls"),
    q("How are password resets protected from abuse?", "Recovery flows are a common attack point.", "High", "security_controls"),
    q("Can AI-generated content introduce unsafe or biased advice?", "This is an ethics challenge.", "High", "ai_grounded"),
    q("What is your fallback if an AI output is harmful or misleading?", "Incident response at a capstone level.", "High", "ai_grounded"),
    q("Do you claim legal compliance is already complete?", "This is a trap around overclaiming compliance.", "Critical", "security_controls"),
    q("What security claim should you never say during defense?", "A panelist may ask bluntly about overclaiming.", "High", "security_controls"),
]

DATABASE_QUESTIONS = [
    q("Why did you choose PostgreSQL for Nexora?", "They want foundational architecture reasoning.", "Medium", "overview_honest"),
    q("Why add pgvector instead of a separate vector store?", "This tests practical design tradeoffs.", "Medium", "ai_grounded"),
    q("What does Drizzle ORM give you over raw SQL everywhere?", "ORM choice needs rationale.", "Medium", "overview_honest"),
    q("How many major table families does the system have?", "They want schema awareness.", "Medium", "overview_honest", evidence_extra=["Repo inventory identified auth, classes, lessons, assessments, performance, intervention, reporting, and AI-related table groups."]),
    q("Which tables support intervention logic?", "This ties DB design to the thesis claim.", "High", "threshold_74", evidence_extra=["Repo: intervention_cases, intervention_assignments, lxp_progress, performance_snapshots, performance_logs."]),
    q("Which tables support AI logging and extraction?", "This checks AI persistence design.", "High", "security_controls"),
    q("How is threshold information persisted?", "They want data-level rule traceability.", "High", "threshold_74"),
    q("How do performance snapshots differ from performance logs?", "A good schema-comprehension question.", "Medium", "threshold_74"),
    q("How is LXP progress recorded?", "Core data model question.", "Medium", "lxp_defense"),
    q("How are system evaluations stored?", "This links methodology claims to schema.", "High", "methodology_limited"),
    q("How are audit records stored and queried?", "Auditability question.", "Medium", "security_controls"),
    q("How do you store extracted modules before teacher approval?", "Important extraction-state question.", "Medium", "extraction_review"),
    q("How do you store embeddings?", "Another direct RAG database question.", "High", "ai_grounded"),
    q("What happens to official class records when intervention occurs?", "Data boundary question.", "High", "not_official_records"),
    q("How do you keep grade-level scope constrained in the database?", "Schema-level scope control question.", "Medium", "scope_narrow"),
    q("How do you handle indexing and query performance for retrieval?", "This is a fair vector-query performance question.", "Medium", "ai_grounded"),
    q("What data would you seed before a defense demo?", "Demo planning often comes down to data quality.", "High", "reports_seeded"),
    q("What table would be the first place to inspect if threshold gating looked wrong?", "A practical debugging question.", "Medium", "threshold_74"),
    q("What table would prove that AI interactions are really logged?", "They want concrete evidence paths.", "Medium", "security_controls"),
    q("If the panel asks whether your schema is already stable for production, what do you say?", "Schema maturity is another overclaim trap.", "High", "deployment_limited"),
]

MOBILE_WEB_QUESTIONS = [
    q("What web roles are strongest today?", "They are checking practical role parity.", "Medium", "overview_honest"),
    q("What mobile roles are strongest today?", "Parity challenge.", "High", "mobile_partial"),
    q("Can students authenticate on mobile with OTP and password recovery flows?", "Paper-to-code parity question.", "High", "mobile_partial"),
    q("Can teachers work on mobile?", "This requires a nuanced current-state answer.", "High", "teacher_mobile_partial"),
    q("Can administrators work fully on mobile?", "A likely parity trap.", "High", "admin_mobile_placeholder"),
    q("Why is the web app still the primary admin surface?", "This tests channel prioritization.", "Medium", "admin_mobile_placeholder"),
    q("How do web and mobile differ in scope today?", "Scope clarity question.", "High", "mobile_partial"),
    q("Can students access JA/LXP on mobile?", "A specific capability question.", "Medium", "mobile_partial"),
    q("How are mobile routes organized?", "This is a technical cross-platform question.", "Medium", "mobile_partial"),
    q("What happens if a teacher signs into the mobile app?", "This checks current role routing.", "Medium", "teacher_mobile_partial"),
    q("What happens if an admin signs into the mobile app?", "Current admin limitation question.", "Medium", "admin_mobile_placeholder"),
    q("Does the mobile app depend on the same backend as the web app?", "Integration basics.", "Medium", "overview_honest"),
    q("How do you justify claiming both web and mobile accessibility?", "This attacks breadth claims.", "High", "mobile_partial"),
    q("What mobile feature should not be overpromised during defense?", "Strategic communication question.", "High", "admin_mobile_placeholder"),
    q("If a panelist asks to switch roles live on mobile, what should you do?", "Demo survival question.", "High", "mobile_partial"),
    q("How do mobile assessment flows compare with web assessment flows?", "Cross-platform behavior question.", "Medium", "mobile_partial"),
    q("How do mobile auth screens reflect the backend contract?", "Contract traceability question.", "Medium", "security_controls"),
    q("How should you describe teacher mobile if not every action has been live-verified today?", "This tests safe phrasing discipline.", "High", "teacher_mobile_partial"),
    q("How should you describe admin mobile if asked directly?", "Another safe-phrasing test.", "High", "admin_mobile_placeholder"),
    q("If the web app works but mobile is unstable, what is the honest defense answer?", "A realistic live-demo contingency question.", "High", "mobile_partial"),
]

DEMO_ATTACK_QUESTIONS = [
    q("Show me the login flow right now.", "A demo-focused panelist starts with basic credibility.", "High", "backend_runtime_risk"),
    q("Show me the role-based dashboards.", "They want to know if roles are real.", "High", "overview_honest"),
    q("Show me where the 74% threshold is visible.", "They want to verify the core intervention rule.", "High", "threshold_74"),
    q("Show me a student who triggered intervention.", "This tests whether seeded data is ready.", "Critical", "reports_seeded"),
    q("Show me JAKIPIR answering from class material.", "This is a likely demo request.", "High", "ai_grounded"),
    q("Show me a PDF extraction and what happens after it finishes.", "This is an attractive but risky AI demo request.", "High", "extraction_review"),
    q("Show me that teachers can review AI output before applying it.", "They are checking governance, not just AI flash.", "High", "extraction_review"),
    q("Show me a report export.", "Export paths often break under pressure.", "Medium", "reports_seeded"),
    q("Show me the audit trail.", "A good panelist likes accountability features.", "Medium", "security_controls"),
    q("Show me what happens when a student scores exactly 74%.", "This is a rule edge case demonstration.", "High", "threshold_74"),
    q("Show me the mobile app for an administrator.", "This is a dangerous parity request.", "Critical", "admin_mobile_placeholder"),
    q("Show me the mobile app for a teacher.", "This can work only if the demo device is prepared.", "High", "teacher_mobile_partial"),
    q("Show me system health or diagnostics.", "They want proof that architecture claims are real.", "Medium", "deployment_limited"),
    q("Show me what happens if the AI is down.", "This probes resilience and honesty.", "High", "deployment_limited"),
    q("Show me how the system avoids AI-based grade tampering.", "This is a conceptual demo question.", "High", "not_official_records"),
    q("Use a random account instead of your prepared seed data.", "This is a classic demo ambush.", "Critical", "reports_seeded"),
    q("Can you switch quickly from student to teacher to admin live?", "Role switching can expose session or data issues.", "High", "backend_runtime_risk"),
    q("What if the backend suddenly fails during the demo?", "They want to see whether the team can recover calmly.", "Critical", "backend_runtime_risk"),
    q("What if Ollama is still loading the model?", "AI latency is a real demo risk.", "High", "deployment_limited"),
    q("If something breaks live, what can you still prove from the repo?", "A fair but sharp defense question.", "High", "overview_honest"),
]

DEPLOYMENT_QUESTIONS = [
    q("Can a public high school realistically run this system?", "Feasibility matters.", "High", "deployment_limited"),
    q("What hardware is required for the AI-enabled version?", "This is a practical cost question.", "High", "deployment_limited"),
    q("Can the core LMS still work without the AI stack?", "This tests graceful degradation.", "High", "deployment_limited"),
    q("How expensive is long-term AI hosting?", "A skeptical business panelist will ask this.", "Medium", "deployment_limited"),
    q("Why use local Ollama instead of a pure cloud API?", "This probes cost, privacy, and speed tradeoffs.", "Medium", "deployment_limited"),
    q("Who maintains the stack after the capstone?", "Operational sustainability question.", "High", "deployment_limited"),
    q("How much technical skill would school staff need?", "Training burden question.", "Medium", "deployment_limited"),
    q("What is your fallback when internet or local networking fails?", "Resilience question.", "High", "deployment_limited"),
    q("Can the school adopt Nexora without GPU hardware?", "Another practical feasibility question.", "High", "deployment_limited"),
    q("Is your observability stack necessary for school deployment or mainly for technical operations?", "This tests whether the team understands operator priorities.", "Medium", "deployment_limited"),
    q("Why is backend readiness more important than frontend appearance during deployment?", "A systems-thinking question.", "Medium", "backend_runtime_risk"),
    q("What would you deploy first in a pilot?", "They want phased adoption reasoning.", "Medium", "deployment_limited"),
    q("What is the safest school rollout strategy?", "Change management question.", "Medium", "deployment_limited"),
    q("What part of the architecture is most likely to raise cost or maintenance burden?", "They want honest prioritization.", "High", "deployment_limited"),
    q("What is your strongest honest deployment claim today?", "This is a strategic closeout question.", "High", "deployment_limited"),
]

TRAP_QUESTIONS = [
    q("Is the AI fully accurate?", "A careless yes would seriously damage credibility.", "Critical", "ai_grounded"),
    q("Is Nexora ready for production deployment today?", "Overclaim trap.", "Critical", "deployment_limited"),
    q("Does the AI replace the teacher during remediation?", "Ethics and pedagogy trap.", "Critical", "not_official_records"),
    q("Can your system support all grade levels and all subjects right now?", "Scope overclaim trap.", "Critical", "scope_narrow"),
    q("Is the LXP already proven effective at improving grades?", "Methodology trap.", "Critical", "methodology_limited"),
    q("Can you guarantee there will never be hallucinations?", "AI overclaim trap.", "Critical", "ai_grounded"),
    q("Does mobile fully match the web system?", "Parity overclaim trap.", "Critical", "mobile_partial"),
    q("Does the system already comply completely with all privacy laws?", "Compliance overclaim trap.", "Critical", "security_controls"),
    q("Can the AI directly change a student's academic record?", "Domain-boundary trap.", "Critical", "not_official_records"),
    q("If a student gets 74%, are they automatically failing?", "Terminology trap around threshold meaning.", "High", "threshold_74"),
    q("Is the 74% threshold universally correct?", "Policy absolutism trap.", "High", "threshold_74"),
    q("Is every chart in your analytics based on large real-world school data already?", "Data-validity trap.", "High", "reports_seeded"),
    q("If the backend is currently down, does that mean the system is fake?", "A hostile reframing trap.", "High", "backend_runtime_risk"),
    q("Can you promise the AI is safe for all minors under every circumstance?", "Absolute safety trap.", "Critical", "security_controls"),
    q("Should schools trust the AI without teacher review?", "Oversimplification trap.", "Critical", "ai_grounded"),
]


QUESTION_BANK = OrderedDict(
    [
        ("General Defense Questions", GENERAL_QUESTIONS),
        ("Technical Questions", TECHNICAL_QUESTIONS),
        ("Research and Methodology Questions", RESEARCH_QUESTIONS),
        ("LMS, LXP, and Pedagogy Questions", PEDAGOGY_QUESTIONS),
        ("AI, JAKIPIR, and RAG Questions", AI_QUESTIONS),
        ("Privacy and Security Questions", PRIVACY_QUESTIONS),
        ("Database Questions", DATABASE_QUESTIONS),
        ("Mobile and Web Implementation Questions", MOBILE_WEB_QUESTIONS),
        ("Demo Attack Questions", DEMO_ATTACK_QUESTIONS),
        ("Deployment and Feasibility Questions", DEPLOYMENT_QUESTIONS),
        ("Trap Questions", TRAP_QUESTIONS),
    ]
)


PANELIST_PERSONAS = [
    {
        "name": "Technical Panelist",
        "focus": "architecture, database, APIs, security, queues, observability, AI separation, deployment realism",
        "hidden_concern": "The repo sounds mature, so any incorrect technical claim will be punished fast.",
        "danger_level": "Critical",
        "likely_questions": [
            "Where is the 74% threshold actually enforced?",
            "Show me where pgvector is really used.",
            "How does the backend talk to the AI service?",
            "What happens if the AI service or backend is down?",
            "How do you prevent AI outputs from mutating official grades?",
        ],
        "best_answer": "Keep answers anchored to concrete modules, schemas, and routes. Admit runtime status honestly if a service is currently down, then pivot to the verified code boundary and the prepared fallback evidence.",
        "answer_to_avoid": "Do not answer with vague architecture buzzwords like 'full-stack integration' without naming NestJS, Next.js, FastAPI, PostgreSQL, Redis, and Ollama clearly.",
        "evidence": [
            "backend/src/modules/lxp/lxp.service.ts",
            "backend/src/app.module.ts",
            "backend/src/main.ts",
            "ai-service/app/retrieval_service.py",
            "docker-compose.yml",
        ],
        "show_if_asked": "Show the teacher performance/intervention logic, health controller, AI policy schema, and retrieval code.",
    },
    {
        "name": "Research and Academic Panelist",
        "focus": "problem statement, objectives, gap, methodology, scope, evaluation, validity, literature",
        "hidden_concern": "The implementation may be stronger than the research design, so overclaiming impact will be exposed.",
        "danger_level": "Critical",
        "likely_questions": [
            "What exact gap does Nexora solve that existing LMS tools do not?",
            "How is the school context specifically justified?",
            "How many respondents evaluated the system?",
            "What statistical treatment did you use?",
            "Is Chapter 4 really results and discussion or mostly design documentation?",
        ],
        "best_answer": "Separate implementation proof from outcome proof. Say the capstone demonstrates feasibility, system integration, and prototype behavior first, while broader educational effectiveness needs larger-scale study.",
        "answer_to_avoid": "Do not claim proven effectiveness unless you can defend respondents, instrument, and statistics immediately.",
        "evidence": [
            "Concept paper objectives and scope",
            "Paper extract lines on promised evaluation dimensions",
            "Repo system_evaluations schema and routes",
        ],
        "show_if_asked": "Show the evaluation module and then admit that implementation readiness is stronger than formal impact evaluation evidence.",
    },
    {
        "name": "Education and Pedagogy Panelist",
        "focus": "LMS vs LXP, mastery learning, intervention fairness, teacher workload, learner support design",
        "hidden_concern": "They suspect the system is just an LMS with a gated remedial tab.",
        "danger_level": "High",
        "likely_questions": [
            "Why do you call it an LXP?",
            "How is 74% justified pedagogically?",
            "Could high-performing students also benefit from the same support?",
            "How does the system personalize learning rather than just restrict access?",
            "Does the AI undermine teacher judgment?",
        ],
        "best_answer": "Defend Nexora as an LMS with intervention-oriented LXP features. Emphasize targeted guidance, review paths, teacher control, and AI as support rather than replacement.",
        "answer_to_avoid": "Do not insist that the system is a full commercial-grade LXP if the panel challenges personalization depth.",
        "evidence": [
            "Paper scope statements about targeted intervention",
            "intervention_cases and intervention_assignments schema",
            "student JA/LXP routes",
        ],
        "show_if_asked": "Show the intervention queue, LXP progress concepts, and JA hub guidance around grounded help.",
    },
    {
        "name": "AI Ethics and Data Privacy Panelist",
        "focus": "hallucination, minors, consent, harmful outputs, logging, privacy law, abuse prevention",
        "hidden_concern": "AI for minors is risky unless governance is explicit.",
        "danger_level": "Critical",
        "likely_questions": [
            "Can JAKIPIR hallucinate?",
            "What happens if a student enters private information into the AI mentor?",
            "Are AI chats logged and who can see them?",
            "How do you answer a Data Privacy Act question?",
            "Why should minors trust an anthropomorphic AI mentor?",
        ],
        "best_answer": "Acknowledge risk, then show the mitigations: grounding, AI policy controls, chat logs, role boundaries, and the separation of AI outputs from official records.",
        "answer_to_avoid": "Do not say the AI is perfectly safe or completely compliant by default.",
        "evidence": [
            "ai_interaction_logs schema",
            "class_ai_policies schema and teacher controls",
            "OTP, JWT, audit, and validation setup",
        ],
        "show_if_asked": "Show AI policy toggles and explain that official grades remain outside AI control.",
    },
    {
        "name": "Demo-Focused Panelist",
        "focus": "live behavior, broken routes, empty datasets, random account switching, latency, edge cases",
        "hidden_concern": "They assume documentation can be polished while live systems still break.",
        "danger_level": "Critical",
        "likely_questions": [
            "Show me the threshold rule live.",
            "Show me a student who actually triggered intervention.",
            "Use a random account.",
            "Show me mobile admin.",
            "What if the AI takes too long?",
        ],
        "best_answer": "Control the demo sequence tightly, start with stable web routes, use seeded accounts intentionally, and keep a fallback screenshot/video pack ready for AI-heavy flows.",
        "answer_to_avoid": "Do not improvise with unprepared accounts or high-latency AI flows first.",
        "evidence": [
            "Current runtime check showing backend unavailable but frontend and AI service reachable",
            "Seeded report and intervention routes in the repo",
            "Demo plan prepared in this package",
        ],
        "show_if_asked": "Show login, role dashboards, teacher classes, student assessment results, intervention trigger, then only stable AI surfaces.",
    },
    {
        "name": "Skeptical Business and Deployment Panelist",
        "focus": "cost, sustainability, school feasibility, training burden, infrastructure, maintenance",
        "hidden_concern": "A public school may not realistically maintain a multi-service AI stack.",
        "danger_level": "High",
        "likely_questions": [
            "Can the school actually run this hardware?",
            "Who maintains Ollama, Docker, and the database?",
            "What if there is no GPU?",
            "Is this still useful without the AI?",
            "Why not just use existing school tools and a simpler workflow?",
        ],
        "best_answer": "Position the current system as technically deployable for pilot use, with the core LMS still useful even if AI is degraded, and describe AI-enabled deployment as a staged decision rather than a mandatory baseline.",
        "answer_to_avoid": "Do not promise easy school-wide rollout with no training, no cost, and no operational support.",
        "evidence": [
            "docker-compose.yml and README deployment notes",
            "ai-service readiness output with available models",
            "hardware notes in chapter-3 documentation",
        ],
        "show_if_asked": "Show the deployment stack diagram and then explain the staged rollout path."
    },
]


SYSTEM_WEAKPOINTS = [
    Weakpoint("Backend runtime not currently healthy", "The current audit could reach the frontend and ai-service, but backend port 3000 was unreachable.", "A live defense can collapse immediately if login or data-backed routes fail.", "High", "Critical", "Critical", "State that the implementation exists but demo readiness depends on verified backend startup on the defense machine.", "Restore backend health, confirm /api/health/live and /api/health/ready, and test seeded logins.", "Mention proactively during internal prep, not during the defense unless a runtime issue appears."),
    Weakpoint("74% threshold justification is thinner than the implementation", "The code consistently uses 74, but the policy defense is weaker than the technical enforcement.", "Panelists can attack the rule as arbitrary.", "High", "High", "High", "Call it the current configurable project cutoff and not a universal law.", "Prepare a one-minute justification tied to mastery learning and school remediation policy discussion.", "Mention only if asked; do not spotlight weakness first."),
    Weakpoint("LXP label is only partially defensible", "Nexora is strongest as an LMS with intervention-oriented LXP features, not a full enterprise LXP.", "A pedagogy panelist may say the system is just gated remediation.", "High", "High", "High", "Narrow the claim to intervention-focused LXP features.", "Rewrite defense script language to avoid overclaiming enterprise-level personalization.", "Mention only if asked."),
    Weakpoint("System evaluation evidence is weaker than system implementation", "The repo has evaluation features, but strong measured outcome evidence is not obvious from the current runtime context.", "Research panelists may say Chapter 4 proves design, not effectiveness.", "High", "Critical", "Critical", "Differentiate feasibility validation from educational outcome validation.", "Prepare respondent counts, instruments, and statistics if they exist; otherwise narrow claims.", "Mention carefully when discussing limitations."),
    Weakpoint("Analytics quality depends heavily on seeded or live data", "Reports and charts are only persuasive if data is populated well.", "Empty or thin dashboards look incomplete even when the code is real.", "High", "High", "High", "Say the workflow is implemented and currently demonstrated with controlled data.", "Seed realistic classes, attempts, intervention cases, reports, and evaluations.", "Do not mention unless a sparse page is visible."),
    Weakpoint("Admin mobile remains limited", "Admin mobile routes resolve to generic workspace sections rather than full operations.", "A role-parity claim can be disproven quickly.", "Medium", "High", "High", "State that administration remains strongest on web.", "Do not demo admin mobile; keep admin workflow on web.", "Mention only if asked."),
    Weakpoint("Teacher mobile exists but should not be sold as fully verified parity", "Teacher tabs and detail screens exist in code, but full defense-device verification was not done in this run.", "An unexpected mobile flow bug could damage confidence.", "Medium", "High", "High", "Describe teacher mobile as prototype-scope support.", "Test teacher mobile routes end to end on the defense device.", "Mention only if asked."),
    Weakpoint("AI mentor safety depends on disciplined framing", "Grounding and policy controls exist, but AI remains probabilistic.", "Absolute claims invite credibility collapse.", "High", "Critical", "Critical", "Always say assistive, grounded, and review-aware.", "Prepare a short, honest AI limitation statement and show policy controls.", "Mention proactively in AI explanation."),
    Weakpoint("Extraction quality is variable for messy or scanned PDFs", "The pipeline is real, but not magical.", "Live failure on a bad file can embarrass the team.", "High", "High", "High", "Say extraction accelerates authoring and remains teacher-reviewed.", "Use one known-good PDF for the demo and prepare screenshots of prior successful runs.", "Mention only if asked."),
    Weakpoint("Local AI infrastructure increases demo fragility", "Ollama, models, and the AI service add startup and latency risk.", "A slow model can kill momentum.", "High", "High", "High", "Set expectations that AI-heavy tasks can be slower than ordinary LMS actions.", "Warm the models ahead of time and prepare fallback artifacts.", "Mention only if latency appears."),
    Weakpoint("Paper still uses broad 'all subjects and grade levels' language", "The repo structurally supports Grades 7 to 10, while the paper still contains broader wording.", "This creates scope mismatch under questioning.", "High", "High", "High", "Reframe those lines as intended school scope, not fully validated scope.", "Revise the paper wording before defense.", "Mention only if challenged."),
    Weakpoint("The title is grammatically incomplete", "The paper title still omits the word 'System'.", "A panelist can catch it on page one and lower confidence immediately.", "High", "Medium", "Medium", "Acknowledge and correct it quickly.", "Fix the title before printing or submission.", "Mention proactively only in revision, not in defense."),
    Weakpoint("Methodology chapter may be read as design-heavy", "The paper has many process flows and technical diagrams.", "A research panelist may say 'results and discussion' are under-evidenced.", "High", "Critical", "Critical", "Stress prototype validation and keep outcome claims narrow.", "Insert or strengthen actual measured evaluation content.", "Mention only in limitations or if challenged."),
    Weakpoint("Anthropomorphic NPC language can sound overbranded", "The paper leans into NPC framing.", "A skeptical panelist may treat it as gimmick language.", "Medium", "Medium", "Medium", "Translate NPC into 'guided AI mentor persona' when answering live.", "Tone down branding-heavy wording in the script.", "Mention only if asked."),
    Weakpoint("No offline mode", "The paper excludes offline functionality and the stack expects online service coordination.", "School feasibility questions can land here.", "Medium", "High", "High", "Say offline mode is outside current scope to preserve real-time data integrity.", "Prepare a staged rollout answer that assumes reliable connectivity.", "Mention only if asked."),
    Weakpoint("Operational maintenance burden is non-trivial", "The stack includes backend, frontend, postgres, redis, ollama, and ai-service.", "A business panelist may say this is too heavy for a public school.", "High", "High", "High", "Defend pilot deployment first, not turnkey mass deployment.", "Prepare a smaller deployment pathway or AI-degraded mode explanation.", "Mention only if asked."),
    Weakpoint("AI logs may contain sensitive student prompts", "The system logs AI interactions for auditability.", "Privacy questions become sharper because minors are involved.", "High", "Critical", "Critical", "Say logs exist for accountability and must be governed under school policy.", "Prepare a privacy/governance answer and consider retention policy wording.", "Mention only if asked."),
    Weakpoint("Live demo depends on seeded accounts and data hygiene", "Without realistic records, strong features look weak.", "The panel may request random flows.", "High", "High", "High", "Use controlled seed data and say the demo uses prepared academic records for consistency.", "Validate accounts, classes, attempts, interventions, and reports before defense.", "Mention only if the panel asks about demo data."),
    Weakpoint("Random-account demo attacks are dangerous", "Role routes and data completeness vary by account.", "Improvised switching can expose empty or invalid surfaces.", "High", "High", "High", "Politely steer to prepared accounts while explaining each role consistently.", "Prepare at least one stable account per role and one fallback student.", "Do not mention proactively."),
    Weakpoint("The paper promises evaluation dimensions like reliability and portability, but proof may be thin", "The wording sounds stronger than the visible dataset.", "A research panelist can ask for hard numbers.", "High", "High", "High", "Say those dimensions are part of the evaluation framework, but claims are bounded to prototype validation unless the data is ready.", "Collect and print the evaluation summary if available.", "Mention only if asked."),
    Weakpoint("The system is broad enough that one broken flow can overshadow many working ones", "There are many modules, routes, and roles.", "Defense impressions are often shaped by the first failure.", "High", "High", "High", "Control the order and show the most stable flows first.", "Use a disciplined demo plan with a do-not-demo list.", "Mention only in team prep."),
    Weakpoint("AI feature naming can confuse the panel", "JA, JAKIPIR, AI tutor, AI mentor, and LXP can blur together verbally.", "Confusion can make the project sound less coherent than it is.", "Medium", "Medium", "Medium", "Standardize one naming ladder in the defense script.", "Use 'JA/JAKIPIR mentor inside Nexora' consistently and define it once.", "Mention proactively in the opening explanation."),
    Weakpoint("Teacher override and governance language must be precise", "Intervention cases and approval semantics can be misunderstood.", "Governance overclaim is easy to catch.", "Medium", "High", "High", "Say teacher management remains central even when the system flags risk automatically.", "Confirm exact pending/active workflow and phrase it safely.", "Mention only if asked."),
    Weakpoint("Observability stack may look like overengineering if explained badly", "Grafana, Prometheus, Loki, and Tempo are real, but not the main educational contribution.", "Panelists may think the project is too infrastructure-heavy for a school capstone.", "Medium", "Medium", "Medium", "Position observability as operator support, not the headline feature.", "Keep observability out of the main demo unless a technical panelist asks.", "Do not mention proactively."),
    Weakpoint("Frontend reachable while backend is down can confuse the panel", "A visible login page may create false confidence until API actions fail.", "This can create a more embarrassing failure than a visibly down system.", "High", "High", "High", "Verify backend before opening the frontend in defense.", "Run health checks first and keep screenshots ready if the backend is unstable.", "Mention only in team prep."),
]


PAPER_WEAKPOINTS = [
    Weakpoint("Title omits the word 'System'", "The title reads 'A Learning Management With Learning Experience Platform Features' instead of 'A Learning Management System...'.", "It is a first-page credibility hit.", "High", "Medium", "Medium", "Acknowledge it as a wording correction, not a conceptual flaw.", "Fix the title everywhere before final printing.", "Mention proactively in revision only."),
    Weakpoint("Broad scope wording remains in places", "The paper still includes wording about all subjects and grade levels.", "It can be contrasted against the repo's tighter grade-level enforcement.", "High", "High", "High", "Clarify that broad wording reflects intended institutional scope, while implemented scope is narrower and prototype-bounded.", "Replace broad statements with Grades 7 to 10 and carefully qualified subject coverage.", "Mention only if asked."),
    Weakpoint("Need statement can sound dramatic", "Some narrative passages use high-intensity language about collapse, severe failure, and impossible teacher workload.", "Panels may see this as rhetorical overreach.", "Medium", "Medium", "Medium", "Bring the answer back to practical school workflow pain points.", "Tone down dramatic wording in the paper.", "Mention only if challenged."),
    Weakpoint("LXP definition can still be challenged", "The paper may sound like it equates targeted remediation with a full LXP identity.", "This invites a conceptual attack.", "High", "High", "High", "Defend it as LXP features inside LMS scope.", "Revise lines that sound like full enterprise LXP equivalence.", "Mention only if asked."),
    Weakpoint("74% justification is not yet airtight", "The rule is consistent, but the literature-to-policy bridge remains vulnerable.", "Panelists can attack it as arbitrary.", "High", "High", "High", "Say it is the project's current configurable cutoff.", "Add a clearer justification paragraph or policy note.", "Mention only if asked."),
    Weakpoint("Methodology promises may exceed visible evidence", "The text promises evaluation of usability, functionality, reliability, and portability.", "Without ready tables and statistics, this becomes a research weakness.", "High", "Critical", "Critical", "Separate prototype validation from broader effectiveness claims.", "Strengthen evaluation data presentation.", "Mention only if challenged."),
    Weakpoint("Chapter 4 may read as design-heavy rather than result-heavy", "Many process figures dominate the discussion.", "A research panelist can say the chapter is mostly system documentation.", "High", "Critical", "Critical", "Frame it as prototype validation and implementation discussion.", "Insert clearer evaluation results, observations, or measured findings.", "Mention only if asked."),
    Weakpoint("Related literature uses strong statistics that may invite citation scrutiny", "Claims like 74% greater engagement and similar figures can draw attention.", "If a panelist asks for source precision, weak citation recall will hurt.", "Medium", "High", "High", "Use only the claims you can defend confidently from the paper.", "Recheck every high-impact statistic and citation pair.", "Mention only if asked."),
    Weakpoint("Problem gap versus existing LMS tools may feel underproven", "The paper can be challenged on whether Google Classroom or Moodle plus process changes could solve enough of the problem.", "Novelty pressure increases.", "High", "High", "High", "Focus on integrated intervention workflow, not generic content hosting.", "Sharpen the comparative gap language.", "Mention only if challenged."),
    Weakpoint("The paper mixes institutional deployment language with prototype language", "Some sections read like full rollout, others like capstone scope.", "This can sound internally inconsistent.", "Medium", "High", "High", "Use the phrase 'prototype scope' consistently during defense.", "Standardize deployment wording throughout the paper.", "Mention only if asked."),
    Weakpoint("AI NPC framing can sound more like a design motif than a validated educational construct", "The anthropomorphic mentor idea is interesting but vulnerable.", "A pedagogy or ethics panelist may push back.", "Medium", "Medium", "Medium", "Translate it into supportive mentor persona language.", "Reduce branding-heavy phrasing where needed.", "Mention only if asked."),
    Weakpoint("Privacy and minors are not foregrounded strongly enough in the narrative", "The system clearly involves student data and AI interaction.", "A privacy panelist may say governance is underwritten.", "High", "Critical", "Critical", "Explain the controls in defense even if the paper section is lighter.", "Add stronger privacy, consent, and governance language.", "Mention only if asked."),
    Weakpoint("The paper could overstate educational effectiveness if spoken carelessly", "Implementation is real, but long-term outcome proof is not the same thing.", "This is a classic capstone trap.", "High", "Critical", "Critical", "Say the study demonstrates system feasibility and workflow integration first.", "Tighten conclusion and results wording.", "Mention proactively in closing limitations."),
    Weakpoint("Teacher workload claims need cautious phrasing", "The paper argues Nexora reduces workload.", "Panelists may ask for measured proof.", "High", "High", "High", "Say it is designed to reduce repetitive manual work, not already statistically proven to do so.", "Adjust wording from proven effect to intended support.", "Mention only if asked."),
    Weakpoint("All-subject rhetoric is stronger than the repo evidence", "Even if the platform is structurally extensible, real subject coverage proof may be uneven.", "This opens a direct mismatch line.", "High", "High", "High", "Treat it as intended scope and design extensibility.", "Narrow or qualify these statements.", "Mention only if asked."),
    Weakpoint("Architecture maturity can be mistaken for deployment maturity", "The paper cites modern infrastructure and observability.", "That can tempt overclaiming readiness.", "Medium", "High", "High", "Call it production-minded architecture for prototype scope.", "Add a more explicit deployment-limitation statement.", "Mention only if asked."),
    Weakpoint("Evaluation significance to the school may sound assumed rather than measured", "School impact claims are appealing but may be under-measured.", "A research panelist can separate significance from evidence.", "Medium", "High", "High", "Distinguish expected benefit from measured impact.", "Clarify significance as rationale, not proven final effect.", "Mention only if challenged."),
    Weakpoint("If the panel reads the paper literally, they may expect more live parity than the team should promise", "Paper breadth can drive demo expectations.", "This creates demo risk before the laptop is even opened.", "High", "High", "High", "Verbally narrow the demonstrated scope early.", "Align paper, script, and demo plan tightly.", "Mention proactively in the system overview."),
    Weakpoint("Use-case volume can make the paper feel exhaustive but also bloated", "Many use-case tables exist.", "A panelist may ask whether all of them were truly validated.", "Medium", "Medium", "Medium", "Say the use cases map intended workflows, while defense will focus on the core validated flows.", "Trim or group if revision is still possible.", "Mention only if asked."),
    Weakpoint("If citations are not memorized, strong literature claims become a liability", "The paper uses many external claims.", "Panelists may test recall.", "Medium", "High", "High", "Memorize only the most central literature claims and avoid overusing others in oral answers.", "Prepare a citation cheat sheet.", "Mention only in preparation."),
    Weakpoint("The paper may not clearly separate system auditability from educational evaluation", "Audit logs and evaluations are different kinds of evidence.", "Confusing them weakens methodology answers.", "Medium", "Medium", "Medium", "Keep technical accountability and educational evaluation as separate tracks in oral defense.", "Refine wording where those ideas blur.", "Mention only if asked."),
    Weakpoint("Future recommendations may need stronger pruning", "Some claimed features are ambitious enough that they may be safer as future work.", "Overfull capstones draw skepticism.", "Medium", "Medium", "Medium", "Admit staged growth clearly.", "Move weaker claims out of the core scope if revision is possible.", "Mention only if asked."),
    Weakpoint("The paper's strongest confidence exceeds the current runtime health seen today", "Backend was unreachable during this audit run.", "Even if the paper is strong, demo-day technical failure can make it look dishonest.", "High", "High", "High", "Treat this as an environment issue and prepare fallback proof.", "Fix runtime before defense and do not rely on paper polish alone.", "Mention only internally."),
    Weakpoint("Concept-paper evolution is stronger than explicit explanation of that evolution", "The repo has grown beyond the early lightweight AI framing.", "A panelist might ask which version of the project story is authoritative.", "Medium", "Medium", "Medium", "Say the concept paper set the direction and the implementation matured into a richer architecture.", "Prepare a clear 'concept to implementation' summary slide.", "Mention only if asked."),
    Weakpoint("The defense can be hurt if the team answers with marketing language instead of scope language", "The paper has a few branding-heavy passages.", "This is easy for panels to punish.", "High", "High", "High", "Stay concrete, school-specific, and limitation-aware in every answer.", "Rewrite oral script and rehearse realistic language.", "Mention proactively in prep."),
]


DEMO_WEAKPOINTS = [
    Weakpoint("Backend port 3000 currently unreachable", "The biggest immediate live-demo blocker.", "Login and all data-backed routes can fail instantly.", "High", "Critical", "Critical", "Open with honesty if needed and pivot to prepared artifacts only if runtime recovery fails.", "Fix startup and re-verify health endpoints.", "Internal prep only."),
    Weakpoint("AI extraction is attractive but risky as a live first demo", "It depends on file quality and model speed.", "A slow or bad extraction wastes defense time.", "High", "High", "High", "Demo it only if specifically asked or if a known-good extraction is prepared.", "Use a proven file and keep screenshots ready.", "Internal prep only."),
    Weakpoint("Random-account requests can expose empty data", "Prepared accounts matter.", "The panel may interpret empty pages as missing implementation.", "High", "High", "High", "Use curated seeded accounts and explain why.", "Validate each seed account.", "Internal prep only."),
    Weakpoint("Admin mobile is not a safe live demo", "It is not the strongest mobile surface.", "A parity attack can start here.", "Medium", "High", "High", "Keep admin on web.", "Do not offer admin mobile unless explicitly required.", "Internal prep only."),
    Weakpoint("Teacher mobile requires device-specific confidence", "Routes exist, but unverified live parity can still fail.", "A half-working route damages trust.", "Medium", "High", "High", "Treat teacher mobile as optional or backup.", "Run a live teacher-mobile rehearsal.", "Internal prep only."),
    Weakpoint("Sparse analytics can look unimpressive", "Charts and tables need data.", "An empty dashboard feels unfinished.", "High", "High", "High", "Show stable charts only after seeding realistic values.", "Seed reports, attempts, interventions, and evaluations.", "Internal prep only."),
    Weakpoint("Role switching can break flow momentum", "Logging in and out repeatedly is risky under time pressure.", "Session, cache, or seed issues can surface.", "Medium", "Medium", "Medium", "Keep role transitions scripted.", "Pre-open tabs or use separate browsers if possible.", "Internal prep only."),
    Weakpoint("Ollama cold starts can create awkward silence", "AI model warm-up is a real timing issue.", "Panels lose patience quickly.", "High", "High", "High", "Narrate that AI tasks are asynchronous and show pre-generated evidence if needed.", "Warm models before the panel enters.", "Internal prep only."),
    Weakpoint("A visible frontend with a hidden dead backend can mislead the presenters themselves", "The UI can load while APIs fail.", "This leads to false confidence before the first click.", "High", "High", "High", "Run health checks first.", "Add a pre-demo port and health checklist.", "Internal prep only."),
    Weakpoint("Export flows may fail on missing or thin data", "Reports are often brittle in demos.", "An export failure looks bad even if the module is real.", "Medium", "Medium", "Medium", "Only demo one verified export path.", "Test it the same day as defense.", "Internal prep only."),
    Weakpoint("Intervention trigger demo needs carefully chosen seeded scores", "The 74% rule is central and must be visible.", "If no student is at-risk, the LXP story weakens.", "High", "High", "High", "Use a prepared student just below threshold.", "Seed attempts and performance snapshots deliberately.", "Internal prep only."),
    Weakpoint("Evaluation pages may show no data", "The schema exists but records may be absent.", "That weakens methodology perception.", "High", "High", "High", "Do not make evaluation pages central unless data exists.", "Seed or capture evaluation entries before defense.", "Internal prep only."),
    Weakpoint("JA or AI history screens may surface awkward prompt content", "Logs can contain unpredictable text.", "This can distract the panel.", "Medium", "Medium", "Medium", "Use sanitized demo data if showing history.", "Review logs before defense.", "Internal prep only."),
    Weakpoint("Discussion or announcement timing claims may be overread as push notifications", "The wording needs care.", "A panelist may expect native push behavior.", "Medium", "Medium", "Medium", "Say web and in-app real-time updates unless push is specifically implemented and verified.", "Tighten script wording.", "Internal prep only."),
    Weakpoint("If one AI feature fails, the team may panic and overshare", "Defense composure matters as much as software.", "Poor recovery can do more damage than the bug.", "High", "High", "High", "Fall back to the strongest LMS flow and explain the limitation calmly.", "Rehearse failure recovery lines.", "Internal prep only."),
]


AI_WEAKPOINTS = [
    Weakpoint("Hallucination risk can never be zero", "Grounding reduces but does not eliminate wrong answers.", "Absolute claims are easy to destroy.", "High", "Critical", "Critical", "Say assistive and grounded, not perfect.", "Prepare a limitation statement.", "Mention proactively in AI section."),
    Weakpoint("Adaptive-intelligence wording is stronger than measurable proof", "The system is guided and evidence-aware, but not fully adaptive in the research sense.", "Pedagogy and AI panels may push here.", "High", "High", "High", "Call it guided, contextual, and policy-controlled.", "Tone down 'adaptive' where necessary.", "Mention only if asked."),
    Weakpoint("Anthropomorphic mentor framing can trigger ethics questions", "NPC language is not neutral for minors.", "Panels may challenge trust and manipulation.", "Medium", "High", "High", "Translate the feature into mentor persona and supportive UX terms.", "Prepare ethics framing.", "Mention only if asked."),
    Weakpoint("Extraction accuracy varies with source quality", "Messy scans and formulas remain difficult.", "Live failure can undermine trust in AI claims.", "High", "High", "High", "Describe it as teacher-reviewed AI assistance.", "Use clean known-good demo files.", "Mention only if asked."),
    Weakpoint("AI logs may contain sensitive content", "Logging is useful but privacy-sensitive.", "Panels may ask who can see the content.", "High", "Critical", "Critical", "Emphasize role control and governance need.", "Review retention and access framing.", "Mention only if asked."),
    Weakpoint("Local model performance depends on hardware", "Useful AI still needs resources.", "Slow inference hurts demo quality.", "High", "High", "High", "Set expectations and stage rollout accordingly.", "Warm models and benchmark likely tasks.", "Mention only if asked."),
    Weakpoint("Cloud fallback code complicates a 'strictly local' narrative", "Architecture reality is more nuanced.", "Inconsistency can damage credibility.", "Medium", "Medium", "Medium", "Say the current preferred deployment is local-first, with optional fallback architecture present in code.", "Standardize phrasing across paper and defense.", "Mention only if asked."),
    Weakpoint("Teacher oversight is necessary but can be underexplained", "Panels want to know who remains responsible.", "If unclear, the AI seems too autonomous.", "High", "High", "High", "Show policy controls and review steps.", "Keep one screenshot or route ready.", "Mention proactively in AI explanation."),
    Weakpoint("No large-scale AI quality study is ready", "Implementation does not equal robust evaluation.", "Research panelists may push for metrics.", "High", "High", "High", "Say evaluation of AI quality remains an ongoing area.", "Avoid impact-overclaim language.", "Mention only if asked."),
    Weakpoint("RAG quality depends on the quality of indexed content", "Bad or thin source material weakens answers.", "A poor class dataset leads to weak AI performance.", "Medium", "Medium", "Medium", "Explain that the mentor is only as strong as the available visible class evidence.", "Use seeded classes with strong source content.", "Mention only if asked."),
    Weakpoint("AI can still refuse or degrade when evidence is insufficient", "This is safe behavior but can look like failure to an uninformed panel.", "Demo expectations may clash with safe behavior.", "Medium", "Medium", "Medium", "Frame refusals as a safety feature, not a bug.", "Prepare one example refusal explanation.", "Mention proactively if showing JA."),
    Weakpoint("AI policy terminology may confuse non-technical panelists", "Strict grounding and source scope are helpful but abstract.", "Confusion reduces perceived control.", "Medium", "Medium", "Medium", "Translate settings into plain language during defense.", "Simplify terminology in slides.", "Mention proactively in AI section."),
    Weakpoint("JAKIPIR naming can overshadow the actual educational function", "Branding may distract from pedagogy.", "Panels may ask why the name matters.", "Low", "Low", "Low", "Say it is simply the mentor persona within Nexora.", "Do not overexplain the brand.", "Mention only if asked."),
    Weakpoint("The AI cannot be defended as an official grader", "This would violate the system's safest boundary.", "A careless answer here is very dangerous.", "High", "Critical", "Critical", "State clearly that official grading remains teacher- and record-driven.", "Keep AI out of any official-grade claim.", "Mention proactively if grading comes up."),
    Weakpoint("AI success can be oversold because the repo is technically impressive", "The stronger the architecture, the more tempting the overclaim.", "Panels punish hype.", "High", "High", "High", "Use measured, scope-aware language.", "Rehearse calm non-marketing answers.", "Mention proactively in team prep."),
]


METHODOLOGY_WEAKPOINTS = [
    Weakpoint("Implementation proof is stronger than outcome proof", "This is the core methodology gap.", "Research panels prioritize this distinction.", "High", "Critical", "Critical", "Say the study validates prototype feasibility first.", "Clarify evaluation scope and claims.", "Mention carefully in limitations."),
    Weakpoint("Evaluation dimensions are promised more strongly than evidenced", "Usability, functionality, reliability, and portability need concrete data.", "Without numbers, these remain soft claims.", "High", "High", "High", "Speak cautiously about evaluation status.", "Prepare a printed summary if data exists.", "Mention only if asked."),
    Weakpoint("Respondent and instrument details may not be demo-ready", "The panel may ask for them immediately.", "Weak recall damages academic confidence.", "High", "High", "High", "Memorize only what is actually available and never invent numbers.", "Prepare a cheat sheet.", "Mention only if asked."),
    Weakpoint("Statistical treatment may be thin or absent", "This is a classic thesis defense target.", "A panelist can use it to downgrade the research rigor.", "High", "Critical", "Critical", "Admit prototype-study limits if necessary.", "Add or strengthen statistical treatment before defense if possible.", "Mention only if asked."),
    Weakpoint("Chapter 4 leans heavily on process flow documentation", "Good for software explanation, weaker for results defense.", "Panels may say it reads like system design, not discussion.", "High", "High", "High", "Frame it as implementation-centered capstone evidence.", "Add measured observations where possible.", "Mention only if challenged."),
    Weakpoint("School-specific problem severity claims may need better local evidence", "Broad statements about workload or intervention difficulty invite proof requests.", "If unsupported, they sound inflated.", "Medium", "Medium", "Medium", "Use modest language and connect back to practical workflow need.", "Tighten contextual evidence in the paper.", "Mention only if asked."),
    Weakpoint("Related systems table may overstate competitive gaps", "Comparative tables are often vulnerable.", "Commercial-system comparisons are easy to attack.", "Medium", "High", "High", "Say the table is a scoped academic comparison, not a market-battle claim.", "Recheck every comparison line.", "Mention only if asked."),
    Weakpoint("The 74% threshold combines literature logic and project policy, but not yet formal institutional validation", "That distinction matters methodologically.", "A panel can call it arbitrary if stated too strongly.", "High", "High", "High", "Describe it as current project configuration pending further validation.", "Add school-policy consultation if possible.", "Mention only if asked."),
    Weakpoint("Significance claims may sound stronger than measured outcomes", "Significance is rationale, not final proof.", "Panels may separate these sharply.", "Medium", "Medium", "Medium", "Use significance as intended benefit, not confirmed effect.", "Rewrite aggressive impact lines.", "Mention only if challenged."),
    Weakpoint("Broad scope creates evaluation dilution", "The more features claimed, the harder it is to evaluate them rigorously.", "A panel may say the study tries to prove too much.", "Medium", "High", "High", "Narrow the defended contribution to the intervention workflow.", "Trim oral emphasis to the core features.", "Mention proactively in the overview."),
    Weakpoint("If the defense relies too much on technical complexity, methodology can look secondary", "Capstones still need research discipline.", "The panel may say it is a good product but a weaker study.", "High", "High", "High", "Balance every technical answer with objective and scope language.", "Rehearse research-first answers for research panelists.", "Mention only in prep."),
    Weakpoint("Paper claims about teacher workload reduction need empirical caution", "Intended support is not yet identical to measured reduction.", "A panel can ask for evidence.", "High", "High", "High", "Say the system is designed to reduce repetitive manual work.", "Remove 'proven reduction' phrasing.", "Mention only if asked."),
    Weakpoint("No real school-wide deployment yet", "That limits ecological validity.", "A panelist may ask whether this is still hypothetical.", "High", "Medium", "Medium", "Say the capstone is pre-deployment/pilot-oriented.", "Keep deployment claims narrow.", "Mention proactively in limitations."),
    Weakpoint("Feature richness can make the objective mapping sound fuzzy", "Too many modules can blur the thesis story.", "Panels may ask what the real study is actually about.", "Medium", "High", "High", "Anchor repeatedly on targeted intervention.", "Simplify slides and speaking points.", "Mention proactively in the opening."),
    Weakpoint("If the team improvises methodology answers, credibility will drop fast", "Methodology questions punish uncertainty more than UI questions do.", "A shaky answer can overshadow strong code.", "High", "Critical", "Critical", "Use narrow, memorized, defensible statements only.", "Rehearse methodology responses separately from the demo.", "Internal prep only."),
]


SAFE_INCOMPLETE_ANSWERS = OrderedDict(
    [
        ("AI mentor not fully adaptive", "Within the current capstone scope, the AI mentor is implemented as a grounded support tool that uses visible class evidence and guided prompts. We describe it as context-aware and assistive, not as a fully adaptive intelligence that independently personalizes everything."),
        ("PDF extraction reliability", "The extraction workflow is implemented and useful for accelerating teacher preparation, but it is still treated as teacher-reviewed AI assistance. For complex or poorly formatted PDFs, human review remains part of the safe workflow."),
        ("RAG not perfect", "The RAG pipeline is implemented through pgvector-backed retrieval and source filtering, but retrieval quality still depends on the quality and availability of indexed class material. That is why we do not claim the AI is infallible."),
        ("Mobile app partial", "The mobile application is implemented in prototype scope with strongest support on student workflows and meaningful but not fully parity-proven role expansion beyond that. The web platform remains the primary surface for complete administration."),
        ("Analytics not fully real-time", "Analytics and reporting workflows are implemented, but not every chart should be described as continuously real-time in the strongest sense. Some outputs are best defended as updated data views based on the available records and generated snapshots."),
        ("Dashboards using seed data", "Several dashboard and reporting surfaces are best demonstrated through controlled test data so the workflow can be shown clearly and consistently. That does not make the feature fake; it simply means classroom-scale live data collection is still limited."),
        ("System not deployed to a real school server", "The current capstone demonstrates a deployable prototype architecture and a pilot-ready workflow, but not a finalized school-wide production deployment. We present it as technically ready for controlled pilot preparation, not as already institutionally rolled out."),
        ("No full external API integration", "The current scope focuses on internal LMS, LXP, and AI-service integration inside the Nexora architecture. External third-party integrations were intentionally excluded so the team could validate the core intervention workflow first."),
        ("No offline mode", "Offline functionality is outside the current scope because the project prioritizes synchronized records, real-time role access, and AI-assisted workflows that depend on connected services. That exclusion is deliberate rather than accidental."),
        ("Hardware requirements too high", "The full AI-enabled stack does require stronger infrastructure than a simple web portal, especially for smoother local inference. That is why the safest deployment answer is a staged or pilot-oriented rollout, not an assumption that every school setup can host the full AI stack immediately."),
        ("Teacher override not implemented everywhere", "Teacher judgment remains central to the workflow even where automated risk detection exists. If a specific override behavior is not yet implemented in the exact form the panel imagines, we describe that as a reasonable enhancement rather than pretending it already exists."),
        ("Audit logs partial", "Audit logging is implemented for important mutations and accountability flows, but we should avoid claiming exhaustive enterprise-grade monitoring of every possible system event. The current scope focuses on meaningful administrative and academic action traceability."),
        ("No large-scale user testing", "The present study is stronger as a system-design and prototype-validation capstone than as a large-scale educational impact study. We openly state that broader user testing remains an important next step."),
        ("No real student deployment yet", "The system is designed for the target school context and demonstrated using controlled development data and prototype workflows. Actual live student deployment remains a separate institutional rollout step.")
    ]
)


DANGEROUS_CLAIMS = OrderedDict(
    [
        ("The AI is fully accurate.", "The AI is grounded and assistive, but teacher and student users should still review outputs critically because AI responses are not perfect."),
        ("The system fully replaces teacher intervention.", "The system supports teachers by surfacing learning gaps and assisting remedial planning, but teacher judgment remains central."),
        ("The platform is ready for full production deployment.", "The platform is best described as a technically mature prototype that is strongest for controlled pilot preparation."),
        ("The AI never hallucinates.", "The system uses grounding and policy controls to reduce hallucination risk, but it does not claim zero-error AI behavior."),
        ("The system supports all subjects completely.", "The system is designed for broad secondary-school use, but the safest defense scope is the current implemented and demonstrated coverage."),
        ("The platform is already proven effective.", "The current study demonstrates implementation feasibility and workflow integration; larger-scale outcome validation remains future work."),
        ("Mobile fully matches the web app.", "Mobile access is implemented, but feature parity differs by role and workflow."),
        ("The threshold of 74% is universally correct.", "The current system uses 74% as its configurable mastery cutoff within project scope and it should still be validated against school policy."),
        ("The AI can grade everything safely.", "The AI assists with support and generation workflows, while official academic records and final grading remain outside direct AI authority."),
        ("The school can deploy this with no special maintenance.", "A controlled pilot is a more honest deployment claim, especially for the AI-enabled stack."),
        ("Our charts already prove long-term learning improvement.", "The analytics and reports demonstrate workflow visibility, but not full long-term educational impact by themselves."),
        ("The system is fully compliant with every privacy requirement.", "The system implements concrete security and privacy-aware controls, but institutional compliance still depends on policy, governance, and rollout review."),
    ]
)


LAST_MINUTE_FIX_PLAN = OrderedDict(
    [
        ("Must fix before defense", [
            "Restore backend runtime and verify `http://localhost:3000/api/health/live` and `http://localhost:3000/api/health/ready` on the actual defense machine.",
            "Correct the title to include the missing word `System`.",
            "Replace or narrow all remaining paper lines that still imply `all subjects and grade levels` have already been fully validated.",
            "Prepare seeded accounts and data that visibly show the 74% threshold, an intervention case, teacher class management, student assessment history, and at least one report/export path.",
            "Rehearse one honest answer for AI limitations, one for methodology limitations, and one for deployment limitations."
        ]),
        ("Should fix before defense", [
            "Prepare one known-good PDF extraction file and screenshots of a prior successful extraction and apply flow.",
            "Seed at least a few evaluation records if the team intends to mention evaluation workflows.",
            "Verify teacher mobile on the exact device to be used, or remove it from the live plan.",
            "Create a one-page evidence cheat sheet with key repo file paths and key paper corrections.",
            "Standardize naming so the team consistently says `Nexora`, `JA/JAKIPIR mentor`, `LXP intervention`, and `74% threshold`."
        ]),
        ("Nice to fix", [
            "Tighten literature wording that sounds dramatic or overly absolute.",
            "Add one stronger privacy/governance paragraph for minors and AI usage.",
            "Trim any slide content that overexplains observability or infrastructure before the core thesis story is clear.",
            "Warm the Ollama models and cache key pages/screens before the panel begins."
        ]),
        ("Can be explained as limitation", [
            "No offline mode.",
            "Admin mobile is not a main supported workflow.",
            "Large-scale outcome validation is not yet complete.",
            "The AI is grounded and useful, but not claimed as perfect."
        ]),
    ]
)


EXECUTIVE_SUMMARY_BULLETS = [
    "The codebase is stronger than the average capstone and implements real LMS, intervention, reporting, and AI-service boundaries.",
    "The main defense risk is not lack of features; it is mismatch between broad paper wording, uneven evaluation evidence, and live demo fragility.",
    "The safest thesis posture is: prototype-scope, school-specific, intervention-focused, teacher-guided, and honest about AI and deployment limits.",
]


TOP_PANEL_CONCERNS = [
    "Why 74% and how strongly is it justified?",
    "Is this really an LXP or just an LMS with remedial gating?",
    "Can the AI hallucinate or mislead minors?",
    "Do you have real evaluation evidence or mainly implementation evidence?",
    "Does the paper overclaim scope across subjects, grade levels, or deployment readiness?",
    "Is mobile parity actually real across roles?",
    "What happens if the AI or backend fails during the demo?",
    "How do you protect student data and AI chat content?",
    "Can the school realistically host and maintain the AI-enabled stack?",
    "Are AI outputs separated from official records and grading?"
]


TOP_10_WEAKPOINTS = [
    "Backend currently not reachable on port 3000 in this audit run.",
    "74% threshold is technically consistent but methodologically vulnerable.",
    "LXP terminology is defendable only when narrowed to intervention-focused LXP features.",
    "Paper still contains some overbroad school-wide wording.",
    "Implementation evidence is stronger than evaluation evidence.",
    "Admin mobile remains limited and should not be demoed as full parity.",
    "Teacher mobile should be treated as prototype-scope unless fully rehearsed.",
    "AI extraction and JAKIPIR are useful but must be framed as assistive and review-aware.",
    "Charts and reports need good seeded data to look credible.",
    "The title still needs a grammar fix."
]


SAFE_STRATEGY = [
    "Lead with stable web flows, not AI-heavy or parity-heavy flows.",
    "Frame AI as grounded support with teacher oversight, not as automatic truth.",
    "Treat the study as implementation- and workflow-validation first, outcome-validation second.",
    "If challenged on breadth, narrow to Grades 7 to 10 and controlled pilot scope immediately.",
    "Never promise full production readiness, full mobile parity, or full AI accuracy."
]


SYSTEM_AND_PAPER_CLAIMS = {
    "paper_claims": [
        "Nexora is a web and mobile LMS with LXP features for targeted intervention.",
        "Students below a 74% threshold get remedial access.",
        "JAKIPIR AI mentor, RAG, PDF extraction, quiz drafting, and AI-assisted remediation are part of the system.",
        "The stack includes Next.js, React, NestJS, PostgreSQL with pgvector, Redis, BullMQ, FastAPI, Ollama, Swagger, and observability tooling.",
        "The project is school-specific to Gat Andres Bonifacio High School."
    ],
    "concept_claims": [
        "The concept paper frames the LXP as an intervention component rather than general student use.",
        "It also positions AI as teacher-guided instructional support rather than teacher replacement.",
        "The concept scope still leans broad, but its language is more modest than the most ambitious parts of the paper."
    ],
    "repo_reality": [
        "The repo contains substantive backend, frontend, ai-service, and mobile code rather than placeholder shells.",
        "The 74% threshold exists in code and schema.",
        "RAG, pgvector retrieval, extraction flows, quiz generation, evaluation schema, audit logs, and role-based routes are real.",
        "Mobile supports real student workflows and teacher navigation, while admin mobile remains limited.",
        "Current runtime during this audit: frontend reachable, ai-service ready, backend unreachable."
    ],
    "mismatch_summary": [
        "The repo is strong enough to support the thesis, but the paper still needs narrower wording in some places.",
        "Implementation maturity does not automatically equal measured educational effectiveness.",
        "Demo-day readiness depends more on environment stability and seeded data than on adding more features."
    ],
}


COMMANDS_RUN = [
    "Read AGENTS, subsystem AGENTS, repo README, package manifests, env examples, docker-compose, concept paper text export, chapter docs, prior audit docs, and extracted paper text.",
    "Extracted text from the provided PDF into artifacts/may03_chapters1_4_extracted.txt with pypdf.",
    "Searched repo and paper text for threshold, RAG, pgvector, AI, OTP, mobile, evaluation, privacy, and deployment evidence.",
    "Checked runtime with localhost probes: frontend 3001 responded 200, ai-service 8000/ready responded healthy, backend 3000 was unreachable.",
]


ASSUMPTIONS = [
    "Because backend port 3000 was not reachable during this run, full end-to-end LMS runtime verification was not possible.",
    "Static code, current paper text, repo documentation, and current partial runtime checks were combined for the final defense package.",
    "Where older repo audits conflicted with current code, current code and current extracted paper text were treated as authoritative.",
    "This package is optimized for hostile-but-fair defense preparation, so it intentionally highlights vulnerabilities rather than flattering the project."
]


UNVERIFIED_CLAIMS = [
    "No new classroom-outcome statistics were generated in this run.",
    "Teacher mobile and all report/export flows were not fully re-walked live in the current session.",
    "The exact live seed-data counts in the current database were not re-queried because backend runtime was unavailable.",
]


def flatten_questions() -> list[dict[str, object]]:
    all_questions: list[dict[str, object]] = []
    for category, items in QUESTION_BANK.items():
        for item in items:
            row = render_question(item)
            row["category"] = category
            all_questions.append(row)
    return all_questions


def validate_counts() -> int:
    actual = {
        "general": len(GENERAL_QUESTIONS),
        "technical": len(TECHNICAL_QUESTIONS),
        "research": len(RESEARCH_QUESTIONS),
        "pedagogy": len(PEDAGOGY_QUESTIONS),
        "ai": len(AI_QUESTIONS),
        "privacy": len(PRIVACY_QUESTIONS),
        "database": len(DATABASE_QUESTIONS),
        "mobile_web": len(MOBILE_WEB_QUESTIONS),
        "demo": len(DEMO_ATTACK_QUESTIONS),
        "deployment": len(DEPLOYMENT_QUESTIONS),
        "trap": len(TRAP_QUESTIONS),
    }
    if actual != COUNTS:
        raise RuntimeError(f"Question counts mismatch. expected={COUNTS} actual={actual}")
    return sum(actual.values())


def md_heading(level: int, text: str) -> str:
    return f"{'#' * level} {text}\n"


def bullet_lines(items: Iterable[str]) -> str:
    return "".join(f"- {item}\n" for item in items)


def write_text(path: Path, text: str) -> None:
    path.write_text(text.strip() + "\n", encoding="utf-8")


def build_question_bank_markdown() -> str:
    parts: list[str] = []
    parts.append(md_heading(1, "Defense Question Bank"))
    parts.append(f"Paper source: `{PAPER_FILE.name}`\n")
    parts.append(f"Repository: `{REPO_NAME}`\n")
    parts.append(f"Generated: `{NOW.isoformat(sep=' ', timespec='seconds')}`\n")
    total = validate_counts()
    parts.append(f"Total questions: **{total}**\n")
    for category, items in QUESTION_BANK.items():
        parts.append(md_heading(2, category))
        for idx, item in enumerate(items, 1):
            row = render_question(item)
            parts.append(md_heading(3, f"{idx}. {row['question']}"))
            parts.append(f"1. Why the panelist might ask it: {row['why']}\n")
            parts.append(f"2. Risk level: {row['risk']}\n")
            parts.append(f"3. Best safe answer: {row['best_safe_answer']}\n")
            parts.append(f"4. Short answer version: {row['short_answer']}\n")
            parts.append(f"5. Long answer version: {row['long_answer']}\n")
            parts.append("6. Evidence to cite from paper/repo/system:\n")
            parts.append(bullet_lines(row["evidence"]))
            parts.append(f"7. What not to say: {row['what_not_to_say']}\n")
            parts.append(f"8. If the feature is incomplete, safest honest answer: {row['incomplete_answer']}\n")
            parts.append("9. Follow-up questions they may ask:\n")
            parts.append(bullet_lines(row["followups"]))
    return "".join(parts)


def build_best_answers_markdown() -> str:
    parts = [md_heading(1, "Defense Best Answers"), "High-risk and critical answer pack.\n"]
    categories_of_interest = {"High", "Critical"}
    for category, items in QUESTION_BANK.items():
        filtered = [render_question(item) for item in items if item.risk in categories_of_interest]
        if not filtered:
            continue
        parts.append(md_heading(2, category))
        for idx, row in enumerate(filtered, 1):
            parts.append(md_heading(3, f"{idx}. {row['question']}"))
            parts.append(f"- Safest phrasing: {row['best_safe_answer']}\n")
            parts.append(f"- Short answer: {row['short_answer']}\n")
            parts.append(f"- Long answer: {row['long_answer']}\n")
            parts.append("- Evidence to show:\n")
            parts.append(bullet_lines(row["evidence"]))
            parts.append(f"- What not to say: {row['what_not_to_say']}\n")
            parts.append(f"- Honest fallback if incomplete: {row['incomplete_answer']}\n")
    parts.append(md_heading(2, "Safe Answers for Known Incomplete Areas"))
    for title, answer in SAFE_INCOMPLETE_ANSWERS.items():
        parts.append(f"- **{title}**: {answer}\n")
    return "".join(parts)


def weakpoint_block(title: str, items: list[Weakpoint]) -> str:
    out = [md_heading(2, title)]
    for idx, item in enumerate(items, 1):
        out.append(md_heading(3, f"{idx}. {item.title}"))
        out.append(f"- Description: {item.description}\n")
        out.append(f"- Why it is dangerous: {item.why_dangerous}\n")
        out.append(f"- Probability panelists notice it: {item.probability}\n")
        out.append(f"- Impact if noticed: {item.impact}\n")
        out.append(f"- Severity: {item.severity}\n")
        out.append(f"- How to defend it verbally: {item.verbal_defense}\n")
        out.append(f"- How to fix it before defense: {item.fix_before_defense}\n")
        out.append(f"- Mention strategy: {item.mention_strategy}\n")
    return "".join(out)


def build_weakpoints_markdown() -> str:
    parts = [md_heading(1, "Top Weakpoints"), f"Readiness score: **{READINESS_SCORE}/100**\n", f"Defense risk level: **{RISK_LEVEL}**\n"]
    parts.append(weakpoint_block("Top 25 System Weakpoints", SYSTEM_WEAKPOINTS))
    parts.append(weakpoint_block("Top 25 Paper Weakpoints", PAPER_WEAKPOINTS))
    parts.append(weakpoint_block("Top 15 Demo Weakpoints", DEMO_WEAKPOINTS))
    parts.append(weakpoint_block("Top 15 AI Weakpoints", AI_WEAKPOINTS))
    parts.append(weakpoint_block("Top 15 Methodology Weakpoints", METHODOLOGY_WEAKPOINTS))
    return "".join(parts)


def build_demo_plan_markdown() -> str:
    parts = [md_heading(1, "Demo Battle Plan")]
    parts.append(md_heading(2, "Recommended Demo Order"))
    parts.append(bullet_lines([
        "Start with backend health confirmation and stable login.",
        "Show web role-based dashboards: admin, teacher, student.",
        "Show teacher class, lesson, and assessment management.",
        "Show student assessment result and the path to intervention logic.",
        "Show a prepared below-74% intervention case.",
        "Show JA/JAKIPIR only if the backend and ai-service are both healthy and warm.",
        "Show one stable report/export path.",
        "End with admin governance features such as audit trail or diagnostics."
    ]))
    parts.append(md_heading(2, "Accounts Needed"))
    parts.append(bullet_lines([
        "1 stable admin account with populated reports and audit visibility.",
        "1 stable teacher account with at least one active class, lessons, assessments, and an intervention case.",
        "1 student account above threshold for normal flow.",
        "1 student account below threshold for intervention and JA/LXP flow.",
        "1 backup student account in case the main intervention seed behaves unexpectedly."
    ]))
    parts.append(md_heading(2, "Seed Data Needed"))
    parts.append(bullet_lines([
        "At least one class with published lessons and a published assessment.",
        "At least one failed attempt below 74% that clearly triggers an intervention case.",
        "At least one successful attempt at or above 74% for contrast.",
        "One prepared PDF and one already-completed extraction for fallback proof.",
        "At least one report/export dataset that does not render empty.",
        "Preferably a few system evaluation entries if the evaluation module will be shown."
    ]))
    parts.append(md_heading(2, "Demo Confidently"))
    parts.append(bullet_lines([
        "Web login and role separation.",
        "Teacher class and assessment workflow.",
        "Student assessment result review.",
        "74% threshold explanation using prepared data.",
        "Audit trail or diagnostics on the admin side."
    ]))
    parts.append(md_heading(2, "Demo Only If Asked"))
    parts.append(bullet_lines([
        "Live PDF extraction from a new file.",
        "Teacher mobile workflow.",
        "Any chart or analytics screen that looks sparse.",
        "AI-heavy remedial generation if Ollama is cold.",
        "Admin mobile."
    ]))
    parts.append(md_heading(2, "Avoid Demoing Live Unless Forced"))
    parts.append(bullet_lines([
        "Admin mobile placeholder routes.",
        "Random unprepared accounts.",
        "Any feature requiring a currently unhealthy backend service.",
        "Long extraction or generation jobs without a fallback artifact.",
        "Screens that depend on data you have not checked the same day."
    ]))
    parts.append(md_heading(2, "Backup Artifacts to Prepare"))
    parts.append(bullet_lines([
        "Screenshots of healthy backend readiness, ai-service readiness, and key dashboards.",
        "Screenshots or a short video of a successful extraction review and apply flow.",
        "Screenshots of JA/JAKIPIR responding with grounded help.",
        "Screenshots of an intervention case opened below threshold.",
        "One export file and one analytics/report page with realistic data."
    ]))
    parts.append(md_heading(2, "If AI Is Slow"))
    parts.append("Say: `The AI-assisted features run asynchronously and are the most compute-sensitive part of the stack. We prepared successful outputs beforehand so we can still show the validated workflow without wasting the panel's time waiting on inference.`\n")
    parts.append(md_heading(2, "If Backend Fails"))
    parts.append("Say: `The current local environment needs recovery, but the implementation boundaries are real. We can still show the verified architecture, the current code paths, and prepared runtime evidence while we explain the exact backend dependency that failed.`\n")
    parts.append(md_heading(2, "If Mobile Fails"))
    parts.append("Say: `The mobile prototype is part of the system scope, but the web flow is the stronger validated demonstration surface today. We can still explain the mobile architecture and show prepared proof of the supported mobile routes.`\n")
    parts.append(md_heading(2, "If the Panelist Asks for a Random Account"))
    parts.append("Say: `For consistency and to avoid exposing incomplete test data, we prepared seeded accounts that represent each validated workflow clearly. We can still explain how the same logic applies across roles and users.`\n")
    parts.append(md_heading(2, "If the Panelist Asks Why a Feature Is Missing"))
    parts.append("Say: `The current capstone scope prioritized the integrated LMS-to-intervention workflow first. That missing area is a recognized extension path, but we did not want to overclaim features beyond what we could implement and defend responsibly.`\n")
    parts.append(md_heading(2, "Emergency Fallback"))
    parts.append(bullet_lines([
        "If backend is unhealthy, stop live clicks immediately and pivot to prepared screenshots, paper-code evidence, and the architecture explanation.",
        "If AI is unhealthy, continue with LMS, threshold logic, reports, audit, and intervention explanation using saved outputs.",
        "If mobile is unstable, keep all role demonstrations on web and show only prepared mobile artifacts.",
        "Never improvise with unknown accounts under panel pressure."
    ]))
    return "".join(parts)


def build_persona_markdown() -> str:
    parts = [md_heading(1, "Panelist Persona Simulation")]
    for persona in PANELIST_PERSONAS:
        parts.append(md_heading(2, persona["name"]))
        parts.append(f"- Focus: {persona['focus']}\n")
        parts.append(f"- Hidden concern behind the questions: {persona['hidden_concern']}\n")
        parts.append(f"- Danger level: {persona['danger_level']}\n")
        parts.append("- Likely questions:\n")
        parts.append(bullet_lines(persona["likely_questions"]))
        parts.append(f"- Best answer approach: {persona['best_answer']}\n")
        parts.append(f"- Answer to avoid: {persona['answer_to_avoid']}\n")
        parts.append("- Evidence we should prepare:\n")
        parts.append(bullet_lines(persona["evidence"]))
        parts.append(f"- Feature/file/page to show if asked: {persona['show_if_asked']}\n")
    return "".join(parts)


def build_dangerous_claims_markdown() -> str:
    parts = [md_heading(1, "Dangerous Claims to Avoid")]
    for claim, safer in DANGEROUS_CLAIMS.items():
        parts.append(f"- Dangerous claim: `{claim}`\n")
        parts.append(f"  Safer version: {safer}\n")
    return "".join(parts)


def build_fix_plan_markdown() -> str:
    parts = [md_heading(1, "Last-Minute Fix Plan")]
    for section, items in LAST_MINUTE_FIX_PLAN.items():
        parts.append(md_heading(2, section))
        parts.append(bullet_lines(items))
    return "".join(parts)


def build_main_report_markdown() -> str:
    total_questions = validate_counts()
    parts = []
    parts.append(md_heading(1, "Thesis Defense Simulation and Weakpoint Report"))
    parts.append(f"System: **Nexora**\n")
    parts.append(f"Paper: **{PAPER_FILE.name}**\n")
    parts.append(f"Repository: **{REPO_NAME}**\n")
    parts.append(f"Generated by: **Codex**\n")
    parts.append(f"Generated at: **{NOW.isoformat(sep=' ', timespec='seconds')}**\n")

    parts.append(md_heading(2, "1. Executive Summary"))
    parts.append(f"- Overall defense readiness score: **{READINESS_SCORE}/100**\n")
    parts.append(f"- Defense risk level: **{RISK_LEVEL}**\n")
    parts.append("- Executive reading:\n")
    parts.append(bullet_lines(EXECUTIVE_SUMMARY_BULLETS))
    parts.append("- Top 10 likely panelist concerns:\n")
    parts.append(bullet_lines(TOP_PANEL_CONCERNS))
    parts.append("- Top 10 weakpoints:\n")
    parts.append(bullet_lines(TOP_10_WEAKPOINTS))
    parts.append("- Safest defense strategy:\n")
    parts.append(bullet_lines(SAFE_STRATEGY))

    parts.append(md_heading(2, "2. System and Paper Claim Summary"))
    parts.append("- What the paper claims:\n")
    parts.append(bullet_lines(SYSTEM_AND_PAPER_CLAIMS["paper_claims"]))
    parts.append("- What the concept paper claims:\n")
    parts.append(bullet_lines(SYSTEM_AND_PAPER_CLAIMS["concept_claims"]))
    parts.append("- What the repository actually implements:\n")
    parts.append(bullet_lines(SYSTEM_AND_PAPER_CLAIMS["repo_reality"]))
    parts.append("- Mismatch summary:\n")
    parts.append(bullet_lines(SYSTEM_AND_PAPER_CLAIMS["mismatch_summary"]))

    parts.append(md_heading(2, "3. Panelist Persona Simulation"))
    for persona in PANELIST_PERSONAS:
        parts.append(md_heading(3, persona["name"]))
        parts.append(f"- Focus: {persona['focus']}\n")
        parts.append(f"- Hidden concern: {persona['hidden_concern']}\n")
        parts.append(f"- Danger level: {persona['danger_level']}\n")
        parts.append("- Likely attack questions:\n")
        parts.append(bullet_lines(persona["likely_questions"]))
        parts.append(f"- Best answer: {persona['best_answer']}\n")
        parts.append(f"- Answer to avoid: {persona['answer_to_avoid']}\n")
        parts.append("- Evidence to prepare:\n")
        parts.append(bullet_lines(persona["evidence"]))
        parts.append(f"- What to show if asked: {persona['show_if_asked']}\n")

    parts.append(md_heading(2, "4. Master Defense Question Bank"))
    parts.append(f"Total generated questions: **{total_questions}**\n")
    for category, items in QUESTION_BANK.items():
        parts.append(md_heading(3, category))
        for idx, item in enumerate(items, 1):
            row = render_question(item)
            parts.append(md_heading(4, f"{idx}. {row['question']}"))
            parts.append(f"- Why: {row['why']}\n")
            parts.append(f"- Risk: {row['risk']}\n")
            parts.append(f"- Best safe answer: {row['best_safe_answer']}\n")
            parts.append(f"- Short answer: {row['short_answer']}\n")
            parts.append(f"- Long answer: {row['long_answer']}\n")
            parts.append("- Evidence:\n")
            parts.append(bullet_lines(row["evidence"]))
            parts.append(f"- What not to say: {row['what_not_to_say']}\n")
            parts.append(f"- Honest fallback if incomplete: {row['incomplete_answer']}\n")
            parts.append("- Follow-up questions:\n")
            parts.append(bullet_lines(row["followups"]))

    parts.append(md_heading(2, "5. Best Answers"))
    for category, items in QUESTION_BANK.items():
        high_risk_rows = [render_question(item) for item in items if item.risk in {"High", "Critical"}]
        if not high_risk_rows:
            continue
        parts.append(md_heading(3, category))
        for row in high_risk_rows[:8]:
            parts.append(md_heading(4, row["question"]))
            parts.append(f"- Short answer: {row['short_answer']}\n")
            parts.append(f"- Long answer: {row['long_answer']}\n")
            parts.append(f"- Safest phrasing: {row['best_safe_answer']}\n")
            parts.append("- Evidence to show:\n")
            parts.append(bullet_lines(row["evidence"]))
            parts.append(f"- What not to say: {row['what_not_to_say']}\n")

    parts.append(md_heading(2, "6. Top System Weakpoints"))
    parts.append(weakpoint_block("System Weakpoints", SYSTEM_WEAKPOINTS))
    parts.append(md_heading(2, "7. Top Paper Weakpoints"))
    parts.append(weakpoint_block("Paper Weakpoints", PAPER_WEAKPOINTS))
    parts.append(md_heading(2, "8. Top AI Weakpoints"))
    parts.append(weakpoint_block("AI Weakpoints", AI_WEAKPOINTS))
    parts.append(md_heading(2, "9. Top Methodology Weakpoints"))
    parts.append(weakpoint_block("Methodology Weakpoints", METHODOLOGY_WEAKPOINTS))

    parts.append(md_heading(2, "10. Demo Battle Plan"))
    parts.append(build_demo_plan_markdown())

    parts.append(md_heading(2, "11. Dangerous Claims to Avoid"))
    for claim, safer in DANGEROUS_CLAIMS.items():
        parts.append(f"- `{claim}`\n")
        parts.append(f"  Safer version: {safer}\n")

    parts.append(md_heading(2, "12. Last-Minute Fix Plan"))
    for section, items in LAST_MINUTE_FIX_PLAN.items():
        parts.append(md_heading(3, section))
        parts.append(bullet_lines(items))

    parts.append(md_heading(2, "13. Final Defense Strategy"))
    parts.append("- What to emphasize:\n")
    parts.append(bullet_lines([
        "Integrated LMS-to-intervention workflow.",
        "Real code across backend, frontend, ai-service, and mobile.",
        "74% threshold implementation consistency.",
        "Teacher-guided AI support, not AI replacement.",
        "Role-based governance and auditability."
    ]))
    parts.append("- What to downplay:\n")
    parts.append(bullet_lines([
        "Claims of full production readiness.",
        "Claims of full mobile parity.",
        "Claims of fully adaptive or perfectly accurate AI.",
        "Any promise of school-wide proven impact not backed by evaluation data."
    ]))
    parts.append("- What to admit honestly:\n")
    parts.append(bullet_lines([
        "The study is stronger in implementation than in large-scale effectiveness proof.",
        "The AI is grounded and useful but not infallible.",
        "Deployment beyond controlled pilot conditions still needs operational validation.",
        "Some mobile and analytics areas remain weaker than the core web workflows."
    ]))
    parts.append("- What to demonstrate first:\n")
    parts.append(bullet_lines([
        "Stable web login and dashboards.",
        "Teacher class and assessment workflow.",
        "Prepared intervention trigger and student follow-up path.",
        "One governance surface such as audit or diagnostics."
    ]))
    parts.append("- Evidence to prepare:\n")
    parts.append(bullet_lines([
        "Health-check screenshots.",
        "Prepared seeded account credentials.",
        "Extraction/AI screenshots and one short backup video.",
        "A one-page repo evidence cheat sheet.",
        "A correction list for the paper."
    ]))

    parts.append(md_heading(2, "14. Appendix"))
    parts.append("- Files inspected:\n")
    parts.append(bullet_lines([
        str(PAPER_FILE),
        "Concept paper.txt",
        "README.md",
        "docker-compose.yml",
        "backend/package.json",
        "next-frontend/package.json",
        "test-mobile/package.json",
        "ai-service/README.md",
        "backend/src/modules/lxp/lxp.service.ts",
        "backend/src/main.ts",
        "backend/src/app.module.ts",
        "backend/src/drizzle/schema/performance.schema.ts",
        "backend/src/drizzle/schema/lxp.schema.ts",
        "backend/src/drizzle/schema/ai-mentor.schema.ts",
        "backend/src/modules/otp/otp.service.ts",
        "backend/src/modules/auth/auth.controller.ts",
        "test-mobile/src/navigation/types.ts",
        "test-mobile/src/navigation/AppNavigator.tsx",
        "test-mobile/src/screens/LoginScreen.tsx",
        "test-mobile/src/screens/VerifyEmailScreen.tsx",
        "test-mobile/src/screens/RoleWorkspaceScreen.tsx",
        "ai-service/app/retrieval_service.py",
        "paper_claims_extracted.md",
        "implementation_truth_table.md",
        "repo_feature_inventory.md",
        "docs/system-audit/concept-paper-current-state-audit-2026-04-06.md",
        "chapter4_diagram_audit.md",
        "panelist_risk_checklist.md",
        "artifacts/may03_chapters1_4_extracted.txt",
    ]))
    parts.append("- Commands run:\n")
    parts.append(bullet_lines(COMMANDS_RUN))
    parts.append("- Assumptions made:\n")
    parts.append(bullet_lines(ASSUMPTIONS))
    parts.append("- Setup blockers:\n")
    parts.append(bullet_lines([
        "Backend runtime was unavailable on port 3000 during this run.",
        "Because backend was down, full cross-role live runtime verification was not possible in this session.",
    ]))
    parts.append("- Unverified claims:\n")
    parts.append(bullet_lines(UNVERIFIED_CLAIMS))
    return "".join(parts)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_doc_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    if level == 1:
        p.style = "Heading 1"
    elif level == 2:
        p.style = "Heading 2"
    elif level == 3:
        p.style = "Heading 3"
    else:
        p.style = "Heading 4"
    run = p.add_run(text)
    if level <= 2:
        run.font.color.rgb = RGBColor(0x12, 0x27, 0x4A)


def add_doc_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_doc_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_kv(doc: Document, key: str, value: str) -> None:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{key}: ")
    r1.bold = True
    p.add_run(value)


def style_document(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.name = "Aptos Display"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.bold = True
    styles["Heading 3"].font.name = "Aptos"
    styles["Heading 3"].font.size = Pt(11)
    styles["Heading 3"].font.bold = True
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Thesis Defense Simulation and Weakpoint Report")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x12, 0x27, 0x4A)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("System: Nexora").font.size = Pt(12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Paper: {PAPER_FILE.name}").font.size = Pt(12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Repository: {REPO_NAME}").font.size = Pt(12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Generated by: Codex").font.size = Pt(12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Date/time generated: {NOW.strftime('%Y-%m-%d %H:%M:%S')}").font.size = Pt(12)
    doc.add_page_break()


def add_score_table(doc: Document) -> None:
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Defense Readiness Score"
    table.cell(0, 1).text = f"{READINESS_SCORE}/100"
    table.cell(1, 0).text = "Defense Risk Level"
    table.cell(1, 1).text = RISK_LEVEL
    for row in table.rows:
        set_cell_shading(row.cells[0], "EAF0F8")
        row.cells[0].paragraphs[0].runs[0].bold = True


def build_docx() -> Path:
    doc = Document()
    style_document(doc)
    add_title_page(doc)

    add_doc_heading(doc, "1. Executive Summary", 1)
    add_score_table(doc)
    add_doc_bullets(doc, EXECUTIVE_SUMMARY_BULLETS)
    add_doc_heading(doc, "Top 10 Likely Panelist Concerns", 2)
    add_doc_bullets(doc, TOP_PANEL_CONCERNS)
    add_doc_heading(doc, "Top 10 Weakpoints", 2)
    add_doc_bullets(doc, TOP_10_WEAKPOINTS)
    add_doc_heading(doc, "Safest Defense Strategy", 2)
    add_doc_bullets(doc, SAFE_STRATEGY)

    add_doc_heading(doc, "2. System and Paper Claim Summary", 1)
    add_doc_heading(doc, "What the Paper Claims", 2)
    add_doc_bullets(doc, SYSTEM_AND_PAPER_CLAIMS["paper_claims"])
    add_doc_heading(doc, "What the System Actually Implements", 2)
    add_doc_bullets(doc, SYSTEM_AND_PAPER_CLAIMS["repo_reality"])
    add_doc_heading(doc, "Mismatch Summary", 2)
    add_doc_bullets(doc, SYSTEM_AND_PAPER_CLAIMS["mismatch_summary"])

    add_doc_heading(doc, "3. Panelist Persona Simulation", 1)
    for persona in PANELIST_PERSONAS:
        add_doc_heading(doc, persona["name"], 2)
        add_kv(doc, "Focus", persona["focus"])
        add_kv(doc, "Hidden concern", persona["hidden_concern"])
        add_kv(doc, "Danger level", persona["danger_level"])
        add_doc_heading(doc, "Likely Questions", 3)
        add_doc_bullets(doc, persona["likely_questions"])
        add_kv(doc, "Best answer", persona["best_answer"])
        add_kv(doc, "Answer to avoid", persona["answer_to_avoid"])
        add_doc_heading(doc, "Evidence to Prepare", 3)
        add_doc_bullets(doc, persona["evidence"])
        add_kv(doc, "Show if asked", persona["show_if_asked"])

    add_doc_heading(doc, "4. Master Defense Question Bank", 1)
    for category, items in QUESTION_BANK.items():
        add_doc_heading(doc, category, 2)
        for idx, item in enumerate(items, 1):
            row = render_question(item)
            add_doc_heading(doc, f"{idx}. {row['question']}", 3)
            add_kv(doc, "Why the panelist might ask it", str(row["why"]))
            add_kv(doc, "Risk level", str(row["risk"]))
            add_kv(doc, "Best safe answer", str(row["best_safe_answer"]))
            add_kv(doc, "Short answer", str(row["short_answer"]))
            add_kv(doc, "Long answer", str(row["long_answer"]))
            add_doc_heading(doc, "Evidence", 4)
            add_doc_bullets(doc, row["evidence"])
            add_kv(doc, "What not to say", str(row["what_not_to_say"]))
            add_kv(doc, "If incomplete, safest honest answer", str(row["incomplete_answer"]))
            add_doc_heading(doc, "Likely follow-ups", 4)
            add_doc_bullets(doc, row["followups"])

    add_doc_heading(doc, "5. Best Answers", 1)
    for category, items in QUESTION_BANK.items():
        filtered = [render_question(item) for item in items if item.risk in {"High", "Critical"}]
        if not filtered:
            continue
        add_doc_heading(doc, category, 2)
        for row in filtered[:8]:
            add_doc_heading(doc, str(row["question"]), 3)
            add_kv(doc, "Short answer", str(row["short_answer"]))
            add_kv(doc, "Long answer", str(row["long_answer"]))
            add_kv(doc, "Safest phrasing", str(row["best_safe_answer"]))
            add_doc_heading(doc, "Evidence to show", 4)
            add_doc_bullets(doc, row["evidence"])
            add_kv(doc, "What not to say", str(row["what_not_to_say"]))

    for section_title, block_title, items in [
        ("6. Top System Weakpoints", "System Weakpoints", SYSTEM_WEAKPOINTS),
        ("7. Top Paper Weakpoints", "Paper Weakpoints", PAPER_WEAKPOINTS),
        ("8. Top AI Weakpoints", "AI Weakpoints", AI_WEAKPOINTS),
        ("9. Top Methodology Weakpoints", "Methodology Weakpoints", METHODOLOGY_WEAKPOINTS),
    ]:
        add_doc_heading(doc, section_title, 1)
        add_doc_heading(doc, block_title, 2)
        for idx, item in enumerate(items, 1):
            add_doc_heading(doc, f"{idx}. {item.title}", 3)
            add_kv(doc, "Description", item.description)
            add_kv(doc, "Why it is dangerous", item.why_dangerous)
            add_kv(doc, "Probability panelists notice it", item.probability)
            add_kv(doc, "Impact if noticed", item.impact)
            add_kv(doc, "Severity", item.severity)
            add_kv(doc, "How to defend it verbally", item.verbal_defense)
            add_kv(doc, "How to fix it before defense", item.fix_before_defense)
            add_kv(doc, "Mention strategy", item.mention_strategy)

    add_doc_heading(doc, "10. Demo Battle Plan", 1)
    for title, items in [
        ("Recommended Demo Order", [
            "Start with health checks and stable login.",
            "Show role-based web dashboards.",
            "Show teacher class, lesson, and assessment management.",
            "Show student assessment result and intervention trigger path.",
            "Show JA/JAKIPIR only if backend and ai-service are both healthy.",
            "End with admin audit, diagnostics, or reports."
        ]),
        ("Accounts Needed", [
            "1 admin account",
            "1 teacher account with a prepared class",
            "1 above-threshold student account",
            "1 below-threshold student account",
            "1 backup student account"
        ]),
        ("Features to Demo Confidently", [
            "Web login and role dashboards",
            "Teacher class workflow",
            "Student assessment result flow",
            "Threshold explanation",
            "Audit or diagnostics"
        ]),
        ("Features to Demo Only If Asked", [
            "Live PDF extraction",
            "Teacher mobile",
            "Sparse analytics pages",
            "Admin mobile",
            "Long AI generation jobs"
        ]),
        ("Emergency Fallback", [
            "Pivot to screenshots or video if runtime breaks.",
            "Keep AI as optional rather than essential to proving the thesis.",
            "Use web as the main demonstration surface."
        ]),
    ]:
        add_doc_heading(doc, title, 2)
        add_doc_bullets(doc, items)
    add_doc_heading(doc, "Contingency Lines", 2)
    add_kv(doc, "If AI is slow", "The AI-assisted features run asynchronously, so we prepared verified outputs to keep the defense efficient while still showing the real workflow.")
    add_kv(doc, "If backend fails", "The repository and architecture remain real, but the current local runtime needs recovery. We can still present verified implementation evidence and prepared artifacts.")
    add_kv(doc, "If mobile fails", "The web system is the stronger primary demonstration surface today, and the mobile prototype remains part of the broader system scope.")

    add_doc_heading(doc, "11. Dangerous Claims to Avoid", 1)
    for claim, safer in DANGEROUS_CLAIMS.items():
        p = doc.add_paragraph()
        r = p.add_run("Dangerous claim: ")
        r.bold = True
        p.add_run(claim)
        p = doc.add_paragraph()
        r = p.add_run("Safer version: ")
        r.bold = True
        p.add_run(safer)

    add_doc_heading(doc, "12. Last-Minute Fix Plan", 1)
    for section, items in LAST_MINUTE_FIX_PLAN.items():
        add_doc_heading(doc, section, 2)
        add_doc_bullets(doc, items)

    add_doc_heading(doc, "13. Final Defense Strategy", 1)
    add_doc_heading(doc, "What to Emphasize", 2)
    add_doc_bullets(doc, [
        "Integrated LMS-to-intervention workflow",
        "Real backend, frontend, ai-service, and mobile code",
        "Teacher-guided AI support",
        "74% threshold consistency in code",
        "Role-based governance and auditability",
    ])
    add_doc_heading(doc, "What to Downplay", 2)
    add_doc_bullets(doc, [
        "Full production-readiness claims",
        "Full mobile parity claims",
        "Perfect AI accuracy claims",
        "Any unsupported outcome-improvement claims",
    ])
    add_doc_heading(doc, "What to Admit Honestly", 2)
    add_doc_bullets(doc, [
        "Implementation readiness is stronger than large-scale impact validation.",
        "The AI is useful and grounded, but not perfect.",
        "Deployment still needs pilot-stage operational validation.",
        "Some mobile and analytics areas are weaker than the core web LMS flow.",
    ])

    add_doc_heading(doc, "14. Appendix", 1)
    add_doc_heading(doc, "Commands Run", 2)
    add_doc_bullets(doc, COMMANDS_RUN)
    add_doc_heading(doc, "Assumptions Made", 2)
    add_doc_bullets(doc, ASSUMPTIONS)
    add_doc_heading(doc, "Unverified Claims", 2)
    add_doc_bullets(doc, UNVERIFIED_CLAIMS)

    output = ROOT / "THESIS_DEFENSE_SIMULATION_AND_WEAKPOINT_REPORT.docx"
    doc.save(output)
    return output


def generate_all() -> dict[str, Path]:
    files = {
        "defense_question_bank.md": ROOT / "defense_question_bank.md",
        "defense_best_answers.md": ROOT / "defense_best_answers.md",
        "top_weakpoints.md": ROOT / "top_weakpoints.md",
        "demo_battle_plan.md": ROOT / "demo_battle_plan.md",
        "panelist_persona_simulation.md": ROOT / "panelist_persona_simulation.md",
        "dangerous_claims_to_avoid.md": ROOT / "dangerous_claims_to_avoid.md",
        "last_minute_fix_plan.md": ROOT / "last_minute_fix_plan.md",
    }
    write_text(files["defense_question_bank.md"], build_question_bank_markdown())
    write_text(files["defense_best_answers.md"], build_best_answers_markdown())
    write_text(files["top_weakpoints.md"], build_weakpoints_markdown())
    write_text(files["demo_battle_plan.md"], build_demo_plan_markdown())
    write_text(files["panelist_persona_simulation.md"], build_persona_markdown())
    write_text(files["dangerous_claims_to_avoid.md"], build_dangerous_claims_markdown())
    write_text(files["last_minute_fix_plan.md"], build_fix_plan_markdown())
    docx_path = build_docx()
    files["THESIS_DEFENSE_SIMULATION_AND_WEAKPOINT_REPORT.docx"] = docx_path
    main_report_md = ROOT / "artifacts" / "thesis_defense_simulation_and_weakpoint_report.md"
    write_text(main_report_md, build_main_report_markdown())
    files["artifacts/thesis_defense_simulation_and_weakpoint_report.md"] = main_report_md
    return files


if __name__ == "__main__":
    total = validate_counts()
    files = generate_all()
    print(f"generated_files={len(files)}")
    print(f"total_questions={total}")
    for name, path in files.items():
        print(f"{name} -> {path}")
